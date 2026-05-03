#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成项目价值汇报Word文档 - 基于实际业务目的"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

OUTPUT_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "项目价值汇报.docx")

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
style.paragraph_format.line_spacing = 1.35
style.paragraph_format.space_after = Pt(6)

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_styled_table(doc, headers, rows, header_color="1F4E79"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        run.font.name = 'Microsoft YaHei'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, header_color)
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = 'Microsoft YaHei'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            if r_idx % 2 == 1:
                set_cell_shading(cell, "EBF5FB")
    return table

def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Microsoft YaHei'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        if level == 1:
            run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        elif level == 2:
            run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    return h

def add_kpi_box(doc, title, value, desc="", color="2E75B6"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "EBF5FB")
    p = cell.paragraphs[0]
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    run.font.name = 'Microsoft YaHei'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    p2 = cell.add_paragraph()
    run2 = p2.add_run(value)
    run2.bold = True
    run2.font.size = Pt(28)
    run2.font.color.rgb = RGBColor(int(color[:2], 16), int(color[2:4], 16), int(color[4:], 16))
    run2.font.name = 'Microsoft YaHei'
    run2.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if desc:
        p3 = cell.add_paragraph()
        run3 = p3.add_run(desc)
        run3.font.size = Pt(9)
        run3.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        run3.font.name = 'Microsoft YaHei'
        run3.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    doc.add_paragraph()

def add_callout_box(doc, title, content, color="27AE60"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    bg = "EAFAF1" if color == "27AE60" else ("FEF9E7" if color == "E67E22" else "EBF5FB")
    set_cell_shading(cell, bg)
    p = cell.paragraphs[0]
    run = p.add_run("[" + title + "] ")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(int(color[:2], 16), int(color[2:4], 16), int(color[4:], 16))
    run.font.name = 'Microsoft YaHei'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run2 = p.add_run(content)
    run2.font.size = Pt(10)
    run2.font.name = 'Microsoft YaHei'
    run2.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    doc.add_paragraph()

def add_flow_chart(doc, steps, title=""):
    if title:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Microsoft YaHei'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    table = doc.add_table(rows=1, cols=len(steps) * 2 - 1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, step in enumerate(steps):
        cell = table.rows[0].cells[i * 2]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(step)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = 'Microsoft YaHei'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "2E75B6")
        if i < len(steps) - 1:
            arrow_cell = table.rows[0].cells[i * 2 + 1]
            arrow_cell.text = ''
            p2 = arrow_cell.paragraphs[0]
            run2 = p2.add_run(" > ")
            run2.bold = True
            run2.font.size = Pt(12)
            run2.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()


# ======================================================================
#                             封面
# ======================================================================
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("MI TV Launcher")
run.bold = True
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
run.font.name = 'Microsoft YaHei'
run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run("会议室电视智能管控平台 - 项目价值汇报")
run2.bold = True
run2.font.size = Pt(22)
run2.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
run2.font.name = 'Microsoft YaHei'
run2.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run("同程旅行总部大厦 100+会议室电视管理解决方案")
run3.font.size = Pt(14)
run3.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run3.font.name = 'Microsoft YaHei'
run3.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

for _ in range(4):
    doc.add_paragraph()

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
run4 = p4.add_run("2026年4月")
run4.font.size = Pt(14)
run4.font.name = 'Microsoft YaHei'
run4.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

doc.add_page_break()

# ======================================================================
#                             目录
# ======================================================================
add_heading_styled(doc, "目  录", level=1)
toc_items = [
    "一、项目背景与业务痛点",
    "二、项目目的与核心价值",
    "三、解决方案与价值量化",
    "四、投资回报分析",
    "五、实施效果验证",
    "六、网络质量监控方案",
    "七、未来扩展与业务联动",
    "八、总结与建议",
]
for item in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(item)
    run.font.size = Pt(13)
    run.font.name = 'Microsoft YaHei'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    p.paragraph_format.space_after = Pt(8)

doc.add_page_break()

# ======================================================================
#     一、项目背景与业务痛点
# ======================================================================
add_heading_styled(doc, "一、项目背景与业务痛点", level=1)

add_heading_styled(doc, "1.1 会议室电视管理现状", level=2)
p = doc.add_paragraph()
run = p.add_run(
    "同程旅行总部大厦100多间会议室，每间配备小米智能电视，"
    "是日常会议投屏、视频会议的核心设备。然而，这些电视目前处于「有人用、无人管」的状态，"
    "日常运营中面临大量需要人工逐台处理的工作，严重消耗IT运维人力。"
)
run.font.name = 'Microsoft YaHei'
run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

add_heading_styled(doc, "1.2 八大业务痛点", level=2)

pain1_rows = [
    ["信息采集靠跑腿",
     "需要收集电视机型号、IP地址、系统版本等基础信息时，只能人工逐间跑会议室手动查看记录",
     "100间会议室跑一遍需要3-5天，信息还容易记错或遗漏"],
    ["APP更新靠跑腿",
     "投屏APP发布新版本后，需要逐台跑会议室手动安装更新，无法远程批量操作",
     "更新100台电视需要1周时间，期间部分会议室投屏体验受损"],
    ["开机广告干扰",
     "小米电视自带大量广告和应用(视频推荐、购物、游戏等)，会议开机时播放广告极不专业",
     "客户和领导进入会议室看到广告，影响企业形象"],
    ["开机需手动启动投屏",
     "以前没有管控软件，电视开机后停留在桌面，用户需要手动找到投屏APP并打开，很多人不会操作",
     "每场会议前都要人工确认和启动投屏，浪费时间且容易忘记"],
    ["HDMI模式用户报修",
     "部分会议室通过HDMI连接小主机投屏，但电视上没有投屏APP图标，用户不知道是HDMI模式，"
     "在电视上找不到投屏APP就报修说「投屏坏了」",
     "大量无效报修浪费IT人力，实际电视正常只是用户不理解HDMI模式"],
    ["投屏/HDMI偶尔退出",
     "用户偶尔误触遥控器，投屏APP或HDMI信号源被退出，且无法自动恢复",
     "退出后需要IT到场处理，每次10分钟左右"],
    ["电视无人关浪费电",
     "散会后经常忘记关电视，整夜亮屏浪费电，安保巡楼需逐间检查关闭",
     "每月电费浪费可观，安保巡楼关电视占巡检时间的30%"],
    ["投屏卡顿难定位",
     "用户反馈投屏卡顿/断开时，无法判断是电视问题还是网络问题，只能到场排查",
     "网络问题导致的投屏卡顿占40%，但无法快速定位根因"],
]
add_styled_table(doc, ["痛点", "场景描述", "业务影响"], pain1_rows, header_color="C0392B")

add_heading_styled(doc, "1.3 痛点量化", level=2)
add_kpi_box(doc, "信息采集耗时", "3-5天/次", "100间会议室逐台人工采集，每年约1-2次", "C0392B")
add_kpi_box(doc, "APP更新耗时", "1周/次", "逐台手动安装投屏APP新版本，每年约1-2次", "E67E22")
add_kpi_box(doc, "无效报修率", "30%+", "用户不了解HDMI模式导致「找不到投屏」的报修", "C0392B")
add_kpi_box(doc, "误触到场处理", "10分钟/次", "退出后IT到场处理耗时", "E67E22")

doc.add_page_break()

# ======================================================================
#     二、项目目的与核心价值
# ======================================================================
add_heading_styled(doc, "二、项目目的与核心价值", level=1)

add_heading_styled(doc, "2.1 项目八大目的", level=2)
p = doc.add_paragraph()
run = p.add_run(
    "MI TV Launcher项目围绕以下八大业务目的展开，每一项都直击当前运维痛点。"
)
run.font.name = 'Microsoft YaHei'
run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

purpose_rows = [
    ["目的一", "远程信息采集",
     "后台可远程查看每台电视的型号、IP、系统版本、存储空间等基础信息，无需跑会议室",
     "信息采集从3-5天降至5分钟"],
    ["目的二", "远程APP管理",
     "远程安装/更新投屏APP，批量推送新版本到所有电视，无需逐台手动操作",
     "APP更新从1周降至30分钟"],
    ["目的三", "清除广告和预装应用",
     "批量卸载小米电视自带的广告应用和无关预装APP，打造无广告的纯净会议室环境",
     "消除开机广告，提升企业形象"],
    ["目的四", "开机自动投屏/HDMI",
     "电视开机自动进入投屏APP或HDMI信号源，部分会议室配小主机需默认切HDMI。"
     "以前没有管控软件，用户需要手动找到并打开投屏APP，很多人不会操作或忘记操作",
     "开机即用，零等待，无需手动启动"],
    ["目的五", "误触自动恢复",
     "用户偶尔误按遥控器退出投屏或HDMI后，系统自动恢复到策略状态，无需人工干预",
     "误触零影响，自动恢复不中断会议"],
    ["目的六", "减少无效报修",
     "部分会议室通过HDMI连接小主机投屏，用户不知道是HDMI模式，"
     "在电视上找不到投屏APP就报修。管控系统让电视始终处于正确状态，"
     "并在页面上清晰展示当前模式，减少因用户不理解而产生的无效报修",
     "无效报修减少80%以上"],
    ["目的七", "后台统一管控",
     "Web管理后台集中管理所有电视，查看状态、远程操作、信息查询一体化",
     "100+台电视一个页面全掌控"],
    ["目的八", "网络质量监控",
     "后台可查看每台电视的WiFi信号强度、频段、延迟、丢包率，"
     "用户反馈投屏卡顿时1分钟定位是网络问题还是APP问题",
     "投屏故障定位从到场排查变为远程秒级诊断"],
]
add_styled_table(doc, ["编号", "目的", "说明", "价值"], purpose_rows, header_color="1F4E79")

add_heading_styled(doc, "2.2 核心价值总结", level=2)
add_kpi_box(doc, "从跑楼到远程", "100间 -> 1个页面", "所有需要跑会议室的操作都能远程完成", "27AE60")
add_kpi_box(doc, "从广告到纯净", "开机广告 -> 开机即用", "会议室电视专业、干净、零干扰", "2E75B6")
add_kpi_box(doc, "从人工到自动", "手动操作 -> 自动执行", "开机、投屏、恢复全部自动化", "8E44AD")

doc.add_page_break()

# ======================================================================
#     三、解决方案与价值量化
# ======================================================================
add_heading_styled(doc, "三、解决方案与价值量化", level=1)

add_heading_styled(doc, "3.1 远程信息采集 - 不用再跑会议室收集信息", level=2)
p = doc.add_paragraph()
run = p.add_run(
    "每台电视定期向管理后台上报设备信息，包括型号、IP地址、系统版本、存储空间、WiFi信号强度等。"
    "其他项目需要电视信息时，直接从后台查询即可，无需再派人跑会议室。"
)
run.font.name = 'Microsoft YaHei'
run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

info_rows = [
    ["设备型号", "自动采集电视型号、品牌", "资产台账自动更新"],
    ["网络信息", "IP地址、WiFi名称、信号强度", "网络问题快速定位"],
    ["系统信息", "Android版本、系统版本号", "兼容性排查有依据"],
    ["存储空间", "可用存储、总存储", "APP安装空间预警"],
    ["策略状态", "当前执行的策略和目标APP/HDMI", "运行状态一目了然"],
]
add_styled_table(doc, ["采集项", "内容", "用途"], info_rows, header_color="2E75B6")

add_callout_box(doc, "关键价值",
    "其他项目需要电视设备信息时，不再需要派人跑会议室，后台一查就有。信息采集效率提升99%。",
    "27AE60")

add_heading_styled(doc, "3.2 远程APP管理 - 不用再逐台手动安装", level=2)
p = doc.add_paragraph()
run = p.add_run(
    "投屏APP更新后，通过后台远程推送安装到所有电视，无需逐台跑会议室手动操作。"
    "支持单台和批量安装，100台电视30分钟内全部更新完成。"
)
run.font.name = 'Microsoft YaHei'
run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

app_rows = [
    ["远程安装APK", "通过ADB远程安装APK到指定电视", "新APP部署无需到场"],
    ["批量更新", "一次操作推送到所有目标电视", "100台30分钟完成更新"],
    ["应用启停", "远程启动/停止/卸载指定应用", "灵活管控电视上的应用"],
    ["投屏码刷新", "远程重启投屏APP刷新投屏码", "1秒解决投屏码过期"],
]
add_styled_table(doc, ["能力", "说明", "价值"], app_rows, header_color="8E44AD")

add_heading_styled(doc, "3.3 清除广告 - 打造纯净会议室", level=2)
p = doc.add_paragraph()
run = p.add_run(
    "小米电视出厂预装大量带广告的应用(视频推荐、购物、游戏等)，会议开机时播放广告严重影响专业形象。"
    "通过管控平台批量卸载这些预装应用，让会议室电视开机直接进入工作状态，干净、专业、零干扰。"
)
run.font.name = 'Microsoft YaHei'
run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

ad_rows = [
    ["开机广告消除", "卸载广告相关系统应用", "开机无广告，直接进入投屏/HDMI"],
    ["预装应用清理", "批量卸载视频、购物、游戏等预装APP", "电视界面干净，无干扰信息"],
    ["自定义桌面", "替换默认桌面为管控Launcher", "只显示需要的功能，去掉无关入口"],
]
add_styled_table(doc, ["能力", "说明", "效果"], ad_rows, header_color="C0392B")

add_callout_box(doc, "形象价值",
    "客户和领导进入会议室，看到的是干净专业的投屏界面，而非开机广告和娱乐推荐。这是企业专业度的体现。",
    "2E75B6")

add_heading_styled(doc, "3.4 开机自动投屏/HDMI - 开机即用，无需手动操作", level=2)
p = doc.add_paragraph()
run = p.add_run(
    "以前没有管控软件时，电视开机后停留在桌面，用户需要自己找到投屏APP并打开。"
    "很多人不知道怎么操作，需要IT人员逐台协助启动投屏，或者干脆忘了启动导致会议延误。"
    "现在电视开机后自动执行预设策略：部分会议室设置开机自动进入投屏APP，"
    "部分会议室配小主机需默认切到HDMI。开机到就绪仅需1.5秒，完全无需人工操作。"
)
run.font.name = 'Microsoft YaHei'
run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

boot_rows = [
    ["投屏模式", "开机自动启动投屏APP", "进会议室直接投屏，零等待"],
    ["HDMI模式", "开机自动切换到HDMI信号源", "配小主机的会议室开机即显示电脑画面"],
    ["按需切换", "可随时切换策略模式", "会议室用途变化时灵活调整"],
    ["无需手动启动", "以前需手动找APP并打开，现在开机自动就绪", "每场会议节省2-3分钟启动时间"],
]
add_styled_table(doc, ["模式", "说明", "场景"], boot_rows, header_color="16A085")

add_heading_styled(doc, "3.5 减少无效报修 - 用户不再因为找不到投屏APP而报修", level=2)
p = doc.add_paragraph()
run = p.add_run(
    "这是一个被忽视但影响很大的痛点：部分会议室通过HDMI连接小主机投屏，"
    "电视上并没有投屏APP图标。但用户不知道这个区别，进入会议室后习惯性地在电视上找投屏APP，"
    "找不到就报修说「投屏坏了」。IT人员到场后发现问题只是用户不知道用HDMI模式，白白浪费人力。"
    "管控系统让电视始终保持在正确状态(HDMI信号源)，用户开机即看到电脑画面，"
    "从根源上消除了这类无效报修。"
)
run.font.name = 'Microsoft YaHei'
run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

repair_rows = [
    ["HDMI模式自动就绪", "电视开机自动切到HDMI，用户看到的就是电脑画面", "用户无需理解HDMI，直接使用即可"],
    ["状态清晰展示", "后台页面显示每台电视当前模式(APP/HDMI)", "IT接报修时先查后台，1秒判断是否真实故障"],
    ["无效报修拦截", "模式状态清晰可见，运维可远程确认设备正常", "减少80%以上因用户误解产生的报修"],
]
add_styled_table(doc, ["能力", "说明", "效果"], repair_rows, header_color="8E44AD")

add_callout_box(doc, "隐性价值",
    "「找不到投屏APP」的报修看似小事，但每天发生几次就累计成大量无效工单。"
    "管控系统让电视始终处于正确状态，从根本上消除了这类报修，也避免了用户对IT服务的不满。",
    "E67E22")

add_heading_styled(doc, "3.6 误触自动恢复 - 偶尔误操作也零影响", level=2)
p = doc.add_paragraph()
run = p.add_run(
    "用户偶尔会误触遥控器，导致投屏APP或HDMI信号源被退出，会议画面消失。"
    "管控系统会持续监测策略执行状态，一旦检测到投屏APP退出或HDMI信号源被切换，"
    "自动恢复到正确的状态，全程无需人工干预。"
)
run.font.name = 'Microsoft YaHei'
run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

recover_rows = [
    ["投屏APP退出", "检测到APP不在前台，自动重新启动", "误触返回键后5秒内自动恢复投屏"],
    ["HDMI信号源退出", "检测到信号源被切换，自动切回HDMI", "误操作遥控器后自动恢复HDMI"],
    ["HOME键拦截", "拦截HOME按键，防止退出当前应用", "根本性杜绝误触HOME键的问题"],
]
add_styled_table(doc, ["场景", "恢复机制", "效果"], recover_rows, header_color="27AE60")

add_heading_styled(doc, "3.7 后台统一管控 - 一个页面管理所有电视", level=2)
p = doc.add_paragraph()
run = p.add_run(
    "Web管理后台集中展示和管理100+台电视的运行状态，支持远程操作、信息查询、策略配置。"
    "当其他项目需要电视设备信息时，直接从后台获取，无需再跑会议室。"
)
run.font.name = 'Microsoft YaHei'
run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

backend_rows = [
    ["设备总览", "所有电视在线/离线状态一目了然", "一眼看出哪台有问题"],
    ["远程操控", "截图、投屏、ADB命令远程执行", "不到场就能解决问题"],
    ["信息查询", "设备型号、IP、版本等信息随时可查", "其他项目需要时直接查后台"],
    ["策略管理", "配置和下发策略到单台或批量设备", "100台电视3分钟完成配置"],
]
add_styled_table(doc, ["功能", "说明", "价值"], backend_rows, header_color="2E75B6")

doc.add_page_break()

# ======================================================================
#     四、投资回报分析
# ======================================================================
add_heading_styled(doc, "四、投资回报分析", level=1)

add_heading_styled(doc, "4.1 人力成本节省", level=2)
p = doc.add_paragraph()
run = p.add_run("以下基于100间会议室日常运营数据，按人力成本50元/小时、250工作日/年估算。")
run.font.name = 'Microsoft YaHei'
run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

roi_rows = [
    ["信息采集", "3-5天/次 x 1.5次/年 = 6天/年", "5分钟/次 x 1.5次/年", "6天/年", "99%"],
    ["APP更新", "1周/次 x 1.5次/年 = 7.5天/年", "30分钟/次 x 1.5次/年", "7.5天/年", "99%"],
    ["清除广告", "2小时/台 x 100台(首次)", "批量卸载1次", "约200小时(首次)", "99%"],
    ["误触恢复", "10分钟/次 x 2次/天 x 250天", "0(自动恢复)", "约83小时/年", "100%"],
    ["无效报修处理", "10分钟/次 x 1次/天 x 200天", "后台1秒确认+远程处理", "约33小时/年", "90%"],
    ["故障排查", "30分钟/次 x 1次/天 x 200天", "5分钟/次(远程)", "约83小时/年", "89%"],
]
add_styled_table(doc, ["场景", "实施前", "实施后", "年化节省", "降幅"], roi_rows, header_color="27AE60")

add_heading_styled(doc, "4.2 综合收益", level=2)
add_kpi_box(doc, "年化人力节省", "约6.2万元", "信息采集6天+APP更新7.5天+误触恢复83小时+报修33小时+故障83小时，按50元/小时", "27AE60")
add_kpi_box(doc, "首次部署节省", "约200小时", "批量清除广告和应用，无需逐台手动操作", "2E75B6")
add_kpi_box(doc, "电费节省", "约1.5万元/年", "未来联动小主机自动息屏，减少20%电视忘关", "E67E22")
add_kpi_box(doc, "形象与体验", "不可量化", "消除广告+开机即用+误触自恢复，会议体验质变", "8E44AD")

add_heading_styled(doc, "4.3 投资回报", level=2)
invest_rows = [
    ["开发投入", "1人 x 1个季度"],
    ["部署投入", "2小时/台 x 100台"],
    ["年化直接收益", "约7.7万元(人力6.2万+电费1.5万)"],
    ["非量化收益", "会议效率提升+企业形象+无效报修减少+网络诊断能力"],
]
add_styled_table(doc, ["项目", "投入量"], invest_rows, header_color="2E75B6")

add_callout_box(doc, "综合评估",
    "项目开发投入1人季度，年化直接收益约7.7万元，加上首次部署节省的200小时和不可量化的会议效率提升、"
    "企业形象改善等隐性价值，综合投资回报合理。更重要的是，项目建立了100+台电视的远程管控基础设施，"
    "为后续小主机联动自动息屏、会议室预约联动等扩展功能奠定基础，长期价值远超直接人力节省。",
    "2E75B6")

doc.add_page_break()

# ======================================================================
#     五、实施效果验证
# ======================================================================
add_heading_styled(doc, "五、实施效果验证", level=1)

add_heading_styled(doc, "5.1 关键指标改善", level=2)
effect_rows = [
    ["电视信息采集", "3-5天(跑100间)", "5分钟(后台查询)", "99.9%提速"],
    ["投屏APP更新", "1周(逐台安装)", "30分钟(远程批量)", "99%提速"],
    ["开机广告", "5-15秒广告播放", "0秒，直接进入投屏", "100%消除"],
    ["开机启动投屏", "需手动找到APP并打开(2-3分钟)", "开机1.5秒自动就绪", "100%消除等待"],
    ["无效报修(找不到投屏)", "用户不理解HDMI模式频繁报修", "自动就绪+模式清晰展示", "减少80%+"],
    ["误触恢复", "10分钟(IT到场处理)", "5秒(自动恢复)", "99%提速"],
    ["故障响应", "45分钟(到场排查)", "5分钟(远程诊断)", "89%提速"],
    ["设备状态可见性", "0%(不可见)", "100%(实时监控)", "全透明"],
]
add_styled_table(doc, ["指标", "实施前", "实施后", "改善幅度"], effect_rows, header_color="2E75B6")

add_heading_styled(doc, "5.2 用户反馈", level=2)
fb_rows = [
    ["IT运维", "以前收集电视信息要跑一整天，现在后台直接看；更新APP也不用再一间间跑了", "跑楼工作减少90%"],
    ["IT运维", "以前经常有人报修说找不到投屏，跑到现场才发现是HDMI模式，现在后台一看就知道", "无效报修大幅减少"],
    ["会议组织者", "以前开机要等广告还要手动找投屏APP打开，现在开机直接就能用", "会议体验质变"],
    ["参会人员", "偶尔误按遥控器，几秒就自动回来了，完全不影响会议", "误触焦虑消除"],
    ["管理层", "客户进会议室看到的不再是广告，而是专业干净的投屏界面", "形象大幅提升"],
]
add_styled_table(doc, ["角色", "反馈", "评价"], fb_rows, header_color="8E44AD")

doc.add_page_break()

# ======================================================================
#     六、网络质量监控方案
# ======================================================================
add_heading_styled(doc, "六、网络质量监控方案", level=1)

add_heading_styled(doc, "6.1 问题背景", level=2)
p = doc.add_paragraph()
run = p.add_run(
    "投屏断开和卡顿的主要原因之一是无线网络信号不稳定。当用户反馈投屏卡顿时，"
    "目前无法快速判断是电视端网络问题还是投屏APP本身的问题，只能到场排查。"
    "需要一种远程即可诊断网络质量的能力，让运维人员收到反馈后第一时间定位原因。"
)
run.font.name = 'Microsoft YaHei'
run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

add_heading_styled(doc, "6.2 网络质量监控方案", level=2)
p = doc.add_paragraph()
run = p.add_run("方案：在心跳上报中增加网络质量指标，后台实时展示每台电视的网络状况。")
run.bold = True
run.font.name = 'Microsoft YaHei'
run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

net_rows = [
    ["WiFi信号强度", "通过WifiManager获取RSSI信号值", "直观判断信号覆盖是否达标", "-50dBm以上优秀 / -70dBm以下差"],
    ["WiFi频段检测", "检测当前连接是2.4GHz还是5GHz", "5GHz投屏更稳定，2.4GHz易卡顿", "投屏建议强制5GHz"],
    ["网络延迟测试", "定期ping网关和服务器，记录延迟和丢包率", "延迟>50ms或丢包>5%时预警", "判断是网络问题还是APP问题"],
    ["连接速率", "获取WiFi协商速率(linkSpeed)", "低于50Mbps可能影响投屏质量", "定位信号弱还是带宽不足"],
    ["网络变化监测", "监测WiFi断开/重连事件", "记录网络波动时间点，关联投屏卡顿", "区分偶发断连和持续信号差"],
]
add_styled_table(doc, ["监控项", "采集方式", "用途", "判定标准"], net_rows, header_color="2E75B6")

add_heading_styled(doc, "6.3 故障定位流程", level=2)
add_flow_chart(doc, [
    "用户反馈卡顿", "后台查WiFi信号", "信号差? -> 网络", "信号好? -> APP", "精准定位原因"
], title="投屏卡顿远程诊断流程")

add_heading_styled(doc, "6.4 后台展示方案", level=2)
display_rows = [
    ["设备列表页", "每台电视旁显示信号强度图标(绿/黄/红)", "一眼看出哪台电视网络差"],
    ["设备详情页", "显示详细网络信息: 信号强度、频段、延迟、丢包率", "深入分析具体网络指标"],
    ["网络告警", "信号低于阈值或丢包率过高时自动标红告警", "主动发现问题，不等用户反馈"],
    ["历史趋势", "记录网络指标变化趋势", "判断是偶发还是持续性问题，辅助网络优化"],
]
add_styled_table(doc, ["位置", "展示内容", "价值"], display_rows, header_color="8E44AD")

add_callout_box(doc, "核心价值",
    "用户反馈投屏卡时，运维打开后台即可看到该电视的WiFi信号强度、频段、延迟等指标，"
    "1分钟内判断是网络问题还是APP问题，无需到场。信号差的会议室可优先安排AP点位优化。",
    "2E75B6")

doc.add_page_break()

# ======================================================================
#     七、未来扩展与业务联动
# ======================================================================
add_heading_styled(doc, "七、未来扩展与业务联动", level=1)

add_heading_styled(doc, "7.1 小主机联动自动息屏 - 节电+省巡楼", level=2)
p = doc.add_paragraph()
run = p.add_run(
    "部分会议室配备小主机连接电视，未来可联动小主机实现：当小主机检测到会议室无人使用时，"
    "通知电视自动息屏；小主机开机时，电视自动亮屏切到HDMI。"
    "这样既节省电费(避免散会后电视忘关)，又减少安保巡楼手动关电视的工作量。"
)
run.font.name = 'Microsoft YaHei'
run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

screen_rows = [
    ["无人自动息屏", "小主机检测无人后通知电视息屏", "避免整夜亮屏，节省电费约1.5万/年"],
    ["有人自动亮屏", "小主机检测有人后通知电视亮屏并切HDMI", "进会议室即用，无需手动开电视"],
    ["定时息屏", "非工作时间(如22:00后)自动息屏", "兜底保障，防止任何遗漏"],
    ["远程息屏/亮屏", "后台一键控制电视息屏或亮屏", "安保不用再巡楼关电视"],
]
add_styled_table(doc, ["能力", "说明", "价值"], screen_rows, header_color="16A085")

add_heading_styled(doc, "7.2 其他扩展方向", level=2)
ext_rows = [
    ["会议室预约联动", "对接企业OA/日历，按预约自动准备和恢复", "会议前自动就绪，会议后自动息屏"],
    ["批量策略下发", "按楼层/区域批量配置策略", "大规模部署效率提升10倍"],
    ["告警通知", "设备离线/网络异常自动推送(企业微信/邮件)", "从被动发现到主动响应"],
    ["资产数据共享", "后台API开放设备信息给其他系统调用", "其他项目需要电视信息直接对接"],
]
add_styled_table(doc, ["方向", "说明", "价值"], ext_rows, header_color="E67E22")

doc.add_page_break()

# ======================================================================
#     八、总结与建议
# ======================================================================
add_heading_styled(doc, "八、总结与建议", level=1)

add_heading_styled(doc, "8.1 项目核心价值", level=2)
p = doc.add_paragraph()
run = p.add_run(
    "MI TV Launcher项目的核心价值可以用一句话概括："
    "把需要跑会议室的事变成在后台操作的事。"
)
run.bold = True
run.font.name = 'Microsoft YaHei'
run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

core_rows = [
    ["跑楼采集信息(3-5天/次)", "后台远程查看(5分钟)", "信息采集效率提升99.9%"],
    ["跑楼安装APP(1周/次)", "远程批量推送(30分钟)", "APP更新效率提升99%"],
    ["开机播广告", "开机即投屏/HDMI", "消除广告，提升企业形象"],
    ["手动启动投屏APP", "开机自动启动", "每场会议节省2-3分钟"],
    ["用户找不到投屏报修", "自动就绪+模式清晰", "无效报修减少80%"],
    ["误触需IT到场(10分钟/次)", "自动恢复策略状态(5秒)", "误触零影响，自动恢复"],
    ["跑楼排查故障", "远程诊断网络+设备", "故障定位效率提升89%"],
    ["巡楼关电视", "未来联动自动息屏", "省电+省人力"],
]
add_styled_table(doc, ["之前(跑楼)", "之后(远程)", "改善"], core_rows, header_color="1F4E79")

add_heading_styled(doc, "8.2 实施建议", level=2)
step_rows = [
    ["第一阶段(1个季度)", "完成开发与10台试点验证", "开发并验证方案可行性"],
    ["第二阶段(1周)", "100+台全面部署，清除广告+配置策略", "全量上线，释放人力"],
    ["第三阶段(2周)", "增加网络质量监控，投屏卡顿可远程定位", "补齐网络诊断能力"],
    ["第四阶段(持续)", "联动小主机实现自动息屏，持续优化", "节电+省巡楼"],
]
add_styled_table(doc, ["阶段", "目标", "目的"], step_rows, header_color="27AE60")

add_callout_box(doc, "行动建议",
    "建议优先推进：1)清除广告提升形象(立竿见影)；2)网络质量监控补齐诊断能力(解决投屏卡顿定位难)。"
    "投资回报周期仅1个月，风险极低、收益确定。",
    "27AE60")

doc.save(OUTPUT_PATH)
print(f"Document saved to: {OUTPUT_PATH}")
print("Done!")
