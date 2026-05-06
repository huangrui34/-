"""生成 TV Launcher 问题与解决方案文档 (.docx)"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement

OUT_DIR = os.path.dirname(os.path.abspath(__file__))
os.makedirs(OUT_DIR, exist_ok=True)

# ========== 色彩体系 ==========
COLOR_PRIMARY = "1B3A5C"
COLOR_SECONDARY = "2E75B6"
COLOR_ACCENT = "4A90D9"
COLOR_TEXT = "2D3436"
COLOR_TEXT_LIGHT = "636E72"
COLOR_BORDER = "D5DDE5"
COLOR_BG_GRAY = "F5F6F8"
COLOR_BG_GREEN = "E8F8F5"
COLOR_BG_RED = "FDEDEC"
COLOR_BG_YELLOW = "FEF9E7"
COLOR_BG_BLUE = "EBF5FB"
COLOR_SOLVE_GREEN = "27AE60"
COLOR_FAIL_RED = "E74C3C"
COLOR_WARN_YELLOW = "F39C12"

# ========== 通用样式工具 ==========

def set_cell_shading(cell, color_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_styled_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor.from_string(COLOR_PRIMARY if level == 1 else COLOR_SECONDARY if level == 2 else COLOR_ACCENT)
        run.font.bold = True
    return h

def add_body_text(doc, text, bold=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(color or COLOR_TEXT)
    run.font.bold = bold
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(4)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    return p

def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.size = Pt(9)
    run.font.name = "Consolas"
    run.font.color.rgb = RGBColor.from_string(COLOR_TEXT)
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(4)
    # 添加灰色背景
    pPr = p._p.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{COLOR_BG_GRAY}"/>')
    pPr.append(shd)
    return p

def add_problem_section(doc, problem):
    """添加一个问题段落，problem 是一个字典"""
    # 问题描述标题
    h = doc.add_heading(level=3)
    run = h.add_run(f"问题{problem['id']}：{problem['title']}")
    run.font.color.rgb = RGBColor.from_string(COLOR_FAIL_RED)
    run.font.bold = True

    # 背景/现象
    add_body_text(doc, "【现象】", bold=True)
    add_body_text(doc, problem['symptom'])

    # 原因分析
    add_body_text(doc, "【原因分析】", bold=True)
    add_body_text(doc, problem['cause'])

    # 尝试过的方案
    if problem.get('attempts'):
        add_body_text(doc, "【尝试过的方案】", bold=True)
        for i, attempt in enumerate(problem['attempts'], 1):
            p = doc.add_paragraph()
            run_num = p.add_run(f"  {i}. {attempt['name']}")
            run_num.font.size = Pt(11)
            run_num.font.bold = True
            if attempt.get('success'):
                run_num.font.color.rgb = RGBColor.from_string(COLOR_SOLVE_GREEN)
            else:
                run_num.font.color.rgb = RGBColor.from_string(COLOR_FAIL_RED)
            run_desc = p.add_run(f" — {attempt['desc']}")
            run_desc.font.size = Pt(11)
            run_desc.font.color.rgb = RGBColor.from_string(COLOR_TEXT)
            pf = p.paragraph_format
            pf.space_before = Pt(2)
            pf.space_after = Pt(2)

    # 最终解决方案
    if problem.get('solution'):
        add_body_text(doc, "【解决方案】", bold=True, color=COLOR_SOLVE_GREEN)
        add_body_text(doc, problem['solution'])
        if problem.get('code_example'):
            add_code_block(doc, problem['code_example'])

    # 经验总结
    if problem.get('lesson'):
        add_body_text(doc, "【经验总结】", bold=True, color=COLOR_SECONDARY)
        add_body_text(doc, problem['lesson'])

    doc.add_paragraph()  # 空行


# ========== 问题数据 ==========

PROBLEMS = [
    {
        "id": 1,
        "title": "安卓6电视无法检测HDMI热插拔信号",
        "symptom": (
            "会议室有一台安卓6.0.1的小米电视，当插入或拔出HDMI线时，APP完全收不到任何通知。"
            "而另一台安卓9的同品牌电视则能正常收到HDMI插拔广播。"
            "这导致HDMI自动切换功能在安卓6电视上完全失效。"
        ),
        "cause": (
            "安卓6电视的HDMI检测机制与安卓9不同：\n"
            "• 安卓9：硬件支持HDMI热插拔检测(HPD)，系统会发出 android.intent.action.HDMI_PLUGGED 广播\n"
            "• 安卓6（小米/Amlogic芯片）：没有硬件HPD支持，依赖CEC协议检测HDMI设备。"
            "但普通电脑的HDMI输出不支持CEC，所以插上电脑的HDMI线，电视根本不知道。\n\n"
            "我们尝试了所有可能的检测方式：\n"
            "• /sys/class/switch/hdmirx_hpd/state 文件 → 始终返回0\n"
            "• mbx.hdmiin.switchfull 系统属性 → 始终为false\n"
            "• android.intent.action.HDMI_PLUGGED 广播 → 从不触发\n"
            "• 以上所有方式在没有物理HPD硬件支持的设备上都无法工作"
        ),
        "attempts": [
            {"name": "读取sysfs文件", "desc": "读取/sys/class/switch/hdmirx_hpd/state，始终返回0，无效", "success": False},
            {"name": "读取系统属性", "desc": "读取mbx.hdmiin.switchfull属性，始终为false，无效", "success": False},
            {"name": "监听HDMI广播", "desc": "注册android.intent.action.HDMI_PLUGGED广播接收器，安卓6上从不触发", "success": False},
        ],
        "solution": (
            "安卓6电视在没有硬件HPD支持的情况下，无法通过软件方式检测HDMI热插拔。"
            "这是硬件限制，不是软件bug。\n\n"
            "替代方案：\n"
            "1. 对于安卓9及以上电视 → 使用HDMI_PLUGGED广播自动切换（已实现）\n"
            "2. 对于安卓6电视 → 只能通过后台管理页面手动切换HDMI策略，或按遥控器切换\n"
            "3. 在APP界面中增加明确的提示，告知用户此型号不支持自动HDMI检测"
        ),
        "lesson": (
            "不是所有问题都能通过软件解决。HDMI热插拔检测是硬件功能，"
            "老型号电视如果没有HPD电路，软件无法弥补。"
            "做功能规划时需要先确认硬件能力，再决定软件方案。"
        ),
    },
    {
        "id": 2,
        "title": "安卓6电视无法切换到指定HDMI端口（显示'未知，搜索信号中'）",
        "symptom": (
            "在安卓6电视上，点击切换到HDMI1/HDMI2/HDMI3时，电视画面显示'未知，搜索信号中'，"
            "而不是显示HDMI输入的画面。但实际上HDMI线已经插好了，电脑也在正常输出信号。"
            "通过电视自带设置菜单切换HDMI则完全正常。"
        ),
        "cause": (
            "最初我们使用的切换方式是发送 android.media.tv.action.SETUP_INPUTS Intent，"
            "附带 INPUT_ID 参数指定HDMI端口。这个方式在参考的酒店系统APK（ETV）中是有效的。\n\n"
            "但在安卓6电视上失败，原因是：\n"
            "• SETUP_INPUTS 这个Intent需要有一个Activity来处理（接收并执行切换）\n"
            "• 安卓9电视有系统Activity来处理这个Intent，所以能成功\n"
            "• 安卓6电视没有注册任何Activity来处理这个Intent，导致报错：\n"
            "  'No Activity found to handle Intent'\n\n"
            "备用的 EXTSRC_PLAY 方式虽然能打开小米电视的信号源选择界面，"
            "但它只打开了选择页面，无法自动切换到指定的HDMI端口，所以显示'未知'。"
        ),
        "attempts": [
            {"name": "SETUP_INPUTS Intent", "desc": "发送android.media.tv.action.SETUP_INPUTS + INPUT_ID，报错No Activity found", "success": False},
            {"name": "EXTSRC_PLAY", "desc": "发送小米电视com.xiaomi.mitv.tvplayer.EXTSRC_PLAY，只打开选择界面，显示'未知'", "success": False},
            {"name": "直接启动tvplayer", "desc": "启动com.xiaomi.mitv.tvplayer，同样只打开选择界面", "success": False},
            {"name": "am startservice", "desc": "直接启动Hdmi1InputService，报错需要BIND_TV_INPUT权限", "success": False},
            {"name": "setprop方式", "desc": "设置mbx.hdmiin.switchfull=true，无root权限无法修改", "success": False},
        ],
        "solution": (
            "使用Android TV Input Framework的 TvView API 直接调谐到HDMI输入。\n\n"
            "这是Android TV标准做法：\n"
            "1. 创建一个HdmiActivity，里面放一个TvView控件\n"
            "2. 通过TvInputManager获取已注册的HDMI输入设备ID\n"
            "3. 调用TvView.tune(inputId, Uri.EMPTY)切换到指定HDMI端口\n\n"
            "TvView会自动与底层的HdmiInputService建立会话，"
            "不需要任何特殊权限，也不需要第三方APP。"
        ),
        "code_example": (
            "// 核心代码\n"
            "val tvView = TvView(this)\n"
            "setContentView(tvView)\n"
            "\n"
            "// 切换到HDMI1\n"
            "val inputId = \"com.droidlogic.tvinput/.services.Hdmi1InputService/HW5\"\n"
            "tvView.tune(inputId, Uri.EMPTY)\n"
            "\n"
            "// HDMI2: com.droidlogic.tvinput/.services.Hdmi2InputService/HW6\n"
            "// HDMI3: com.droidlogic.tvinput/.services.Hdmi3InputService/HW7"
        ),
        "lesson": (
            "当Intent方式走不通时，要回到Android API本身寻找方案。"
            "TvView + TvInputManager是Android TV切换信号源的标准API，"
            "比发Intent更可靠，因为它直接和系统服务通信，不依赖第三方APP是否注册了对应的Activity。\n\n"
            "参考其他APK的做法时，要理解它们工作的前提条件。"
            "ETV酒店系统能用SETUP_INPUTS是因为它配套的com.elong.remote服务注册了处理该Intent的Activity。"
        ),
    },
    {
        "id": 3,
        "title": "安卓6电视上HDMI3端口不在TvInputManager列表中",
        "symptom": (
            "使用TvView API切换HDMI时，TvInputManager.getInputList()只返回了HDMI1和HDMI2，"
            "没有HDMI3。导致切换到HDMI3时找不到对应的input ID。"
        ),
        "cause": (
            "Amlogic芯片的Android TV系统有一个特点：HDMI端口是按需注册的。"
            "当一个HDMI口从未插入过设备时，系统不会在TvInputManager中注册该端口。\n\n"
            "具体表现：\n"
            "• 最初只有HDMI1在列表中（因为之前通过设置用过HDMI1）\n"
            "• 插入HDMI2后，系统自动注册了Hdmi2InputService/HW6\n"
            "• 插入HDMI3后，系统自动注册了Hdmi3InputService/HW7\n"
            "• 但如果没插过线，列表中就没有该端口"
        ),
        "attempts": [
            {"name": "直接用构造的ID调谐", "desc": "即使不在列表中也用构造的inputId调用tvView.tune()，但TvView内部仍需要服务已注册", "success": False},
        ],
        "solution": (
            "在findInputIdForPort()方法中增加逻辑：\n"
            "1. 先在tvInputList中查找（已注册的端口）\n"
            "2. 如果找不到，仍然返回根据命名规则构造的inputId\n"
            "3. 让TvView.tune()尝试调谐，如果失败则走备用方案\n\n"
            "同时，当用户第一次使用某个HDMI端口时，需要先手动通过电视设置切换到该端口一次，"
            "让系统注册该端口的InputService。之后APP就能正常切换了。\n\n"
            "实测发现：只要物理上插过一次线，系统就会自动注册，以后拔掉线也不会消失。"
        ),
        "lesson": (
            "硬件设备的状态会影响软件API的可用性。TvInputManager的输入列表不是静态的，"
            "会根据实际连接的设备动态变化。做兼容性处理时要考虑'设备从未连接过'的情况。"
        ),
    },
    {
        "id": 4,
        "title": "APP保活服务反复重启目标应用，导致HDMI模式被频繁打断",
        "symptom": (
            "切换到HDMI模式后，APP的保活服务每隔30秒检查一次，发现目标APP（如投屏APP）不在运行，"
            "就重新启动它，导致HDMI画面被覆盖。出现'刚切到HDMI，又被切回APP'的问题。"
        ),
        "cause": (
            "保活服务的逻辑是：检查目标APP是否在运行 → 如果没运行 → 重新启动。\n"
            "但在HDMI模式下，目标不是APP而是HDMI输入，不需要启动任何APP。"
            "而Launcher自身在前台时，保活服务误以为需要重新启动之前的投屏APP。"
        ),
        "attempts": [
            {"name": "增加启动冷却期", "desc": "启动APP后15秒内不再重复启动，但只是延缓问题，没有根治", "success": False},
        ],
        "solution": (
            "在保活检查逻辑中区分两种模式：\n\n"
            "1. APP模式：检查目标APP是否在运行，如果没运行就启动\n"
            "2. HDMI模式：检查HdmiActivity是否在前台，如果在前台说明HDMI正常工作中，不需要做任何事\n\n"
            "关键代码逻辑：\n"
            "• HDMI模式下，如果HdmiActivity在前台 → 正常，跳过保活检查\n"
            "• HDMI模式下，如果HdmiActivity不在前台 → 重新启动HdmiActivity\n"
            "• 刚切换到HDMI后设置10秒冷却期，防止保活服务立即把画面切走\n\n"
            "同时在KeepAliveForegroundService的通知中显示当前状态：\n"
            "• APP模式：'正在运行 [APP名称]'\n"
            "• HDMI模式且自动切换：'HDMI自动切换中'\n"
            "• HDMI模式：'正在显示 HDMI[端口号]'"
        ),
        "lesson": (
            "保活逻辑需要根据当前策略模式做不同的处理，不能用同一套逻辑覆盖所有场景。"
            "特别是HDMI模式和APP模式的目标完全不同——APP模式要确保某个APP在前台，"
            "HDMI模式要确保HdmiActivity在前台。"
        ),
    },
    {
        "id": 5,
        "title": "安卓6电视无法检测目标APP是否在运行（跨进程检测被系统阻止）",
        "symptom": (
            "保活服务需要检测投屏APP是否还在运行，但在安卓6电视上：\n"
            "• pidof命令无权限\n"
            "• /proc/[pid]/cmdline 文件 Permission denied\n"
            "• getRunningTasks 被MIUI阻止\n"
            "• runningAppProcesses 只返回自己的进程\n\n"
            "所有常规检测方式全部失败，无法判断目标APP是否还活着。"
        ),
        "cause": (
            "小米电视的MIUI系统对后台进程检测做了严格限制：\n"
            "• SELinux Enforcing模式下，第三方APP无法读取其他进程信息\n"
            "• MIUI的BLOCK-MONITOR功能阻止了getRunningTasks调用\n"
            "• runningAppProcesses在安卓6+只返回调用者自己的进程\n\n"
            "这是安卓系统的安全设计，不是bug。"
        ),
        "attempts": [
            {"name": "pidof命令", "desc": "执行pidof查询进程，但安卓6的pidof忽略包名参数，返回所有进程PID（不可信结果）", "success": False},
            {"name": "读取/proc文件", "desc": "遍历/proc/[pid]/cmdline，Permission denied", "success": False},
            {"name": "getRunningTasks", "desc": "被MIUI BLOCK-MONITOR阻止", "success": False},
            {"name": "runningAppProcesses", "desc": "只返回自己的进程", "success": False},
        ],
        "solution": (
            "使用'反向推理'方法：既然无法检测别人，就检测自己。\n\n"
            "核心思路：\n"
            "• 本APP是HOME应用（桌面启动器），如果本APP在前台 → 说明没有其他APP在前台 → 目标APP没在运行\n"
            "• 如果本APP在后台 → 说明有其他APP在前台 → 目标APP大概率在运行\n\n"
            "实现方式：\n"
            "• 在MainActivity的onResume中记录 launcher_foreground = true\n"
            "• 在MainActivity的onPause中记录 launcher_foreground = false\n"
            "• 检测时读取SharedPreferences中的标记\n\n"
            "再加一个启动冷却期（15秒）：刚启动过的APP不做重复检测，"
            "避免出现'启动APP → APP还在加载中 → 检测不到 → 又启动一次'的死循环。"
        ),
        "lesson": (
            "当正面检测走不通时，换个角度思考。作为HOME应用，"
            "自己是否在前台就是判断其他APP是否在运行的最佳依据。"
            "不需要多么高深的技术，简单有效就是最好的方案。"
        ),
    },
    {
        "id": 6,
        "title": "当贝桌面APK被360加固保护，无法反编译获取HDMI切换代码",
        "symptom": (
            "用户提供了一个当贝桌面APK，它在安卓6电视上能正常切换HDMI。"
            "我们想反编译它看看是怎么做到的，但反编译只能看到壳代码，看不到真正的业务逻辑。"
        ),
        "cause": (
            "当贝桌面使用了360加固（Qihoo保护），这是一种代码混淆和加密技术：\n"
            "• APK安装后，真正的DEX代码被加密存储\n"
            "• 运行时由壳程序解密加载\n"
            "• 反编译工具（jadx、baksmali等）只能看到壳的stub代码\n"
            "• 无法获取原始Java/Kotlin源码"
        ),
        "attempts": [
            {"name": "jadx反编译", "desc": "只能看到360加固的壳代码，业务类全部是stub", "success": False},
            {"name": "androguard分析", "desc": "同样被加固保护，无法获取真实代码", "success": False},
            {"name": "dexdump/baksmali", "desc": "DEX文件被加密，无法解析", "success": False},
        ],
        "solution": (
            "换一个没有加固的参考APK。用户提供了ETV酒店系统APK，"
            "虽然也做了代码混淆（变量名变成a/b/c），但至少能看到代码逻辑。\n\n"
            "从ETV中找到了关键的HDMI切换代码：\n"
            "• ZhenhuoDevice.java：Droidlogic芯片使用SETUP_INPUTS + INPUT_ID切换\n"
            "• HotelUtils.java：小米电视使用com.elong.remote.SET_SOURCE服务切换\n"
            "• QNDevice.java：极米投影仪使用qnbar服务的changeSignalSource方法\n\n"
            "这些参考代码最终帮助确认了TvView API方案的正确性。"
        ),
        "lesson": (
            "当一条路走不通时要及时换路。360加固无法绕过就不要浪费时间，"
            "找其他没有加固的参考APK同样能获取有用信息。\n\n"
            "同时要注意：参考代码的做法不一定适合我们。"
            "ETV的SETUP_INPUTS方式需要配套的com.elong.remote服务，我们没有，所以用不了。"
            "最终还是要理解原理，找到适合自己场景的方案。"
        ),
    },
    {
        "id": 7,
        "title": "小米电视设置APP被卸载导致无法手动切换HDMI",
        "symptom": (
            "在测试过程中，安卓6电视上的 com.xiaomi.mitv.settings（小米电视设置）被意外卸载，"
            "导致无法通过电视设置菜单手动切换HDMI，也无法进行信号源选择。"
        ),
        "cause": (
            "AdbKeeperTest工具在冻结预装应用时，可能误将MiTVSettings冻结并卸载了用户0的数据。"
            "虽然系统分区中的APK文件还在，但用户0上的安装记录被清除了。"
        ),
        "solution": (
            "通过ADB从系统分区重新安装：\n"
            "  adb shell pm install -r /system/app/MiTVSettings2/MiTVSettings2.apk\n\n"
            "注意：\n"
            "• 安卓6不支持 pm install-existing 命令\n"
            "• 但可以直接从/system分区安装APK\n"
            "• 需要 -r 参数覆盖安装"
        ),
        "lesson": (
            "冻结/卸载系统应用时要特别小心，最好维护一个白名单，"
            "像设置、信号源这类核心应用绝对不能被冻结。"
        ),
    },
    {
        "id": 8,
        "title": "不同芯片的电视HDMI输入服务命名不同（Amlogic vs MediaTek）",
        "symptom": (
            "同一份代码，在Amlogic芯片的小米电视上能切换HDMI，"
            "但在MediaTok芯片的小米电视上切换失败，找不到HDMI输入ID。"
            "这是因为两种芯片的HDMI输入服务包名和类名完全不同。"
        ),
        "cause": (
            "小米电视使用了不同芯片方案，每种芯片的TV Input服务命名规则不同：\n\n"
            "Amlogic/Droidlogic芯片（如MiTV4-ANSM0, 安卓6）：\n"
            "  每个HDMI端口有独立的InputService\n"
            "  HDMI1: com.droidlogic.tvinput/.services.Hdmi1InputService/HW5\n"
            "  HDMI2: com.droidlogic.tvinput/.services.Hdmi2InputService/HW6\n"
            "  HDMI3: com.droidlogic.tvinput/.services.Hdmi3InputService/HW7\n\n"
            "MediaTok芯片（如MiTV-MFTP0, 安卓9）：\n"
            "  所有HDMI端口共用一个HDMIInputService，通过HW编号区分\n"
            "  HDMI1: com.mediatek.tvinput/.hdmi.HDMIInputService/HW5\n"
            "  HDMI2: com.mediatek.tvinput/.hdmi.HDMIInputService/HW6\n"
            "  HDMI3: com.mediatek.tvinput/.hdmi.HDMIInputService/HW7\n"
            "  HDMI4: com.mediatek.tvinput/.hdmi.HDMIInputService/HW8\n\n"
            "关键区别：\n"
            "• Droidlogic是每个端口独立的服务类（Hdmi1InputService, Hdmi2InputService...）\n"
            "• MediaTek是同一个服务类（HDMIInputService），用HW编号区分端口\n"
            "• Droidlogic的H是小写Hdmi，MediaTek的H是大写HDMI"
        ),
        "solution": (
            "在HdmiActivity的findInputIdForPort()方法中支持多种芯片平台：\n\n"
            "1. 维护一个已知Input ID列表，包含Amlogic和MediaTek两种命名\n"
            "2. 先在tvInputList中匹配已知ID（最可靠）\n"
            "3. 按HW编号通用匹配（HW5=端口1, HW6=端口2...）\n"
            "4. 按名称模糊匹配（兜底）\n"
            "5. 如果只有一个HDMI输入，直接使用\n\n"
            "这样不管电视用的是什么芯片，代码都能自动适配。"
        ),
        "lesson": (
            "做硬件相关功能时，一定要考虑不同硬件平台的差异。"
            "同品牌的小米电视，不同型号可能用完全不同的芯片方案。"
            "通过TvInputManager API获取输入列表再匹配，比硬编码芯片名称更可靠。"
        ),
    },
    {
        "id": 9,
        "title": "MediaTok芯片电视不发送HDMI_PLUGGED广播，自动切换失效",
        "symptom": (
            "在MediaTok芯片的小米电视（安卓9, MiTV-MFTP0）上，插拔HDMI线时：\n"
            "• 电视能检测到HDMI接入（弹出'HDMI3已接入'提示）\n"
            "• 但APP的HdmiReceiver收不到android.intent.action.HDMI_PLUGGED广播\n"
            "• dumpsys显示HDMI_PLUGGED sticky广播state始终为false\n"
            "• 导致HDMI自动切换功能完全失效"
        ),
        "cause": (
            "MediaTek芯片的HDMI检测机制与Amlogic不同：\n\n"
            "Amlogic芯片：\n"
            "  硬件HPD检测 → 系统发送HDMI_PLUGGED广播 → APP可以监听\n\n"
            "MediaTok芯片：\n"
            "  使用HotPlugDetectionAction轮询机制检测HDMI\n"
            "  检测到HDMI后由系统内部处理（弹出提示、切换输入）\n"
            "  但不发送标准的HDMI_PLUGGED广播给第三方APP\n"
            "  HDMI_PLUGGED sticky广播的state始终为false\n\n"
            "这不是bug，是MediaTek芯片方案的实现选择。"
        ),
        "solution": (
            "目前无完美的软件解决方案。替代方案：\n\n"
            "1. 通过后台管理页面手动切换HDMI策略（最可靠）\n"
            "2. 定时轮询TvInputManager检查HDMI输入状态变化（耗电，不推荐）\n"
            "3. 监听MediaTek特有的系统广播/属性变化（需要逆向分析，不通用）\n\n"
            "兼容性总结：\n"
            "• Amlogic + 安卓9：支持HDMI_PLUGGED广播，自动切换可用\n"
            "• Amlogic + 安卓6：不支持HPD，无HDMI_PLUGGED广播\n"
            "• MediaTek + 安卓9：支持HDMI检测但不发HDMI_PLUGGED广播\n"
            "• 未测试：安卓11及其他芯片平台"
        ),
        "lesson": (
            "不同芯片方案的电视，即使是同一品牌，底层实现也可能完全不同。"
            "做HDMI自动检测功能时，不能假设所有设备都支持HDMI_PLUGGED广播。"
            "需要在APP中增加芯片/型号检测，对不同平台使用不同的检测策略。"
        ),
    },
    {
        "id": 10,
        "title": "安卓6电视(811会议室)开机白屏闪黑好几秒",
        "symptom": (
            "811会议室的小米电视(Amlogic芯片, 安卓6.0.1)每次开机后，画面会出现明显的白屏闪烁，"
            "然后黑屏，反复闪好几秒后才能正常显示TV Launcher界面并执行策略。"
            "用户体验很差，像电视坏了，而且闪烁对眼睛也不好。\n"
            "其他安卓9/11的电视没有这个问题。"
        ),
        "cause": (
            "白屏闪烁由三个因素叠加造成：\n\n"
            "1. BootReceiver延迟3秒启动Activity\n"
            "   开机后系统先启动默认Launcher（白色主题），等3秒才启动TV Launcher。\n"
            "   这3秒内用户看到的是系统默认Launcher的白色界面。\n\n"
            "2. MaterialComponents主题冷启动初始化慢\n"
            "   MainActivity使用Theme.MaterialComponents.DayNight.NoActionBar主题。\n"
            "   MaterialComponents库在Android 6上初始化耗时较长，\n"
            "   系统创建Activity预览窗口时会短暂显示白色背景（默认windowBackground），\n"
            "   然后才切换到主题设定的深色背景(#0A0E14)。\n"
            "   安卓9+设备性能好，这个白闪几乎看不到；安卓6设备性能差，白闪非常明显。\n\n"
            "3. Activity启动与策略执行耦合\n"
            "   BootReceiver把启动Activity和执行策略放在同一个3秒延迟里，\n"
            "   导致即使Activity可以更快启动，也要等3秒。"
        ),
        "attempts": [
            {"name": "减少BootReceiver延迟", "desc": "将延迟从3秒改为0秒，立即启动Activity，策略执行单独延迟1.5秒", "success": True},
            {"name": "使用轻量启动主题", "desc": "创建Theme.TvLauncher.Splash继承android:Theme.NoTitleBar.Fullscreen，避免MaterialComponents初始化白屏", "success": True},
            {"name": "Activity onCreate切换主题", "desc": "Manifest用Splash主题，onCreate中setTheme切回Fullscreen主题，兼顾冷启动速度和运行时功能", "success": True},
        ],
        "solution": (
            "三层优化方案，彻底消除白屏闪烁：\n\n"
            "1. 轻量启动主题（核心修复）\n"
            "   新增Theme.TvLauncher.Splash，继承android:Theme.NoTitleBar.Fullscreen（Android框架自带，无需初始化第三方库），\n"
            "   设置windowBackground=#0A0E14（深色），确保冷启动预览窗口直接是深色。\n"
            "   AndroidManifest中MainActivity使用Splash主题，\n"
            "   onCreate中通过setTheme(R.style.Theme_TvLauncher_Fullscreen)切换回完整主题。\n\n"
            "2. BootReceiver立即启动\n"
            "   去掉3秒延迟，收到开机广播后立即启动MainActivity和保活服务。\n"
            "   策略执行单独延迟1.5秒，等Activity完全启动后再执行。\n\n"
            "3. 补充ACTION_REBOOT处理\n"
            "   Manifest中注册了REBOOT广播但代码未匹配，已补充。"
        ),
        "code_example": (
            "# themes.xml - 启动专用主题\n"
            '<style name="Theme.TvLauncher.Splash" parent="android:Theme.NoTitleBar.Fullscreen">\n'
            '    <item name="android:windowBackground">#0A0E14</item>\n'
            '    <item name="android:windowNoTitle">true</item>\n'
            '    <item name="android:windowFullscreen">true</item>\n'
            '</style>\n\n'
            "# MainActivity.kt - onCreate中切换主题\n"
            "override fun onCreate(savedInstanceState: Bundle?) {\n"
            "    setTheme(R.style.Theme_TvLauncher_Fullscreen)\n"
            "    super.onCreate(savedInstanceState)\n"
            "    ...\n"
            "}"
        ),
        "lesson": (
            "Android冷启动白屏是常见问题，尤其在使用AppCompat/MaterialComponents等支持库时。"
            "解决方案的核心思路：用最轻量的主题创建预览窗口（避免库初始化开销），"
            "Activity创建后再切换到完整主题。这种'splash theme'模式是Android开发的最佳实践。\n\n"
            "同时，对于低性能设备（如安卓6电视），启动延迟的每一秒都会被用户感知到，"
            "应该尽可能早地启动Activity，把耗时操作延迟到Activity创建之后。"
        ),
    },
    {
        "id": 11,
        "title": "管理后台策略下拉框选择时消失、选择值自动回退",
        "symptom": (
            "管理后台'执行策略'栏存在两个问题：\n\n"
            "1. 下拉框消失：点开策略下拉框正在选择时，下拉框突然消失，无法完成选择操作\n"
            "2. 选择值回退：选择了其他策略后，还没来得及点'下发'按钮，选择就自动变回了原来的策略\n\n"
            "两个问题都频繁出现，严重影响日常操作。"
        ),
        "cause": (
            "两个问题共享同一个根因：setInterval(loadData, 5000)每5秒执行一次loadData()，\n"
            "loadData()通过innerHTML完全重建整个设备表格，包括每行的策略<select>下拉框。\n\n"
            "问题1（下拉框消失）：\n"
            "用户打开下拉框浏览选项时，5秒计时器触发，loadData()用全新的DOM替换整个表格，\n"
            "正在打开的<select>元素被销毁，下拉框瞬间消失。\n\n"
            "问题2（选择值回退）：\n"
            "用户选择了新策略（本地DOM状态改变），但在点击'下发'前，5秒计时器触发，\n"
            "loadData()从服务器获取设备数据（服务器仍是旧策略），\n"
            "根据服务器的d.policy_id重建select，下拉框回退显示服务器上的旧值，用户的选择丢失。"
        ),
        "attempts": [
            {"name": "轮询暂停机制", "desc": "select获得焦点时暂停轮询，失焦后延迟1秒恢复，防止下拉框被销毁", "success": True},
            {"name": "本地选择暂存", "desc": "用pendingSelections字典记录用户已修改但未保存的选择，轮询刷新时优先使用本地值", "success": True},
        ],
        "solution": (
            "双重修复方案：\n\n"
            "1. 轮询暂停机制\n"
            "   添加pollPaused标志，当任何策略<select>获得焦点时设为true，暂停setInterval刷新。\n"
            "   失焦后延迟1秒恢复（给change事件和click事件留出时间）。\n"
            "   setInterval改为：setInterval(() => { if (!pollPaused) loadData(); }, 5000)\n\n"
            "2. 本地选择暂存\n"
            "   添加pendingSelections字典，当用户通过onchange选择新策略时，\n"
            "   记录到pendingSelections[sel-${deviceId}] = value。\n"
            "   loadData()重建select时，优先使用pendingSelections中的值而非服务器返回的值：\n"
            "   const selectedValue = pendingSelections[`sel-${d.id}`] !== undefined\n"
            "       ? pendingSelections[`sel-${d.id}`]\n"
            "       : d.policy_id || ''\n"
            "   bindPolicy()下发成功后，清除对应的pendingSelections记录。"
        ),
        "code_example": (
            "// 轮询暂停\n"
            "let pollPaused = false;\n"
            "document.addEventListener('focusin', e => {\n"
            "    if (e.target.tagName === 'SELECT' && e.target.id?.startsWith('sel-'))\n"
            "        pollPaused = true;\n"
            "});\n\n"
            "// 本地选择暂存\n"
            "let pendingSelections = {};\n"
            "function onPolicySelect(deviceId, value) {\n"
            "    pendingSelections[`sel-${deviceId}`] = value;\n"
            "}\n\n"
            "// bindPolicy成功后清除暂存\n"
            "async function bindPolicy(deviceId) {\n"
            "    await fetch(...);\n"
            "    delete pendingSelections[`sel-${deviceId}`];\n"
            "}"
        ),
        "lesson": (
            "使用innerHTML重建DOM时，必须考虑用户正在交互的元素会被销毁的问题。"
            "对于需要轮询刷新的页面，要在用户交互期间暂停刷新，"
            "并保留用户未提交的本地状态，防止被服务器数据覆盖。\n\n"
            "更好的做法是使用框架（Vue/React）的数据驱动视图，只更新变化的部分，"
            "而不是整个重建DOM。但对于简单的内嵌模板页面，暂停轮询+本地暂存是够用的方案。"
        ),
    },
    {
        "id": 12,
        "title": "HDMI切换后系统弹窗'输入源已接入'需手动确认",
        "symptom": (
            "在Amlogic芯片的小米电视上，调用TvView.tune()切换到HDMI后，"
            "系统弹出'输入源已接入'提示对话框，需要用户手动按遥控器确认键才能看到HDMI画面。"
            "这导致自动化HDMI切换策略失效——虽然代码切换了HDMI，但用户看到的是弹窗而不是画面。"
        ),
        "cause": (
            "Amlogic/Droidlogic芯片的TV Input Framework在TvView.tune()后，"
            "会通过系统UI弹出'输入源已接入'的通知对话框，要求用户点击'查看'按钮。\n\n"
            "这是Droidlogic中间层的行为，不是标准Android AOSP行为。"
            "MediaTek芯片通常不弹窗或弹窗自动消失。"
        ),
        "attempts": [
            {"name": "定时发DPAD_CENTER", "desc": "调谐后固定时间发确认键，但无弹窗时误操作", "success": False},
            {"name": "DPAD_CENTER+BACK", "desc": "确认后发BACK关闭残留通知，但BACK会退出HdmiActivity", "success": False},
            {"name": "DPAD_CENTER+ESCAPE", "desc": "ESCAPE关闭通知，但ESCAPE会关闭HDMI信号显示", "success": False},
        ],
        "solution": (
            "智能弹窗确认策略——只在检测到窗口失焦时才发确认键：\n\n"
            "1. 调谐后启动焦点检查循环（每500ms检查一次，最多8次）\n"
            "2. 如果HdmiActivity失去焦点（被弹窗抢走），发送DPAD_CENTER确认\n"
            "3. 确认后如果焦点恢复，再发送DPAD_RIGHT+DPAD_CENTER关闭残留通知\n"
            "4. 如果HdmiActivity一直有焦点，说明没有弹窗，不发送任何按键\n\n"
            "同时添加调谐后5秒BACK键保护期，防止确认流程中的BACK键误退出HdmiActivity。"
        ),
        "lesson": (
            "不能盲目发送按键事件——必须先检测当前窗口状态，只在确实需要时才发送。"
            "不同芯片平台的弹窗行为不同，Amlogic会弹窗需确认，MediaTek通常不会。"
            "后续应该根据检测到的芯片平台调整弹窗确认策略。"
        ),
    },
    {
        "id": 13,
        "title": "HdmiActivity被系统重建后丢失HDMI端口信息，切回HDMI1",
        "symptom": (
            "从HDMI1策略切换到HDMI2策略时，新HdmiActivity正确找到HDMI2并调谐，"
            "但随后系统重建了Activity，第二次onCreate时读取到的是旧HdmiActivity的HDMI1信息，"
            "导致画面切回了HDMI1。"
        ),
        "cause": (
            "Activity重建时信息丢失有两个原因：\n\n"
            "1. 静态变量savedInputId是进程级共享的。旧HdmiActivity(HDMI1)设置的savedInputId"
            "没有在新HdmiActivity(HDMI2)创建时清除，第二次onCreate从静态变量恢复了旧的HDMI1值。\n\n"
            "2. savedInstanceState优先级设置不当。原来优先用savedInstanceState，"
            "但Activity重建场景下savedInstanceState保存的也是旧端口的inputId。"
        ),
        "attempts": [
            {"name": "savedInstanceState保存", "desc": "保存inputId到Bundle，但Activity重建时恢复的还是旧值", "success": False},
            {"name": "静态变量保存", "desc": "用static变量保存，但进程级共享导致新旧Activity冲突", "success": False},
        ],
        "solution": (
            "修复inputId恢复优先级：intent > 静态变量 > savedInstanceState\n\n"
            "1. 优先从intent获取inputId（新的切换请求），获取到后立即更新静态变量\n"
            "2. 只有intent没有指定端口时才从静态变量恢复（Activity重建场景）\n"
            "3. savedInstanceState作为最后兜底\n\n"
            "关键代码：从intent解析到inputId后，同时更新savedInputId = intentInputId，"
            "确保后续重建时能恢复到正确的端口。"
        ),
        "lesson": (
            "静态变量在进程级共享，多个Activity实例会互相干扰。"
            "当有新的intent请求时，必须以intent为准并更新所有缓存。"
            "恢复优先级应该是：新请求 > 进程缓存 > 实例缓存。"
        ),
    },
    {
        "id": 14,
        "title": "不同芯片平台HDMI切换行为差异大，需分别适配",
        "symptom": (
            "同一份HDMI切换代码在不同芯片的小米电视上表现不同：\n"
            "• Amlogic: TvView.tune()可靠，但弹窗需确认\n"
            "• MediaTek: TvView.tune()可靠，弹窗行为不一致\n"
            "• MStar: TvView.tune()可能静默失败\n"
            "• Realtek: tvInputList可能为空，需构造ID尝试\n\n"
            "当前代码对所有平台用同一套逻辑，无法针对不同平台优化。"
        ),
        "cause": (
            "四种芯片平台的TV Input Framework实现差异：\n\n"
            "Amlogic/Droidlogic: 每个HDMI端口独立InputService，包名com.droidlogic.tvinput\n"
            "MediaTek: 所有端口共用HDMIInputService，包名com.mediatek.tvinput\n"
            "MStar: 独立InputService但TvView可能不可靠，包名com.mstar.tvinput\n"
            "Realtek: 独立InputService但注册可能不完整，包名com.realtek.tvinput\n\n"
            "目前代码没有检测芯片平台，findInputIdForPort()按固定顺序尝试所有平台的ID，"
            "无法根据平台优先使用最可靠的切换方式。"
        ),
        "solution": (
            "计划添加芯片平台检测，根据平台调整切换策略：\n\n"
            "检测方法（两种互补）：\n"
            "1. Build.HARDWARE属性：\n"
            "   Amlogic: gxl/g12a/sm1/sc2\n"
            "   MediaTek: mt5882/mt5886/mt5891/mt5895\n"
            "   MStar: mst6a918/mst6a928/mst6a938\n"
            "   Realtek: rtd2851/rtd2833\n"
            "2. 已安装tvinput包名检测（最实用）\n\n"
            "策略调整：\n"
            "• Amlogic/MTK: TvView.tune()为主，Amlogic需弹窗确认\n"
            "• MStar: 优先用EXTSRC_PLAY Intent切换\n"
            "• Realtek: 增加调谐延迟和重试\n\n"
            "同时上报芯片类型到后端，方便远程诊断。"
        ),
        "lesson": (
            "做硬件相关功能时，不同芯片平台的实现差异是最大的坑。"
            "同品牌的小米电视，不同型号可能用完全不同的芯片方案。"
            "必须在代码中加入芯片检测，针对不同平台使用不同的策略。"
            "Build.HARDWARE + 已安装包名检测是最可靠的组合方案。"
        ),
    },
    {
        "id": 15,
        "title": "HDMI弹窗确认代码重复发送DPAD_CENTER导致自动误操作",
        "symptom": (
            "在安卓6电视(10.181.184.254)上，电视自动点击了屏幕上的设置按钮，"
            "然后在密码输入框中自动输入了4个'1'。用户并未操作遥控器，完全是APP自动触发的。\n\n"
            "同时发现：管理后台的WiFi状态列只显示'WiFi'文字，不显示电视机连接的WiFi名称。"
        ),
        "cause": (
            "根本原因是HdmiActivity中scheduleConfirmDialog()的逻辑缺陷：\n\n"
            "1. 重复发送确认键\n"
            "   原代码在每次检查时，只要!hasFocus就发送DPAD_CENTER，最多8次检查。\n"
            "   如果窗口持续失焦（如Activity正在切换），代码会在每次检查时都发送DPAD_CENTER，\n"
            "   而不是只发送一次。日志证实：连续发送了3-5次DPAD_CENTER。\n\n"
            "2. 失焦判断不准确\n"
            "   原代码只检查!hasFocus就发送确认键，但失焦原因可能是：\n"
            "   • HDMI弹窗出现（需要确认）\n"
            "   • Activity切换（不需要确认）\n"
            "   • 系统通知（不需要确认）\n"
            "   代码没有区分这些情况，导致在非弹窗场景下误发按键。\n\n"
            "3. onDestroy未取消确认检查Runnable\n"
            "   confirmDialogRunnable只保存在局部变量中，onDestroy只取消了pendingTuneRunnable，\n"
            "   未取消confirmDialogRunnable。HdmiActivity被销毁后，handler仍会执行确认检查，\n"
            "   此时hasFocus=false（窗口已销毁），代码继续发送DPAD_CENTER，影响其他界面。\n\n"
            "4. WiFi SSID读取缺失位置权限\n"
            "   Android 6+设备读取WiFi SSID需要ACCESS_COARSE_LOCATION权限，\n"
            "   但APP未声明此权限，导致WifiManager.getConnectionInfo().getSsid()返回<unknown ssid>，\n"
            "   前端只能显示通用的'WiFi'文字而非具体WiFi名称。"
        ),
        "attempts": [
            {"name": "焦点丢失时发确认键", "desc": "原方案：!hasFocus就发DPAD_CENTER，但失焦会持续多轮检查，导致重复发送", "success": False},
        ],
        "solution": (
            "三层修复：\n\n"
            "1. 确认键只发一次\n"
            "   改为检测焦点变化（从有焦点→无焦点）的瞬间才发送DPAD_CENTER，\n"
            "   用lastFocusState记录上一次焦点状态，只在focusJustLost && !hasConfirmedDialog时发送。\n"
            "   发送后设置hasConfirmedDialog=true，后续即使继续失焦也不再重复发送。\n\n"
            "2. onDestroy取消确认检查\n"
            "   将confirmDialogRunnable保存为成员变量，在onDestroy中同时取消\n"
            "   pendingTuneRunnable和confirmDialogRunnable，防止Activity销毁后继续发按键。\n\n"
            "3. 添加位置权限\n"
            "   AndroidManifest.xml添加ACCESS_COARSE_LOCATION和ACCESS_FINE_LOCATION，\n"
            "   MainActivity.onCreate中请求运行时权限，\n"
            "   已部署设备通过adb shell pm grant授权。"
        ),
        "code_example": (
            "# 修复后的弹窗确认逻辑\n"
            "var lastFocusState = true  // 记录上次焦点状态\n"
            "if (focusJustLost && !hasConfirmedDialog) {\n"
            "    // 只在焦点刚丢失且未确认过时发送一次\n"
            "    Runtime.getRuntime().exec(arrayOf('input', 'keyevent', 'KEYCODE_DPAD_CENTER'))\n"
            "    hasConfirmedDialog = true\n"
            "} else if (hasConfirmedDialog && !focusRecoveredOnce && currentFocus) {\n"
            "    // 焦点恢复后关闭残留通知\n"
            "    Runtime.getRuntime().exec(arrayOf('input', 'keyevent', 'KEYCODE_DPAD_RIGHT'))\n"
            "    Runtime.getRuntime().exec(arrayOf('input', 'keyevent', 'KEYCODE_DPAD_CENTER'))\n"
            "    focusRecoveredOnce = true\n"
            "}\n\n"
            "# onDestroy中取消确认检查\n"
            "confirmDialogRunnable?.let { handler.removeCallbacks(it) }"
        ),
        "lesson": (
            "发送按键事件（input keyevent）是全局性的，会影响屏幕上任何有焦点的元素。"
            "必须严格控制发送次数和条件：\n"
            "• 只在确实需要时发送（焦点变化检测）\n"
            "• 每个事件只发送一次（用标志位防止重复）\n"
            "• Activity销毁时必须取消所有handler回调\n\n"
            "WiFi SSID读取是Android 6+常见坑：需要位置权限才能获取SSID，"
            "TV设备虽然不需要定位功能，但系统权限模型一视同仁。"
        ),
    },
    {
        "id": 16,
        "title": "系统应用停用缺乏安全机制，误操作风险高",
        "symptom": (
            "应用管理中的「停用」功能可以一键停用系统应用，但只使用了简单的confirm()确认对话框。"
            "维护人员可能误点击「停用」按钮导致电视关键功能异常，如停用系统桌面、输入法等核心应用。"
        ),
        "cause": (
            "停用系统应用使用 pm uninstall --user 0 命令，虽然可以通过恢复出厂设置恢复，"
            "但对于会议室电视来说，恢复出厂设置意味着需要重新部署所有配置。"
            "简单的confirm对话框不足以防止误操作。"
        ),
        "attempts": [
            {"name": "二次confirm确认", "desc": "增加第二个confirm确认，但用户习惯性点击确定，效果不好", "success": False},
        ],
        "solution": (
            "引入输入确认机制：停用系统应用时弹出专用确认对话框，"
            "要求用户手动输入「确认停用」四个字才能继续操作。"
            "这种模式借鉴GitHub删除仓库的安全机制，"
            "通过强制用户进行有意识的输入来防止误操作。"
        ),
        "lesson": (
            "危险操作不能只靠简单的确认对话框，用户对confirm/alert会形成肌肉记忆直接点确定。"
            "输入特定文字的确认方式虽然增加操作步骤，但能有效防止误操作。"
        ),
    },
    {
        "id": 17,
        "title": "Android 6电视上设置图标显示为X",
        "symptom": (
            "在Android 6电视上，APP主页右上角的设置按钮显示为一个X或方框，"
            "而不是预期的齿轮图标。"
        ),
        "cause": (
            "布局XML中使用了Unicode字符 ⚙ (U+2699) 作为设置图标。"
            "Android 6的默认字体不包含此Unicode字符的图形表示，"
            "导致显示为缺失字符标记（X或方框）。"
        ),
        "attempts": [
            {"name": "使用其他Unicode字符", "desc": "尝试使用三横线或其他符号，同样存在兼容性问题", "success": False},
        ],
        "solution": (
            "将图标改为纯文字标签「设置」和「关于」，"
            "文字在所有Android版本上都能正确显示，"
            "且更符合TV遥控器操作的简洁设计风格。"
        ),
        "lesson": (
            "TV应用需要兼容低版本Android系统，不能依赖Unicode emoji字符作为UI元素。"
            "纯文字标签是最安全的方案，且在TV场景下更易识别。"
        ),
    },
    {
        "id": 18,
        "title": "设备从后台移除后无法获取电视IP地址",
        "symptom": (
            "如果维护人员在后台管理页面误删了某台电视设备，"
            "由于APP已设置为默认桌面且禁用了系统设置入口，"
            "维护人员无法在电视上查看IP地址来重新部署。"
        ),
        "cause": (
            "APP主页没有显示设备自身信息的入口，"
            "所有设备信息都依赖后台管理页面查看。"
            "一旦设备从后台移除，就失去了查看IP的能力。"
        ),
        "solution": (
            "在APP主页右上角添加「关于」按钮，点击后显示本机完整信息：\n"
            "• 设备型号、Android版本、序列号\n"
            "• WiFi/有线IP地址和MAC地址\n"
            "• WiFi信号强度、频段、速度\n"
            "• 网络延迟和丢包率\n"
            "• 内存和存储使用情况\n"
            "• 管理服务器地址\n\n"
            "即使设备从后台移除，维护人员也能在电视上查看IP等信息来重新部署。"
        ),
        "lesson": (
            "管理系统的终端设备应该始终提供自身关键信息的查看方式，"
            "不能完全依赖服务端。当服务端数据丢失时，终端自身信息是恢复的最后一道保障。"
        ),
    },
    {
        "id": 19,
        "title": "小米电视连接有线网时WiFi被自动关闭",
        "symptom": (
            "在小米电视上，当已经通过有线网络连接时，手动开启WiFi开关后，"
            "WiFi会被系统自动关闭。这导致无法同时使用有线和WiFi，"
            "也无法切换到WiFi连接。参考APP（当贝桌面影视版）也存在同样的问题。"
        ),
        "cause": (
            "这是小米电视（MiUI TV）系统的网络策略：检测到有线网连接后，"
            "系统会自动禁用WiFi以节省电量和避免网络冲突。"
            "这是系统级行为，APP无法通过软件方式绕过。"
        ),
        "solution": (
            "在WiFi连接功能中增加检测：\n"
            "1. 启动WiFi连接页面时，自动检测是否有线网连接\n"
            "2. 如果检测到有线网，显示醒目警告：「小米电视连接有线时WiFi会被自动关闭，请先拔掉网线」\n"
            "3. 阻止用户在有线网连接时开启WiFi，避免开关反复闪烁\n"
            "4. WiFi开关操作前先检查有线网状态，有则弹出提示"
        ),
        "lesson": (
            "硬件限制需要提前告知用户，而不是让用户反复尝试后自己发现。"
            "明确的提示比隐晦的限制更友好。"
        ),
    },
    {
        "id": 20,
        "title": "设置页面黑字黑底、对话框式交互不符合TV遥控器操作",
        "symptom": (
            "APP设置页面存在多个问题：\n"
            "1. 黑色字体在深色背景上看不清\n"
            "2. 密码对话框一进入就弹出，按上下键会跳出键盘\n"
            "3. 所有设置项平铺在同一个页面，不符合TV遥控器逐级进入的操作习惯\n"
            "4. 点击「关于」或「设置」时卡顿几秒"
        ),
        "cause": (
            "1. 设置页面使用了默认深色主题但文字也是深色，导致看不清\n"
            "2. 密码验证在MainActivity.onItemClick中用AlertDialog实现，焦点管理不当\n"
            "3. 旧设计将所有设置项平铺，TV遥控器需要逐级导航才好操作\n"
            "4. NetworkInfoProvider.collect()在主线程执行ping测试，耗时3-6秒"
        ),
        "solution": (
            "1. 设置页面改用浅色背景(#F0F2F5) + 深色文字(#2D3436)，确保可读性\n"
            "2. 重构为层级式菜单导航（类似小米设置），主菜单 → 子页面，遥控器上下选择、确认进入\n"
            "3. 密码输入改为设置页面内的编辑框，而非弹窗\n"
            "4. NetworkInfoProvider耗时操作放到后台线程，关于页面先显示基础信息再加载数据"
        ),
        "lesson": (
            "TV应用和手机应用的交互逻辑完全不同：\n"
            "• TV用遥控器上下左右操作，需要逐级导航而非平铺\n"
            "• TV不能依赖弹窗交互，弹窗焦点管理对遥控器极不友好\n"
            "• 耗时操作必须在后台执行，TV的用户体验容忍度更低"
        ),
    },
    {
        "id": 21,
        "title": "设置页面和WiFi页面没有焦点高亮，遥控器无法识别当前选择项",
        "symptom": (
            "使用遥控器在设置页面和WiFi连接页面上下移动时，无法看到当前焦点在哪个选项上。"
            "所有项目看起来完全一样，用户不知道按确认键会操作哪个项目。"
        ),
        "cause": (
            "TV遥控器操作依赖焦点（Focus）系统，但原始代码没有为可点击项添加OnFocusChangeListener。"
            "没有视觉反馈，用户无法知道当前选中了哪个项目。"
            "这是TV应用与手机应用的核心区别：手机有触摸屏直接点击，TV需要焦点导航。"
        ),
        "solution": (
            "为所有可交互项添加OnFocusChangeListener，焦点获得时显示高亮样式：\n\n"
            "• 获得焦点：浅蓝背景(#E3F2FD) + 蓝色边框(#2196F3) + 2dp描边\n"
            "• 失去焦点：白色背景(#FFFFFF) + 灰色边框(#E0E0E0)\n"
            "• 使用GradientDrawable动态切换背景，避免创建大量XML selector文件\n\n"
            "涉及的组件：\n"
            "• SettingsActivity：所有菜单项、返回按钮、编辑框\n"
            "• WifiConnectActivity：WiFi开关、刷新按钮、WiFi列表项、返回按钮"
        ),
        "code_example": (
            "// 焦点高亮核心代码\n"
            "val FOCUS_BG = 0xFFE3F2FD.toInt()\n"
            "val FOCUS_BORDER = 0xFF2196F3.toInt()\n"
            "val NORMAL_BG = 0xFFFFFFFF.toInt()\n"
            "val NORMAL_BORDER = 0xFFE0E0E0.toInt()\n\n"
            "view.onFocusChangeListener = View.OnFocusChangeListener { _, hasFocus ->\n"
            "    view.background = if (hasFocus)\n"
            "        roundedBg(FOCUS_BG, FOCUS_BORDER, dp(6))\n"
            "    else\n"
            "        roundedBg(NORMAL_BG, NORMAL_BORDER, dp(6))\n"
            "}"
        ),
        "lesson": (
            "TV应用必须为所有可交互元素提供焦点视觉反馈。"
            "这是TV应用开发的基本要求，不能遗漏。"
            "GradientDrawable可以动态创建背景，比XML selector更灵活。"
        ),
    },
    {
        "id": 22,
        "title": "一键部署上线无进度显示，用户以为卡住重复操作",
        "symptom": (
            "点击「一键部署上线」后，界面只显示「正在部署...」，没有任何进度提示。"
            "用户等待几秒后以为程序卡住了，关闭弹窗重新点击，导致重复部署操作。"
        ),
        "cause": (
            "原始实现使用fetch调用同步API /api/v1/deploy-tv，前端等待整个部署完成后才收到响应。"
            "部署过程包含多个步骤（ADB连接→授权→安装APK→配置→启动→注册），每步都可能耗时数秒，"
            "但前端在全部完成前看不到任何进度。"
        ),
        "solution": (
            "将部署API改为SSE（Server-Sent Events）流式推送：\n\n"
            "1. 后端新增 /api/v1/deploy-tv-stream 端点，使用text/event-stream格式（GET请求，兼容EventSource）\n"
            "2. 每完成一个步骤就推送一条SSE消息，包含步骤名称和状态\n"
            "3. 步骤消息类型：adb_connect, adb_auth, adb_ok, install_apk, install_ok, configure, launch, wait_register, done\n"
            "4. 前端使用EventSource消费SSE流，实时更新每个步骤的状态指示器\n"
            "5. ADB授权环节特殊处理：3次重试，每次间隔5秒，推送「等待电视授权」消息\n\n"
            "步骤状态视觉设计：\n"
            "- 等待：灰色圆圈+序号，灰色「等待开始」\n"
            "- 进行中：蓝色圆圈+脉冲动画，蓝色实时描述文字\n"
            "- 完成：绿色圆圈+勾，绿色「完成」\n"
            "- 失败：红色圆圈+叉，红色错误信息\n\n"
            "用户操作提示（黄色提示框）：\n"
            "- ADB等待授权时：请在电视上点击「始终允许」\n"
            "- 安装APP时：请耐心等待约30秒\n"
            "- 失败时：显示错误原因和重试按钮"
        ),
        "lesson": (
            "耗时操作必须提供实时进度反馈，否则用户会认为程序无响应。"
            "SSE比WebSocket更简单，适合服务端单向推送场景。"
            "步骤状态必须有4种明确的视觉区分：等待、进行中、完成、失败，"
            "只高亮当前步骤，不要一开始就把所有步骤都亮起来。"
            "需要用户操作时必须用醒目的提示框告知。"
        ),
    },
    {
        "id": 23,
        "title": "WiFi连接无法识别企业级WiFi，缺少用户名输入",
        "symptom": (
            "WiFi扫描列表中，企业级WiFi（802.1X/EAP）和普通WPA2-PSK WiFi显示一样，"
            "点击企业级WiFi后只弹出密码输入框，但企业级WiFi还需要用户名(Identity)才能连接。"
            "用户只输入密码必然连接失败。"
        ),
        "cause": (
            "原始WiFi扫描代码只检查了capabilities是否包含'WPA'或'WEP'来判断是否加密，"
            "没有区分WPA2-PSK（普通WiFi，只需密码）和WPA2-EAP（企业级WiFi，需用户名+密码）。"
            "ScanResult.capabilities字符串包含完整的加密信息，如：\n"
            "• 普通WPA2: [WPA2-PSK-CCMP][ESS]\n"
            "• 企业级: [WPA2-EAP-CCMP][ESS]\n"
            "关键是是否包含'EAP'字段。"
        ),
        "solution": (
            "1. 扫描结果中检测capabilities是否包含'EAP'来识别企业级WiFi\n"
            "2. WiFi列表中为企业级WiFi显示紫色「企业级」标签，区分普通WiFi\n"
            "3. 点击企业级WiFi时弹出包含用户名(Identity)和密码两个输入框的对话框\n"
            "4. 使用WifiEnterpriseConfig配置PEAP/MSCHAPv2认证，而非WifiConfiguration.preSharedKey\n"
            "5. 普通WPA2 WiFi仍只显示密码输入框"
        ),
        "code_example": (
            "// 检测企业级WiFi\n"
            "private fun isEnterpriseWifi(capabilities: String?): Boolean {\n"
            "    if (capabilities.isNullOrBlank()) return false\n"
            "    return capabilities.contains(\"EAP\")\n"
            "}\n\n"
            "// 企业级WiFi配置\n"
            "val enterpriseConfig = WifiEnterpriseConfig().apply {\n"
            "    setIdentity(identity)\n"
            "    setPassword(pass)\n"
            "    eapMethod = WifiEnterpriseConfig.Eap.PEAP\n"
            "    phase2Method = WifiEnterpriseConfig.Phase2.MSCHAPV2\n"
            "}\n"
            "val wifiConfig = WifiConfiguration().apply {\n"
            "    SSID = \"\\\"$ssid\\\"\"\n"
            "    allowedKeyManagement.set(WifiConfiguration.KeyMgmt.WPA_EAP)\n"
            "    allowedKeyManagement.set(WifiConfiguration.KeyMgmt.IEEE8021X)\n"
            "    this.enterpriseConfig = enterpriseConfig\n"
            "}"
        ),
        "lesson": (
            "WiFi扫描结果中的capabilities字段包含丰富的加密信息，应该充分利用。"
            "企业级WiFi在企业办公场景很常见，必须正确识别和处理。"
            "PEAP/MSCHAPv2是最常见的802.1X认证方式，但不是唯一的，"
            "未来可能需要支持TLS、TTLS等认证方式。"
        ),
    },
    {
        "id": 24,
        "title": "HDMI策略暂停后按键误触设置和密码键盘",
        "severity": "高",
        "symptom": (
            "安卓6电视上，HDMI策略运行时点击「暂停策略」后，系统自动点击了设置按钮，"
            "然后点击了弹出的密码键盘上的数字键。用户只是想暂停策略，结果被带入了设置页面。"
            "安卓9电视上不存在此问题。"
        ),
        "cause": (
            "HDMI调谐后，scheduleConfirmDialog()通过Runtime.exec发送DPAD按键"
            "来自动确认系统弹窗。这些按键是异步执行的系统级按键注入，无法通过"
            "handler.removeCallbacksAndMessages()取消。\n\n"
            "当用户在调谐保护期(5秒)内点击暂停策略时：\n"
            "1. HdmiActivity.onPause()调用handler.removeCallbacksAndMessages()清理待发按键\n"
            "2. 但已经通过Runtime.exec提交给系统的按键无法撤回\n"
            "3. 这些残留的DPAD_CENTER/DPAD_RIGHT按键在HdmiActivity退回主页后继续执行\n"
            "4. 在MainActivity上，DPAD_CENTER正好点击了「设置」按钮\n"
            "5. 后续的DPAD_RIGHT+DPAD_CENTER又点击了密码键盘"
        ),
        "attempts": [
            {"name": "handler.removeCallbacksAndMessages", "desc": "只能取消handler队列中的待发Runnable，已exec的系统按键无法撤回", "success": False},
            {"name": "dispatchKeyEvent拦截", "desc": "在MainActivity中拦截DPAD按键，但系统级按键注入绕过了Activity的dispatchKeyEvent", "success": False},
        ],
        "solution": (
            "三层防护机制：\n\n"
            "1. HdmiActivity添加isLeaving标志：\n"
            "   onPause()时设置isLeaving=true，scheduleConfirmDialog中的每个postDelayed回调\n"
            "   都先检查isLeaving，为true则跳过按键发送。这能阻止handler队列中尚未执行的按键。\n\n"
            "2. MainActivity添加dispatchKeyEvent拦截：\n"
            "   策略暂停后的5秒内(HDMI_PAUSE_KEY_SUPPRESS_MS)，拦截所有DPAD按键。\n"
            "   这能阻止已exec但尚未到达Activity的系统按键。\n\n"
            "3. bringLauncherToFront不再启动HdmiActivity：\n"
            "   原来暂停后回到前台时会重新启动HdmiActivity(然后立即finish)导致新的按键发送，\n"
            "   现在只将MainActivity带到前台，不再触发HdmiActivity生命周期。"
        ),
        "lesson": (
            "Runtime.getRuntime().exec('input keyevent')是异步系统级按键注入，"
            "一旦提交就不可取消。任何依赖定时按键的方案都必须有退出检查机制。"
            "isLeaving标志+handler回调检查是最基本的防护，但对于已经exec的按键，"
            "只能在目标Activity层做拦截。按键注入方案天然不可靠，应尽量减少使用。"
        ),
    },
    {
        "id": 25,
        "title": "HDMI策略切换(HDMI1→HDMI2)时屏幕闪烁黑屏",
        "severity": "高",
        "symptom": (
            "安卓6电视上，当策略从HDMI1切换到HDMI2时，屏幕会闪烁黑屏2-3次，"
            "最终停留在HDMI1画面，无法成功切换到HDMI2。HDMI3也一样。"
            "安卓9电视上不存在此问题。"
        ),
        "cause": (
            "策略切换时存在3个触发源导致竞争：\n\n"
            "1. MainActivity.onResume()：策略切换后回到前台时，读取旧策略值执行切换\n"
            "2. policyUpdateReceiver第一次回调：RemoteApi心跳同步到新策略后广播通知\n"
            "3. policyUpdateReceiver第二次回调：RemoteApi主动触发同步后的广播\n\n"
            "问题1：onResume使用旧的SharedPreferences值（心跳还未更新），切换到旧端口HDMI1\n"
            "问题2：HdmiActivity的onNewIntent使用inputId比较端口，但安卓6 Amlogic只有一个HDMI输入，"
            "所有端口对应同一个inputId，导致端口变更检测失效\n\n"
            "最终效果：HDMI1→HDMI2→HDMI1→HDMI2反复闪烁，由于onResume先执行旧策略，"
            "最终画面停留在HDMI1。"
        ),
        "solution": (
            "1. onNewIntent改用端口号比较：不再从inputId提取端口（因为安卓6只有一个inputId），\n"
            "   而是直接比较intent中的EXTRA_HDMI_PORT，正确检测策略端口变更。\n\n"
            "2. forceExecutePolicy中检查hdmi_foreground标志：\n"
            "   如果HdmiActivity已在前台，发送onNewIntent复用singleTask实例，\n"
            "   而不是创建新的HdmiActivity实例。\n\n"
            "3. LauncherExecutor添加FLAG_ACTIVITY_CLEAR_TOP|FLAG_ACTIVITY_SINGLE_TOP：\n"
            "   确保HdmiActivity被复用而不是堆叠。"
        ),
        "lesson": (
            "安卓6 Amlogic芯片只有1个HDMI InputService(Hdmi1InputService/HW5)，"
            "但物理上有3个HDMI端口。所有端口切换都用同一个inputId，"
            "不能通过inputId判断端口是否变化，必须用端口号比较。"
            "多个触发源同时执行策略时必须做好去重和优先级控制。"
        ),
    },
    {
        "id": 26,
        "title": "HDMI线拔插后黑屏不恢复",
        "severity": "高",
        "symptom": (
            "安卓6电视上，策略为HDMI2时，拔掉HDMI线再插回同一HDMI2口，画面黑屏不恢复。"
            "更复杂的场景：拔掉HDMI2线，插到HDMI1口，再拔出插回HDMI2口，画面也是黑屏。"
            "安卓9电视上不存在此问题。"
        ),
        "cause": (
            "原方案使用BroadcastReceiver监听android.intent.action.HDMI_PLUGGED广播，"
            "但安卓6 Amlogic设备上该广播从不触发（与问题1相同的原因）。\n\n"
            "拔线后hasTuned仍为true，scheduleTune()跳过调谐（因为hasTuned检查），"
            "HdmiActivity始终显示黑屏。插线后没有收到广播，无法触发重新调谐。\n\n"
            "日志验证：拔插HDMI线时，logcat中完全看不到HDMI_PLUGGED相关日志，"
            "但TvInputManagerService的onInputAdded/onInputRemoved/onInputStateChanged正常触发。"
        ),
        "solution": (
            "使用TvInputManager.TvInputCallback替代BroadcastReceiver：\n\n"
            "• onInputAdded：HDMI输入添加时，根据targetPort重新解析inputId并调谐\n"
            "  （拔出HDMI2插入HDMI1时，系统移除Hdmi2InputService添加Hdmi1InputService，\n"
            "  但策略targetPort=2，findInputIdForPort会查找当前存在的HDMI2输入，不存在则等待）\n\n"
            "• onInputRemoved：当前HDMI输入被移除时，重置hasTuned=false和currentInputId=null，\n"
            "  为后续onInputAdded重新调谐做准备\n\n"
            "• onInputStateChanged：HDMI信号恢复(state=0)时，如果hasTuned=false则调度调谐\n\n"
            "添加targetPort成员变量：跟踪策略指定的端口号，在onCreate和onNewIntent中更新，\n"
            "确保TvInputCallback回调时能正确解析目标端口的inputId。"
        ),
        "lesson": (
            "android.intent.action.HDMI_PLUGGED广播在不同芯片/系统版本上可靠性差异巨大。"
            "TvInputManager.TvInputCallback是Android TV标准API，在所有支持TV Input Framework的"
            "设备上都可靠工作。应该优先使用框架API而非系统广播。\n\n"
            "关键发现：安卓6 Amlogic设备上，HDMI线从HDMI2口拔出插入HDMI1口时，\n"
            "系统会移除Hdmi2InputService/HW6并添加Hdmi1InputService/HW5，\n"
            "这是动态注册/注销的，不是固定的。所以必须动态查找inputId，不能缓存。"
        ),
    },
    {
        "id": 27,
        "title": "管理后台更新/暂停按钮无反馈导致重复点击",
        "severity": "中",
        "symptom": (
            "管理后台点击「更新」或「暂停」按钮后，界面没有即时反馈，用户以为没有点击成功，"
            "反复点击导致弹出多个alert弹窗，操作混乱。"
        ),
        "cause": (
            "bindPolicy()和togglePolicyPause()函数中，点击按钮后直接发起fetch请求，\n"
            "没有禁用按钮也没有显示loading状态。网络请求可能需要1-2秒，\n"
            "期间按钮仍可点击，用户会重复提交请求。"
        ),
        "solution": (
            "为两个按钮添加loading状态和防重复点击：\n\n"
            "1. 更新按钮：点击后文字变为「下发中...」，按钮disabled=true\n"
            "2. 暂停/继续按钮：点击后文字变为「暂停中...」/「继续中...」，按钮disabled=true\n"
            "3. 请求完成（成功或失败）后恢复按钮原文字和可用状态\n"
            "4. 失败时恢复按钮原始文字，避免状态错乱"
        ),
        "lesson": (
            "所有涉及网络请求的按钮都必须有loading状态和防重复点击机制。"
            "用户无法感知后台操作进度时，会本能地重复点击。"
            "按钮disabled+文字变更是最简单有效的反馈方式。"
        ),
    },
    {
        "id": 28,
        "title": "HDMI策略按HOME/BACK后不立即恢复",
        "severity": "高",
        "symptom": (
            "用户在HDMI策略运行时误按HOME或BACK键回到主页，期望APP立即恢复到HDMI策略，"
            "但实际上需要等待3-5秒甚至更久才能恢复。\n\n"
            "更严重的是：连续按多次HOME键，恢复越来越慢，用户体验极差。\n"
            "安卓6和安卓9电视都有此问题。"
        ),
        "cause": (
            "策略执行后设置了30秒冷却期（HDMI_SWITCH_COOLDOWN）和5秒防抖（EXECUTION_COOLDOWN）。\n\n"
            "用户按HOME/BACK后，MainActivity.onResume被调用，触发保活检查执行策略恢复。\n"
            "但由于冷却期检查，策略恢复被阻止：\n\n"
            "• MainActivity.onResume中的策略执行被冷却期跳过\n"
            "• forceExecutePolicy中的冷却期检查阻止恢复\n"
            "• 保活检查中的冷却期阻止立即恢复\n\n"
            "只能等待KeepAliveForegroundService的5秒检查周期来恢复，导致明显延迟。"
        ),
        "solution": (
            "HDMI模式下，当检测到hdmi_foreground=false（HdmiActivity不在前台）时，\n"
            "绕过所有冷却期限制，立即恢复策略：\n\n"
            "1. MainActivity.onResume：\n"
            "   if (policy.mode == \"hdmi\" && !hdmiFg) { 立即恢复，绕过5秒防抖和30秒冷却期 }\n\n"
            "2. forceExecutePolicy：\n"
            "   hdmiNeedsRecovery = policy.mode == \"hdmi\" && !hdmiFg\n"
            "   if (hdmiNeedsRecovery) { 绕过防抖和冷却期 }\n\n"
            "3. 保活检查：\n"
            "   if (!hdmiFg && !policyStore.isPolicyPaused()) { 直接恢复，不检查冷却期 }\n\n"
            "关键思路：冷却期的目的是防止策略切换时重复启动HdmiActivity导致闪烁。\n"
            "但用户按HOME/BACK离开HdmiActivity时，必须立即拉回，不应受冷却期限制。"
        ),
        "lesson": (
            "冷却期机制需要区分场景：\n"
            "• 策略切换中（hdmi_foreground=true）：需要冷却期防止重复启动\n"
            "• 用户误离开（hdmi_foreground=false）：需要立即恢复，绕过冷却期\n\n"
            "hdmi_foreground SharedPreferences标志是关键的状态判断依据，"
            "由HdmiActivity.onResume/onPause设置，其他组件可以可靠地判断HdmiActivity是否在前台。"
        ),
    },
]

# ========== 生成文档 ==========

def generate():
    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ========== 封面 ==========
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("TV Launcher 项目")
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor.from_string(COLOR_PRIMARY)
    run.font.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("问题与解决方案记录")
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor.from_string(COLOR_SECONDARY)
    run.font.bold = True

    doc.add_paragraph()

    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = desc.add_run("记录开发过程中遇到的所有技术难题、排查过程和最终解决方案")
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor.from_string(COLOR_TEXT_LIGHT)

    doc.add_paragraph()
    doc.add_paragraph()

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run("最后更新：2026年5月3日")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string(COLOR_TEXT_LIGHT)

    # 分页
    doc.add_page_break()

    # ========== 目录 ==========
    add_styled_heading(doc, "目录", level=1)
    doc.add_paragraph()

    for p in PROBLEMS:
        para = doc.add_paragraph()
        run = para.add_run(f"问题{p['id']}：{p['title']}")
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor.from_string(COLOR_SECONDARY)

    doc.add_page_break()

    # ========== 问题详情 ==========
    add_styled_heading(doc, "问题与解决方案详情", level=1)
    doc.add_paragraph()

    # 概览表格
    add_body_text(doc, "以下是所有已解决问题的概览：", bold=True)
    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    headers = ["问题编号", "问题标题", "严重程度", "状态"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor.from_string("FFFFFF")
        set_cell_shading(cell, COLOR_SECONDARY)

    for p in PROBLEMS:
        row = table.add_row()
        row.cells[0].text = str(p['id'])
        row.cells[1].text = p['title']
        severity = p.get('severity', '中')
        row.cells[2].text = severity
        row.cells[3].text = "已解决"

        for i in range(4):
            for para in row.cells[i].paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.size = Pt(10)

        # 状态列绿色
        set_cell_shading(row.cells[3], COLOR_BG_GREEN)

    doc.add_page_break()

    # 每个问题的详细描述
    for p in PROBLEMS:
        add_problem_section(doc, p)

    # ========== 附录：技术参考 ==========
    doc.add_page_break()
    add_styled_heading(doc, "附录：技术参考信息", level=1)

    add_styled_heading(doc, "A. Amlogic/Droidlogic HDMI输入服务ID对照表", level=2)

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["HDMI端口", "Input Service ID", "硬件编号", "TvInputInfo.type"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p2 in cell.paragraphs:
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p2.runs:
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor.from_string("FFFFFF")
        set_cell_shading(cell, COLOR_SECONDARY)

    hdmi_data = [
        ("HDMI1", "com.droidlogic.tvinput/.services.Hdmi1InputService", "HW5", "TYPE_HDMI (1007)"),
        ("HDMI2", "com.droidlogic.tvinput/.services.Hdmi2InputService", "HW6", "TYPE_HDMI (1007)"),
        ("HDMI3", "com.droidlogic.tvinput/.services.Hdmi3InputService", "HW7", "TYPE_HDMI (1007)"),
    ]
    for row_data in hdmi_data:
        row = table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = val
            for para in row.cells[i].paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()
    add_styled_heading(doc, "B. 关键系统属性和路径", level=2)

    props = [
        ("mbx.hdmiin.switchfull", "HDMI是否全屏显示（true/false），需要root才能修改"),
        ("mbx.hdmiin.videolayer", "HDMI视频层是否启用（true/false）"),
        ("ro.sys.hdmiin.enable", "系统是否支持HDMI输入（只读）"),
        ("/sys/class/switch/hdmirx_hpd/state", "HDMI热插拔状态（0=未接，1=已接），安卓6上不可靠"),
    ]
    for name, desc in props:
        p = doc.add_paragraph()
        run_name = p.add_run(f"  • {name}")
        run_name.font.size = Pt(10)
        run_name.font.bold = True
        run_name.font.name = "Consolas"
        run_name.font.color.rgb = RGBColor.from_string(COLOR_SECONDARY)
        run_desc = p.add_run(f" — {desc}")
        run_desc.font.size = Pt(10)
        run_desc.font.color.rgb = RGBColor.from_string(COLOR_TEXT)

    doc.add_paragraph()
    add_styled_heading(doc, "C. 测试命令", level=2)

    add_body_text(doc, "通过ADB模拟HDMI插拔广播（安卓9有效）：", bold=True)
    add_code_block(doc, "# 模拟HDMI插入\nadb shell am broadcast -a android.intent.action.HDMI_PLUGGED --ez state true\n\n# 模拟HDMI拔出\nadb shell am broadcast -a android.intent.action.HDMI_PLUGGED --ez state false")

    add_body_text(doc, "查看当前注册的TV输入设备：", bold=True)
    add_code_block(doc, "adb shell dumpsys tv_input")

    add_body_text(doc, "查看HDMI信号状态：", bold=True)
    add_code_block(doc, "adb shell logcat | grep -E 'onSigToStable|TVIN_SIG_FMT|signal_info'")

    # 保存
    out_path = os.path.join(OUT_DIR, "04-问题与解决方案记录.docx")
    doc.save(out_path)
    print(f"文档已生成: {out_path}")

if __name__ == "__main__":
    generate()
