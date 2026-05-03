"""生成 TV Launcher 项目专业文档 (.docx) — 深度美化版"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement

OUT_DIR = os.path.dirname(os.path.abspath(__file__))
os.makedirs(OUT_DIR, exist_ok=True)

# ========== 色彩体系 ==========
COLOR_PRIMARY = "1B3A5C"       # 深蓝 - 一级标题
COLOR_SECONDARY = "2E75B6"     # 中蓝 - 二级标题/表头
COLOR_ACCENT = "4A90D9"        # 亮蓝 - 三级标题/强调
COLOR_BG_LIGHT = "F2F7FC"      # 浅蓝背景
COLOR_BG_GRAY = "F5F6F8"       # 浅灰斑马纹
COLOR_TEXT = "2D3436"           # 正文深灰
COLOR_TEXT_LIGHT = "636E72"    # 辅助文字
COLOR_BORDER = "D5DDE5"        # 表格边框
COLOR_COVER_BAR = "1B3A5C"     # 封面装饰条
COLOR_PRIORITY_P0 = "E74C3C"   # 红色
COLOR_PRIORITY_P1 = "F39C12"   # 橙色
COLOR_PRIORITY_P2 = "27AE60"   # 绿色

# ========== 通用样式工具 ==========

def set_cell_shading(cell, color_hex):
    """设置单元格底色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_border(cell, **kwargs):
    """设置单元格边框
    kwargs: top=("single","4","1B3A5C"), bottom=..., left=..., right=...
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge, (style, sz, color) in kwargs.items():
        element = tcBorders.find(qn(f'w:{edge}'))
        if element is None:
            element = OxmlElement(f'w:{edge}')
            tcBorders.append(element)
        element.set(qn('w:val'), style)
        element.set(qn('w:sz'), sz)
        element.set(qn('w:color'), color)
        element.set(qn('w:space'), '0')

def set_cell_margins(cell, top=40, bottom=40, left=80, right=80):
    """设置单元格内边距 (单位: twips)"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.find(qn('w:tcMar'))
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)
    for edge, val in [("top", top), ("bottom", bottom), ("start", left), ("end", right)]:
        element = tcMar.find(qn(f'w:{edge}'))
        if element is None:
            element = OxmlElement(f'w:{edge}')
            tcMar.append(element)
        element.set(qn('w:w'), str(val))
        element.set(qn('w:type'), 'dxa')

def set_run_font(run, name_cn='微软雅黑', name_en='Calibri', size=Pt(10.5), color=None, bold=False):
    """统一设置 run 的字体"""
    run.font.size = size
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    run.font.name = name_en
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), name_cn)

def add_horizontal_line(doc, color=COLOR_PRIMARY, width=2):
    """添加水平分隔线"""
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(width * 4))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_page_header_footer(doc):
    """添加页眉页脚"""
    section = doc.sections[0]
    # 页眉
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run("Meeting TV Launcher")
    set_run_font(run, size=Pt(8), color=COLOR_TEXT_LIGHT)
    # 页眉下方加线
    pPr = hp._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), COLOR_BORDER)
    pBdr.append(bottom)
    pPr.append(pBdr)
    # 页脚 - 页码
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = fp.add_run("— ")
    set_run_font(run1, size=Pt(9), color=COLOR_TEXT_LIGHT)
    # 页码字段
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run2 = fp.add_run()
    run2._element.append(fldChar1)
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    run3 = fp.add_run()
    run3._element.append(instrText)
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run4 = fp.add_run()
    run4._element.append(fldChar2)
    run5 = fp.add_run(" —")
    set_run_font(run5, size=Pt(9), color=COLOR_TEXT_LIGHT)

def setup_page(doc):
    """设置页面布局"""
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    add_page_header_footer(doc)

def setup_styles(doc):
    """设置全局样式"""
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)
    style.font.color.rgb = RGBColor.from_string(COLOR_TEXT)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    pf = style.paragraph_format
    pf.space_before = Pt(3)
    pf.space_after = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.25

    # 一级标题样式
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Calibri'
    h1.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    h1.font.size = Pt(18)
    h1.font.color.rgb = RGBColor.from_string(COLOR_PRIMARY)
    h1.font.bold = True
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(8)
    # 一级标题底部边框
    pPr = h1.element.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        h1.element.insert(0, pPr)
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), COLOR_SECONDARY)
    pBdr.append(bottom)
    pPr.append(pBdr)

    # 二级标题样式
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Calibri'
    h2.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    h2.font.size = Pt(14)
    h2.font.color.rgb = RGBColor.from_string(COLOR_SECONDARY)
    h2.font.bold = True
    h2.paragraph_format.space_before = Pt(16)
    h2.paragraph_format.space_after = Pt(6)

    # 三级标题样式
    h3 = doc.styles['Heading 3']
    h3.font.name = 'Calibri'
    h3.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    h3.font.size = Pt(12)
    h3.font.color.rgb = RGBColor.from_string(COLOR_ACCENT)
    h3.font.bold = True
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(4)

def add_cover(doc, title, subtitle, doc_id, version, date):
    """专业封面页"""
    # 顶部装饰条
    for _ in range(2):
        doc.add_paragraph()

    # 装饰色块
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0)
    cell.text = ""
    set_cell_shading(cell, COLOR_COVER_BAR)
    # 设置色块高度
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("  ")
    run.font.size = Pt(8)
    # 去掉色块边框
    for edge in ['top', 'bottom', 'left', 'right']:
        set_cell_border(cell, **{edge: ("none", "0", "FFFFFF")})

    # 标题
    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    set_run_font(run, size=Pt(32), color=COLOR_PRIMARY, bold=True, name_cn='微软雅黑')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    set_run_font(run, size=Pt(14), color=COLOR_TEXT_LIGHT, name_cn='微软雅黑')

    # 分隔线
    for _ in range(1):
        doc.add_paragraph()
    add_horizontal_line(doc, COLOR_SECONDARY, width=1)
    for _ in range(2):
        doc.add_paragraph()

    # 信息卡片
    info_table = doc.add_table(rows=5, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_data = [
        ("文档编号", doc_id),
        ("版本号", version),
        ("编写日期", date),
        ("密级", "内部"),
        ("状态", "已发布"),
    ]
    for i, (k, v) in enumerate(info_data):
        cell_k = info_table.cell(i, 0)
        cell_v = info_table.cell(i, 1)
        # 清空
        cell_k.paragraphs[0].clear()
        cell_v.paragraphs[0].clear()
        # key
        p = cell_k.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(k)
        set_run_font(run, size=Pt(11), color=COLOR_TEXT_LIGHT, bold=True)
        set_cell_shading(cell_k, COLOR_BG_LIGHT)
        set_cell_margins(cell_k, top=50, bottom=50, left=120, right=120)
        # value
        p = cell_v.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(v)
        set_run_font(run, size=Pt(11), color=COLOR_PRIMARY, bold=True)
        set_cell_margins(cell_v, top=50, bottom=50, left=120, right=120)
        # 边框
        for cell in [cell_k, cell_v]:
            set_cell_border(cell,
                top=("single", "4", COLOR_BORDER),
                bottom=("single", "4", COLOR_BORDER),
                left=("single", "4", COLOR_BORDER),
                right=("single", "4", COLOR_BORDER))

    # 设置信息表列宽
    for row in info_table.rows:
        row.cells[0].width = Cm(3.5)
        row.cells[1].width = Cm(7)

    # 底部装饰条
    for _ in range(4):
        doc.add_paragraph()
    t2 = doc.add_table(rows=1, cols=1)
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell2 = t2.cell(0, 0)
    cell2.text = ""
    set_cell_shading(cell2, COLOR_COVER_BAR)
    p2 = cell2.paragraphs[0]
    p2.paragraph_format.space_before = Pt(2)
    p2.paragraph_format.space_after = Pt(2)
    run2 = p2.add_run("  ")
    run2.font.size = Pt(4)
    for edge in ['top', 'bottom', 'left', 'right']:
        set_cell_border(cell2, **{edge: ("none", "0", "FFFFFF")})

    doc.add_page_break()

def add_revision_table(doc):
    """修订记录表"""
    doc.add_heading("文档修订记录", level=1)
    t = doc.add_table(rows=2, cols=4, style='Table Grid')
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["版本", "日期", "修订人", "修订内容"]
    for i, h in enumerate(headers):
        cell = t.cell(0, i)
        cell.paragraphs[0].clear()
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run_font(run, size=Pt(10), color="FFFFFF", bold=True)
        set_cell_shading(cell, COLOR_SECONDARY)
        set_cell_margins(cell, top=50, bottom=50)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    row = t.rows[1]
    data = ["V1.0", "2026-04-23", "\u2014", "初版发布"]
    for i, v in enumerate(data):
        cell = row.cells[i]
        cell.paragraphs[0].clear()
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(v)
        set_run_font(run, size=Pt(10), color=COLOR_TEXT)
        set_cell_margins(cell, top=50, bottom=50)
    doc.add_paragraph()

def make_table(doc, headers, rows, col_widths=None, highlight_col=None, priority_col=None):
    """创建专业格式化表格
    highlight_col: 需要特殊高亮的列索引列表（如编号列）
    priority_col: 优先级列索引，自动着色 P0/P1/P2
    """
    t = doc.add_table(rows=1 + len(rows), cols=len(headers), style='Table Grid')
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for i, h in enumerate(headers):
        cell = t.cell(0, i)
        cell.paragraphs[0].clear()
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run_font(run, size=Pt(9.5), color="FFFFFF", bold=True)
        set_cell_shading(cell, COLOR_SECONDARY)
        set_cell_margins(cell, top=55, bottom=55, left=60, right=60)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # 数据行
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = t.cell(r_idx + 1, c_idx)
            cell.paragraphs[0].clear()
            p = cell.paragraphs[0]

            # 优先级列特殊着色
            is_priority = priority_col is not None and c_idx == priority_col
            is_highlight = highlight_col is not None and c_idx in highlight_col

            if is_priority and str(val).startswith("P"):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                pcolor = {
                    "P0": COLOR_PRIORITY_P0,
                    "P1": COLOR_PRIORITY_P1,
                    "P2": COLOR_PRIORITY_P2,
                }.get(str(val), COLOR_TEXT)
                run = p.add_run(str(val))
                set_run_font(run, size=Pt(9.5), color=pcolor, bold=True)
            elif is_highlight:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(str(val))
                set_run_font(run, size=Pt(9.5), color=COLOR_SECONDARY, bold=True)
            else:
                run = p.add_run(str(val))
                set_run_font(run, size=Pt(9.5), color=COLOR_TEXT)

            set_cell_margins(cell, top=45, bottom=45, left=60, right=60)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            # 斑马纹
            if r_idx % 2 == 1:
                set_cell_shading(cell, COLOR_BG_GRAY)

            # 边框
            set_cell_border(cell,
                top=("single", "4", COLOR_BORDER),
                bottom=("single", "4", COLOR_BORDER),
                left=("single", "4", COLOR_BORDER),
                right=("single", "4", COLOR_BORDER))

    # 表头边框
    for i in range(len(headers)):
        cell = t.cell(0, i)
        set_cell_border(cell,
            top=("single", "8", COLOR_SECONDARY),
            bottom=("single", "6", COLOR_SECONDARY),
            left=("single", "4", COLOR_SECONDARY),
            right=("single", "4", COLOR_SECONDARY))

    # 列宽
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)

    doc.add_paragraph()
    return t

def add_prototype_box(doc, text, title=None):
    """添加原型图框 - 带标题和灰色背景的等宽字体区域"""
    if title:
        p = doc.add_paragraph()
        run = p.add_run(title)
        set_run_font(run, size=Pt(9.5), color=COLOR_SECONDARY, bold=True)

    # 用单列表格模拟代码框
    lines = text.strip().split('\n')
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0)
    cell.paragraphs[0].clear()
    set_cell_shading(cell, "F7F8FA")
    set_cell_margins(cell, top=80, bottom=80, left=120, right=120)

    for i, line in enumerate(lines):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(COLOR_TEXT)
        r = run._element
        rPr = r.find(qn('w:rPr'))
        if rPr is None:
            rPr = OxmlElement('w:rPr')
            r.insert(0, rPr)
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), 'Consolas')

    # 原型框边框
    set_cell_border(cell,
        top=("single", "6", COLOR_BORDER),
        bottom=("single", "6", COLOR_BORDER),
        left=("single", "6", COLOR_BORDER),
        right=("single", "6", COLOR_BORDER))

    doc.add_paragraph()

def add_flow_steps(doc, steps):
    """添加流程步骤 - 带编号圆圈"""
    for i, step in enumerate(steps, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Cm(0.5)
        # 编号
        run = p.add_run(f"  {i}  ")
        set_run_font(run, size=Pt(9.5), color="FFFFFF", bold=True)
        # 给编号加圆形底色（用shading模拟）
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:fill'), COLOR_SECONDARY)
        run._element.get_or_add_rPr().append(shd)
        # 步骤文字
        run2 = p.add_run(f"  {step}")
        set_run_font(run2, size=Pt(10.5), color=COLOR_TEXT)
    doc.add_paragraph()

def add_body(doc, text):
    """添加正文段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=Pt(10.5), color=COLOR_TEXT)
    p.paragraph_format.first_line_indent = Cm(0.74)
    return p

def add_toc(doc, items):
    """添加目录"""
    doc.add_heading("目录", level=1)
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        # 分割标题和页码
        parts = item.split("  ", 1)
        if len(parts) == 2:
            run = p.add_run(parts[0])
            set_run_font(run, size=Pt(11), color=COLOR_SECONDARY)
            run2 = p.add_run(f"  {'.' * 40}  {parts[1]}")
            set_run_font(run2, size=Pt(9), color=COLOR_TEXT_LIGHT)
        else:
            run = p.add_run(item)
            set_run_font(run, size=Pt(11), color=COLOR_SECONDARY)
    doc.add_page_break()


# ================================================================
# 文档1：需求规格说明书
# ================================================================
def create_requirement_doc():
    doc = Document()
    setup_styles(doc)
    setup_page(doc)

    add_cover(doc, "需求规格说明书", "Meeting TV Launcher\n会议室电视集中管控系统", "TVL-REQ-001", "V1.0", "2026-04-23")
    add_revision_table(doc)

    add_toc(doc, [
        "1  引言  3",
        "2  项目概述  4",
        "3  用户角色  5",
        "4  功能需求  6",
        "5  非功能需求  10",
        "6  约束与假设  11",
        "7  术语表  12",
    ])

    # === 1 引言 ===
    doc.add_heading("1  引言", level=1)
    doc.add_heading("1.1  编写目的", level=2)
    add_body(doc, "本文档定义 Meeting TV Launcher 系统的功能需求和非功能需求，作为系统设计、开发、测试的依据。")
    doc.add_heading("1.2  项目背景", level=2)
    add_body(doc, "企业会议室配备小米智能电视，需要统一管控：开机后自动进入指定应用（投屏软件、视频会议等）或切换到指定 HDMI 信号源，防止被随意切换或关闭。管理员需通过 Web 后台集中管理所有电视，远程下发策略、监控状态、排查故障。")
    doc.add_heading("1.2.1  现状与痛点", level=3)
    add_body(doc, "当前企业会议室电视管理面临以下突出问题：")
    make_table(doc,
        ["痛点", "具体表现", "影响"],
        [
            ["人工操作低效", "每台电视需逐台手动配置投屏软件或 HDMI 信号源", "新电视上线耗时 30 分钟以上，人为操作容易遗漏"],
            ["状态不可见", "无法实时掌握电视在线/离线状态和运行情况", "故障发现滞后，只能等参会人员现场报修"],
            ["故障响应慢", "IT 人员必须到会议室现场排查问题", "会议中断等待时间长，影响业务效率"],
            ["缺乏统一管控", "不同会议室电视配置不一致，策略执行标准不统一", "会议室用途变更后无法快速切换策略"],
            ["安全隐患", "电视可被随意切换应用或关闭管控软件", "重要会议期间电视被误操作，影响会议正常进行"],
        ], col_widths=[2.5, 5.5, 5.5])
    doc.add_heading("1.2.2  项目价值与意义", level=3)
    add_body(doc, "本项目通过技术手段实现企业会议室电视的集中化、自动化、远程化管理，具有以下核心价值：")
    make_table(doc,
        ["价值维度", "说明", "量化收益"],
        [
            ["运维效率提升", "从逐台手动配置升级为一键部署和集中策略下发，新电视上线时间从 30 分钟缩短至 3 分钟", "运维效率提升 10 倍"],
            ["故障响应提速", "通过 Scrcpy 实时投屏和 ADB 远程命令，IT 人员无需到场即可排查和解决问题", "故障平均响应时间从 30 分钟降至 5 分钟"],
            ["管理标准化", "统一策略模板，所有电视自动执行标准配置，杜绝人为操作差异", "配置不一致问题归零"],
            ["实时可观测性", "Web 后台实时展示所有电视的在线状态、网络信息、资源占用和策略执行情况", "故障发现从被动等待变为主动监控"],
            ["安全可控性", "Kiosk 模式锁定电视，防止非授权操作；ADB 命令安全过滤，防止误操作", "会议中误操作风险降至最低"],
            ["规模化扩展", "单后台可管理数十台电视，新增设备零配置自动注册上线", "管理规模从几台扩展到几十台无需增加人力"],
        ], col_widths=[2.5, 7, 4])
    add_body(doc, "本项目的实施将显著降低企业 IT 运维的人力成本和时间成本，提升会议室电视管理的规范性和可靠性，为企业数字化转型提供基础设施层面的有力支撑。")
    doc.add_heading("1.3  适用范围", level=2)
    add_body(doc, "本文档适用于 Meeting TV Launcher 系统的后端服务、Web 管理前端、Android TV 客户端三个子系统。")

    # === 2 项目概述 ===
    doc.add_heading("2  项目概述", level=1)
    doc.add_heading("2.1  系统目标", level=2)
    add_body(doc, "实现企业会议室电视的集中管控，包括策略自动执行、远程监控、远程控制、一键部署四大核心能力。")
    doc.add_heading("2.2  系统组成", level=2)
    make_table(doc,
        ["子系统", "说明", "技术栈"],
        [
            ["后端服务", "提供 REST API 和 Web 管理页面", "Python / FastAPI / SQLite"],
            ["Web 管理前端", "管理员操作界面", "HTML / CSS / JavaScript"],
            ["Android TV 客户端", "电视端策略执行引擎", "Kotlin / Android SDK"],
        ], col_widths=[3.5, 5, 5])

    # === 3 用户角色 ===
    doc.add_heading("3  用户角色", level=1)
    make_table(doc,
        ["角色编号", "角色名称", "描述", "核心场景"],
        [
            ["U01", "系统管理员", "IT 运维人员，负责设备管理和策略配置", "日常管理、故障排查、新设备上线"],
            ["U02", "电视设备", "安装 Launcher APP 的小米电视", "自动注册、心跳上报、策略执行"],
        ], highlight_col=[0], col_widths=[2, 2.5, 5, 4])

    # === 4 功能需求 ===
    doc.add_heading("4  功能需求", level=1)

    doc.add_heading("4.1  设备管理", level=2)
    make_table(doc,
        ["需求编号", "需求名称", "优先级", "需求描述", "验收标准"],
        [
            ["FR-DM-001", "设备自注册", "P0", "电视安装 APP 后首次启动自动注册到后台", "APP 启动后 30 秒内出现在设备列表"],
            ["FR-DM-002", "设备心跳上报", "P0", "设备定时上报网络信息、资源使用、在线状态", "心跳间隔 \u2264 15 分钟，后台实时显示在线/离线"],
            ["FR-DM-003", "设备列表展示", "P0", "后台以表格形式展示所有设备信息", "表格含：会议室、设备名、型号、网络、状态、策略、控制、操作"],
            ["FR-DM-004", "列筛选", "P1", "每列表头支持筛选，类似 Excel", "点击筛选图标弹出复选框列表，支持全选/搜索"],
            ["FR-DM-005", "列排序", "P1", "每列表头支持升序/降序排序", "点击列标题循环切换：升序\u2192降序\u2192取消"],
            ["FR-DM-006", "全局搜索", "P1", "搜索框可搜索设备所有信息", "输入关键词实时过滤，匹配设备名/SN/IP/MAC/型号/应用/策略"],
            ["FR-DM-007", "会议室标签", "P2", "设备显示所属会议室，可编辑", "橙色标签显示会议室名，点击编辑按钮修改"],
            ["FR-DM-008", "设备注销", "P1", "移除设备同时清理电视端 APP 数据", "删除数据库记录 + ADB 清数据 + 停止 APP"],
        ], highlight_col=[0], priority_col=2, col_widths=[2.2, 2, 1.2, 4, 4])

    doc.add_heading("4.2  策略管理", level=2)
    make_table(doc,
        ["需求编号", "需求名称", "优先级", "需求描述", "验收标准"],
        [
            ["FR-PM-001", "创建策略", "P0", "支持创建 APP 启动策略和 HDMI 切换策略", "策略含名称、模式(APP/HDMI)、目标参数"],
            ["FR-PM-002", "删除策略", "P0", "删除不需要的策略", "确认后删除，已绑定的设备策略清空"],
            ["FR-PM-003", "策略绑定", "P0", "将策略绑定到指定设备", "绑定后通过 ADB 推送，设备 10 秒内开始执行"],
            ["FR-PM-004", "策略暂停", "P1", "暂停策略执行", '设备停止自动执行，标签变为\u201c已暂停\u201d'],
            ["FR-PM-005", "策略恢复", "P1", "恢复暂停的策略", '设备恢复执行，标签变为\u201c运行中\u201d'],
            ["FR-PM-006", "策略自动执行", "P0", "设备开机/重启后自动执行绑定的策略", "开机后立即显示深色界面（无白屏闪烁），1.5秒内执行目标策略"],
        ], highlight_col=[0], priority_col=2, col_widths=[2.2, 2, 1.2, 4, 4])

    doc.add_heading("4.3  远程控制", level=2)
    make_table(doc,
        ["需求编号", "需求名称", "优先级", "需求描述", "验收标准"],
        [
            ["FR-RC-001", "Scrcpy 投屏", "P1", "通过 Scrcpy 实现低延迟实时屏幕镜像和交互控制", "点击启动后弹出 Scrcpy 窗口，延迟 < 200ms"],
            ["FR-RC-002", "ADB 遥控器", "P1", "Web 端提供遥控器按钮", "点击按钮发送 ADB keyevent，电视 2 秒内响应"],
            ["FR-RC-003", "ADB 命令执行", "P1", "执行任意 ADB shell 命令", "支持命令输入、常用命令列表、输出显示"],
            ["FR-RC-004", "定时执行", "P2", "周期性执行 ADB 命令", "可设间隔和重复次数，实时显示结果"],
            ["FR-RC-005", "危险命令过滤", "P1", "禁止执行破坏性命令", "阻止 rm -rf /、format、mkfs、dd if=、shutdown"],
            ["FR-RC-006", "ADB 连接管理", "P1", "建立/断开 ADB 无线连接", "支持 connect/disconnect 操作"],
        ], highlight_col=[0], priority_col=2, col_widths=[2.2, 2, 1.2, 4, 4])

    doc.add_heading("4.4  部署与安装", level=2)
    make_table(doc,
        ["需求编号", "需求名称", "优先级", "需求描述", "验收标准"],
        [
            ["FR-DE-001", "一键部署", "P1", "输入 IP 自动完成全流程部署", "自动：ADB 连接\u2192安装 APP\u2192配置\u2192注册\u2192绑定策略"],
            ["FR-DE-002", "APK 远程安装", "P1", "通过 ADB 安装 APK 到电视", "支持安装内置 APK 和上传的 APK"],
            ["FR-DE-003", "应用管理", "P1", "列出已安装应用，支持卸载", "显示第三方应用列表，一键卸载"],
        ], highlight_col=[0], priority_col=2, col_widths=[2.2, 2, 1.2, 4, 4])

    doc.add_heading("4.5  其他功能", level=2)
    make_table(doc,
        ["需求编号", "需求名称", "优先级", "需求描述", "验收标准"],
        [
            ["FR-OT-001", "操作日志", "P2", "记录所有管理操作", "记录操作类型、时间、设备、详情"],
            ["FR-OT-002", "OTA 更新", "P2", "推送 APP 更新到电视", "后台配置新版本，电视自动下载安装"],
            ["FR-OT-003", "Kiosk 模式", "P2", "锁定电视不允许退出 Launcher", "拦截 BACK/HOME，三击 HOME 临时退出 10 分钟"],
        ], highlight_col=[0], priority_col=2, col_widths=[2.2, 2, 1.2, 4, 4])

    # === 5 非功能需求 ===
    doc.add_heading("5  非功能需求", level=1)
    make_table(doc,
        ["需求编号", "类别", "需求描述", "指标"],
        [
            ["NFR-001", "性能", "策略下发到执行的响应时间", "\u2264 10 秒"],
            ["NFR-002", "性能", "管理后台页面加载时间", "\u2264 3 秒"],
            ["NFR-003", "性能", "设备列表自动刷新间隔", "5 秒"],
            ["NFR-004", "可靠性", "心跳超时判定设备离线", "超过 2 次心跳间隔(约30分钟)"],
            ["NFR-005", "可靠性", "APP 崩溃后自动恢复", "前台服务保活 + 开机自启"],
            ["NFR-006", "安全性", "ADB 命令安全过滤", "阻止破坏性命令"],
            ["NFR-007", "安全性", "设备认证", "每个 APP 实例持有唯一 token"],
            ["NFR-008", "可用性", "管理后台操作门槛", "管理员无需培训即可使用"],
            ["NFR-009", "兼容性", "电视端兼容性", "Android 6.0+，小米电视"],
            ["NFR-010", "兼容性", "浏览器兼容性", "Chrome / Edge / Firefox 最新版"],
            ["NFR-011", "可移植性", "项目可移植", "无硬编码路径和 IP，可跨机器部署"],
        ], highlight_col=[0], col_widths=[2.2, 2, 4.5, 4.5])

    # === 6 约束与假设 ===
    doc.add_heading("6  约束与假设", level=1)
    doc.add_heading("6.1  约束", level=2)
    make_table(doc,
        ["编号", "约束内容"],
        [
            ["C-001", "服务器与电视必须在同一局域网"],
            ["C-002", "电视必须开启 ADB 无线调试（端口 5555）"],
            ["C-003", "Scrcpy 仅在服务器本地可见，无法通过浏览器远程使用"],
            ["C-004", "Python 版本需 3.9-3.13（3.14+ 与 Pydantic 不兼容）"],
            ["C-005", "Android APP 最低 SDK 版本为 23（Android 6.0）"],
        ], highlight_col=[0], col_widths=[2, 11.5])
    doc.add_heading("6.2  假设", level=2)
    make_table(doc,
        ["编号", "假设内容"],
        [
            ["A-001", "所有电视均为小米品牌，运行 Android 系统"],
            ["A-002", "电视通过有线或 WiFi 连接到企业内网"],
            ["A-003", "管理员具有基本 IT 运维能力"],
        ], highlight_col=[0], col_widths=[2, 11.5])

    # === 7 术语表 ===
    doc.add_heading("7  术语表", level=1)
    make_table(doc,
        ["术语", "英文", "定义"],
        [
            ["策略", "Policy", "控制电视行为的规则，包括 APP 启动和 HDMI 切换两种模式"],
            ["心跳", "Heartbeat", "设备定时向服务器发送的状态更新请求"],
            ["ADB", "Android Debug Bridge", "Android 调试桥，用于与设备通信"],
            ["Scrcpy", "Screen Copy", "开源 Android 屏幕镜像工具，支持低延迟远程控制"],
            ["Kiosk", "Kiosk Mode", "锁定模式，限制设备只能运行指定应用"],
            ["OTA", "Over-The-Air", "空中升级，通过网络推送应用更新"],
            ["SN", "Serial Number", "设备序列号，本项目使用 Android ANDROID_ID"],
        ], highlight_col=[0], col_widths=[2, 3, 8.5])

    path = os.path.join(OUT_DIR, "01-需求规格说明书.docx")
    doc.save(path)
    print(f"已生成: {path}")


# ================================================================
# 文档2：PRD 产品需求文档
# ================================================================
def create_prd_doc():
    doc = Document()
    setup_styles(doc)
    setup_page(doc)

    add_cover(doc, "产品需求文档 (PRD)", "Meeting TV Launcher\n会议室电视集中管控系统", "TVL-PRD-001", "V1.0", "2026-04-23")
    add_revision_table(doc)

    add_toc(doc, [
        "1  产品概述  3",
        "2  功能模块详细设计  4",
        "3  用户操作流程  10",
        "4  界面原型  12",
        "5  API 接口清单  16",
        "6  数据库设计  18",
    ])

    # === 1 产品概述 ===
    doc.add_heading("1  产品概述", level=1)
    doc.add_heading("1.1  产品定位", level=2)
    add_body(doc, "Meeting TV Launcher 是一套企业会议室电视集中管控平台，实现对多台小米电视的统一策略管理、远程监控和远程控制。")
    doc.add_heading("1.2  目标用户", level=2)
    add_body(doc, "IT 运维管理员，负责会议室电视的日常管理和故障处理。")
    doc.add_heading("1.3  项目价值与意义", level=2)
    add_body(doc, "本产品致力于解决企业会议室电视管理中\u201c看不见、管不住、响应慢\u201d的三大核心问题，将分散的、依赖人力的管理方式升级为集中的、自动化的、远程可控的数字化管理平台。")
    doc.add_heading("1.3.1  业务价值", level=3)
    make_table(doc,
        ["价值点", "说明", "量化指标"],
        [
            ["运维效率", "从逐台手动配置到一键部署，新电视上线全流程自动化", "单台上线时间 30 分钟 \u2192 3 分钟"],
            ["故障响应", "远程投屏 + ADB 命令，无需到场即可诊断和修复", "故障响应 30 分钟 \u2192 5 分钟"],
            ["管理标准化", "统一策略模板，自动执行标准配置，杜绝人为差异", "配置不一致率降至 0"],
            ["实时监控", "Web 后台实时展示所有设备状态、网络和策略执行情况", "故障发现从被动等待变为主动感知"],
            ["安全可控", "Kiosk 锁定 + 危险命令过滤，防止非授权操作", "会议中误操作风险降至最低"],
            ["规模扩展", "零配置自动注册，单后台管理数十台设备", "管理规模 10 倍扩展无需增人"],
        ], col_widths=[2.5, 6, 5])
    doc.add_heading("1.3.2  用户价值", level=3)
    make_table(doc,
        ["用户角色", "价值", "典型场景"],
        [
            ["IT 管理员", "从逐台跑现场到坐在工位完成所有操作", "新电视上线不再需要带笔记本到会议室，后台输入 IP 即可完成部署"],
            ["IT 管理员", "实时掌握所有电视运行状态，问题主动发现", "通过在线状态监控，第一时间发现电视离线并远程排查"],
            ["IT 管理员", "策略一键切换，会议室用途变更快速响应", "会议室从投屏模式切换为 HDMI 模式，后台改策略即可，无需到现场"],
            ["会议使用者", "电视开机即可使用，无需等待人工配置", "到达会议室时电视已自动启动投屏软件，即开即用"],
        ], col_widths=[2.5, 5, 6])
    doc.add_heading("1.3.3  技术价值", level=3)
    make_table(doc,
        ["维度", "说明"],
        [
            ["架构简洁", "Python + SQLite 零依赖轻量架构，单进程部署，运维成本低"],
            ["协议标准化", "REST API + ADB 标准协议，无自定义协议，兼容性和可维护性强"],
            ["扩展灵活", "模块化设计，策略类型可扩展（当前支持 APP/HDMI，未来可扩展音量、亮度等）"],
            ["跨平台兼容", "后端跨平台（Linux/Windows/Mac），电视端兼容 Android 6.0+"],
            ["无外部依赖", "不依赖云服务和公网，纯内网部署，数据安全可控"],
        ], col_widths=[2.5, 10.5])

    # === 2 功能模块详细设计 ===
    doc.add_heading("2  功能模块详细设计", level=1)

    doc.add_heading("2.1  设备管理模块", level=2)
    doc.add_heading("2.1.1  设备列表", level=3)
    add_body(doc, '入口：首页\u201c设备管理（会议室视角）\u201d卡片。表格展示所有已注册设备，按 ID 倒序排列，5 秒自动刷新。')
    make_table(doc,
        ["功能项", "描述", "规则"],
        [
            ["列筛选", "每列标题可筛选，类似 Excel", "点击列标题右侧 \u25bc 图标，弹出复选框列表，支持搜索和全选"],
            ["列排序", "每列标题可排序", "点击列标题文字，循环：升序\u2191\u2192降序\u2193\u2192取消"],
            ["全局搜索", "右上角搜索框，搜索全部信息", "实时过滤，匹配设备名/SN/IP/MAC/型号/应用/策略等"],
            ["会议室标签", "每台设备显示所属会议室", "橙色标签，点击 \u270e 可编辑"],
        ], col_widths=[2.5, 4.5, 6.5])

    doc.add_heading("2.1.2  设备表格列定义", level=3)
    make_table(doc,
        ["列名", "显示内容", "筛选值来源", "排序依据"],
        [
            ["会议室/设备", "会议室名 + 设备名 + SN", "各会议室名、未分配", "会议室名 + 设备名"],
            ["型号/资源状态", "型号 + RAM/存储进度条", "各型号", "型号名"],
            ["网络详情", "有线 IP + WiFi IP + MAC", "各 IP", "IP 地址"],
            ["在线状态", "在线/离线 + SSID", "在线、离线", "在线优先"],
            ["执行策略", "策略下拉框 + 更新按钮", "各策略名、未分配", "策略名"],
            ["策略控制", "运行中/已暂停 + 按钮", "运行中、已暂停", "暂停状态"],
            ["操作", "远程控制/安装APK/应用管理/移除", "\u2014", "\u2014"],
        ], col_widths=[2.8, 4, 3.5, 3.2])

    doc.add_heading("2.1.3  设备操作", level=3)
    make_table(doc,
        ["操作", "入口", "逻辑"],
        [
            ["远程控制", '\u201c\U0001f4fa 远程控制\u201d按钮', "打开远程控制模态框：Scrcpy 控制、ADB 连接、遥控器按钮"],
            ["安装 APK", '\u201c安装APK\u201d按钮', "通过 ADB 安装最新版 APP 到电视"],
            ["应用管理", '\u201c\U0001f4f1 应用管理\u201d按钮', "列出已安装应用，支持卸载和上传安装新 APK"],
            ["移除设备", '\u201c移除\u201d按钮', "确认后删除设备记录 + ADB 清数据 + 停止 APP"],
            ["编辑会议室", '会议室标签旁\u201c\u270e\u201d按钮', "弹出编辑框修改会议室名称"],
        ], col_widths=[2.5, 3.5, 7.5])

    doc.add_heading("2.2  策略管理模块", level=2)
    doc.add_heading("2.2.1  策略库", level=3)
    add_body(doc, '入口：首页\u201c策略库\u201d卡片。表格展示所有策略。')
    make_table(doc,
        ["功能项", "描述", "规则"],
        [
            ["策略列表", "表格展示所有策略", "显示：名称、模式(APP/HDMI)、目标内容"],
            ["新增策略", "\u201c+ 新增策略\u201d按钮", "弹出模态框：名称、模式、目标APP或HDMI端口"],
            ["删除策略", "每行\u201c删除\u201d按钮", "确认后删除"],
        ], col_widths=[2.5, 4.5, 6.5])

    doc.add_heading("2.2.2  策略类型定义", level=3)
    make_table(doc,
        ["类型", "模式", "参数", "执行效果"],
        [
            ["APP 启动", "app", "target_app_package\uff08包名\uff09", "启动指定 APP（投屏、视频会议等）"],
            ["HDMI 切换", "hdmi", "target_hdmi_port\uff08端口号\uff09", "切换到指定 HDMI 输入源"],
        ], col_widths=[2.5, 2, 4.5, 4.5])

    doc.add_heading("2.2.3  策略绑定与控制", level=3)
    make_table(doc,
        ["操作", "入口", "逻辑"],
        [
            ["绑定策略", '设备行\u201c执行策略\u201d下拉框+\u201c下发\u201d按钮', "下拉选择策略（轮询刷新时保留未保存的选择），点击\u201c下发\u201d推送到设备并立即生效"],
            ["暂停策略", '设备行\u201c\u23f8 暂停\u201d按钮', '设备停止自动执行策略，标签变为\u201c已暂停\u201d'],
            ["恢复策略", '设备行\u201c\u25b6 继续\u201d按钮', '设备恢复执行策略，标签变为\u201c运行中\u201d'],
        ], col_widths=[2.5, 5, 5.5])

    doc.add_heading("2.3  远程控制模块", level=2)
    doc.add_heading("2.3.1  Scrcpy 高级控制", level=3)
    make_table(doc,
        ["操作", "描述"],
        [
            ["检查 Scrcpy", "检测服务器是否安装 Scrcpy"],
            ["启动 Scrcpy", "ADB 连接后启动 Scrcpy 窗口（服务器电脑上），低延迟实时控制"],
            ["停止 Scrcpy", "关闭 Scrcpy 进程"],
            ["获取命令", "显示 Scrcpy 启动命令，可手动在命令行执行"],
        ], col_widths=[3, 10])

    doc.add_heading("2.3.2  遥控器按钮", level=3)
    make_table(doc,
        ["按钮", "ADB 按键", "KeyCode", "功能"],
        [
            ["\u2b06 上", "KEYCODE_DPAD_UP", "19", "方向上"],
            ["\u2b07 下", "KEYCODE_DPAD_DOWN", "20", "方向下"],
            ["\u2b05 左", "KEYCODE_DPAD_LEFT", "21", "方向左"],
            ["\u27a1 右", "KEYCODE_DPAD_RIGHT", "22", "方向右"],
            ["确认", "KEYCODE_ENTER", "66", "确认选择"],
            ["返回", "KEYCODE_BACK", "4", "返回上一级"],
            ["主页", "KEYCODE_HOME", "3", "回到主页"],
            ["菜单", "KEYCODE_MENU", "82", "打开菜单"],
            ["音量+", "KEYCODE_VOLUME_UP", "24", "增大音量"],
            ["音量-", "KEYCODE_VOLUME_DOWN", "25", "减小音量"],
            ["电源", "KEYCODE_POWER", "26", "开关屏幕"],
        ], col_widths=[2.5, 4.5, 2.5, 4])

    doc.add_heading("2.4  ADB 命令控制台", level=2)
    make_table(doc,
        ["功能项", "描述"],
        [
            ["设备选择", "下拉框选择目标设备"],
            ["命令输入", "输入 ADB 命令（如 shell dumpsys battery）"],
            ["常用命令", "左侧列表双击填入命令（15 条常用命令）"],
            ["定时执行", "设置间隔秒数和重复次数（0=无限），周期性执行命令"],
            ["输出显示", "等宽字体终端风格，支持清空"],
            ["安全过滤", "禁止执行危险命令：rm -rf /、format、mkfs、dd if=、>/dev/、shutdown"],
        ], col_widths=[2.5, 10.5])

    doc.add_heading("2.4.1  常用命令列表", level=3)
    make_table(doc,
        ["命令", "说明"],
        [
            ["shell dumpsys battery", "电池信息"],
            ["shell dumpsys wifi", "WiFi 信息"],
            ["shell getprop ro.build.version.release", "系统版本"],
            ["shell getprop ro.product.model", "设备型号"],
            ["shell pm list packages -3", "第三方应用列表"],
            ["shell pm clear <包名>", "清除应用数据"],
            ["shell am force-stop <包名>", "强制停止应用"],
            ["shell am start -n <包名/Activity>", "启动应用"],
            ["shell input keyevent <keycode>", "模拟按键"],
            ["shell input tap <x> <y>", "模拟点击"],
            ["shell screencap -p /sdcard/screen.png", "截屏到设备"],
            ["shell settings get global <key>", "读取系统设置"],
            ["shell settings put global <key> <value>", "写入系统设置"],
            ["shell reboot", "重启设备"],
            ["logcat -d -t 50", "最近 50 条日志"],
        ], col_widths=[6, 7])

    doc.add_heading("2.5  快捷功能（顶部栏）", level=2)
    make_table(doc,
        ["功能", "入口", "描述"],
        [
            ["一键部署", 'IP 输入框+\u201c一键部署上线\u201d', "输入电视 IP，自动完成安装、配置、注册、绑定默认策略"],
            ["Scrcpy 快速连接", 'IP+端口输入框+\u201cScrcpy连接\u201d', "输入 IP 直接启动 Scrcpy"],
            ["ADB 控制台", '\u201c\U0001f4bb ADB 控制台\u201d按钮', "打开 ADB 命令控制台模态框"],
            ["操作日志", '\u201c\U0001f4cb 操作日志\u201d按钮', "查看操作历史记录"],
        ], col_widths=[3, 4.5, 5.5])

    doc.add_heading("2.6  操作日志", level=2)
    make_table(doc,
        ["功能项", "描述"],
        [
            ["日志列表", "按时间倒序显示所有操作记录"],
            ["操作类型", "策略绑定、远程控制、安装APK、卸载应用、删除策略、移除设备、修改会议室等"],
            ["清空记录", "一键清空所有日志"],
        ], col_widths=[2.5, 10.5])

    # === 3 用户操作流程 ===
    doc.add_heading("3  用户操作流程", level=1)

    doc.add_heading("3.1  新电视上线流程", level=2)
    add_flow_steps(doc, [
        "电视连网，开启 ADB 调试",
        "管理员在后台输入电视 IP \u2192 点击\u201c一键部署上线\u201d",
        "系统自动：ADB 连接 \u2192 安装 APP \u2192 配置服务器地址 \u2192 启动 APP",
        "APP 自动注册到后台，管理员绑定策略",
        "电视开始执行策略（如启动投屏软件）",
    ])

    doc.add_heading("3.2  切换会议室用途流程", level=2)
    add_flow_steps(doc, [
        "管理员在后台找到目标电视",
        "修改策略：从\u201c投屏软件\u201d切换为\u201cHDMI1\u201d",
        "策略通过 ADB 实时推送到电视，立即生效",
    ])

    doc.add_heading("3.3  会议中断暂停策略流程", level=2)
    add_flow_steps(doc, [
        "管理员点击\u201c暂停\u201d按钮",
        "电视停止自动执行策略，允许自由操作",
        "会议结束后点击\u201c继续\u201d，策略恢复执行",
    ])

    doc.add_heading("3.4  远程排查故障流程", level=2)
    add_flow_steps(doc, [
        "管理员点击\u201c远程控制\u201d \u2192 启动 Scrcpy",
        "实时查看电视画面，用鼠标键盘操作",
        "或打开 ADB 控制台执行诊断命令",
    ])

    # === 4 界面原型 ===
    doc.add_heading("4  界面原型", level=1)
    add_body(doc, "以下为各主要界面的原型图，展示布局和交互元素。")

    doc.add_heading("4.1  管理后台主页", level=2)
    add_prototype_box(doc,
"\u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510\n"
"\u2502 Meeting TV Launcher \u63a7\u5236\u53f0                                       \u2502\n"
"\u2502                                                                  \u2502\n"
"\u2502 [IP] [\u4e00\u952e\u90e8\u7f72\u4e0a\u7ebf]  [IP][5555][Scrcpy\u8fde\u63a5]  [ADB\u63a7\u5236\u53f0] [\u64cd\u4f5c\u65e5\u5fd7] \u2502\n"
"\u251c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2524\n"
"\u2502                                                                  \u2502\n"
"\u2502 \u8bbe\u5907\u7ba1\u7406 (\u4f1a\u8bae\u5ba4\u89c6\u89d2)                                  [\U0001f50d \u641c\u7d22] \u2502\n"
"\u2502 \u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u253c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u253c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u253c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u253c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u253c\u2500\u2500\u2500\u2500\u2500\u253c\u2500\u2500\u2500\u2510 \u2502\n"
"\u2502 \u2502\u4f1a\u8bae\u5ba4/\u8bbe\u5907\u2502\u578b\u53f7/\u8d44\u6e90 \u2502\u7f51\u7edc\u8be6\u60c5  \u2502\u5728\u7ebf\u72b6\u6001  \u2502\u6267\u884c\u7b56\u7565  \u2502\u7b56\u7565\u63a7\u2502\u64cd\u4f5c\u2502 \u2502\n"
"\u2502 \u2502  \u21c5 \u25bc   \u2502  \u21c5 \u25bc   \u2502  \u21c5 \u25bc   \u2502  \u21c5 \u25bc   \u2502  \u21c5 \u25bc   \u2502\u21c5 \u25bc  \u2502    \u2502 \u2502\n"
"\u2502 \u251c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u253c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u253c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u253c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u253c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u253c\u2500\u2500\u2500\u2500\u2500\u253c\u2500\u2500\u2500\u2524 \u2502\n"
"\u2502 \u2502[811]    \u2502MiTV4   \u2502\u6709\u7ebfIP   \u2502\u25cf \u5728\u7ebf   \u2502\u7b56\u7565\u4e0b\u62c9  \u2502\u25b6\u8fd0\u884c \u2502\U0001f4fa  \u2502 \u2502\n"
"\u2502 \u2502TV-DCF8  \u2502RAM:\u2588\u2588\u2590 \u2502WiFi IP \u2502SSID     \u2502[\u66f4\u65b0]    \u2502[\u23f8\u6682\u505c]\u2502\U0001f4e5\U0001f4f1\U0001f5d1\u2502 \u2502\n"
"\u2502 \u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2534\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2534\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2534\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2534\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2534\u2500\u2500\u2500\u2500\u2500\u2534\u2500\u2500\u2500\u2518 \u2502\n"
"\u2502                                                                  \u2502\n"
"\u2502 \u7b56\u7565\u5e93                                            [+ \u65b0\u589e\u7b56\u7565]  \u2502\n"
"\u2502 \u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u253c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u253c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u253c\u2500\u2500\u2500\u2500\u2500\u2510                 \u2502\n"
"\u2502 \u2502\u7b56\u7565\u540d\u79f0   \u2502\u6a21\u5f0f      \u2502\u76ee\u6807\u5185\u5bb9           \u2502\u7ba1\u7406   \u2502                 \u2502\n"
"\u2502 \u2502\u6295\u5c4f\u8f6f\u4ef6   \u2502APP      \u2502com.tcly.share..  \u2502[\u5220\u9664] \u2502                 \u2502\n"
"\u2502 \u2502HDMI1     \u2502HDMI     \u2502HDMI 1            \u2502[\u5220\u9664] \u2502                 \u2502\n"
"\u2502 \u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2534\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2534\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2534\u2500\u2500\u2500\u2500\u2500\u2518                 \u2502\n"
"\u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518"
    )

    doc.add_heading("4.2  列筛选弹出框", level=2)
    add_prototype_box(doc,
"  \u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510\n"
"  \u2502 \U0001f50d \u641c\u7d22...           \u2502\n"
"  \u251c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2524\n"
"  \u2502 \u2611 \uff08\u5168\u9009\uff09            \u2502\n"
"  \u2502 \u2611 811                \u2502\n"
"  \u2502 \u2611 106                \u2502\n"
"  \u2502 \u2611 \u6d4b\u8bd5\u7535\u89c6\u673a          \u2502\n"
"  \u2502 \u2610 \u672a\u5206\u914d\u4f1a\u8bae\u5ba4        \u2502\n"
"  \u251c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2524\n"
"  \u2502  [\u786e\u5b9a]    [\u53d6\u6d88]     \u2502\n"
"  \u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518"
    )

    doc.add_heading("4.3  远程控制模态框", level=2)
    add_prototype_box(doc,
"\u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510\n"
"\u2502 \u8fdc\u7a0b\u63a7\u5236: TV-9E52              [\u2715\u5173\u95ed]\u2502\n"
"\u2502                                      \u2502\n"
"\u2502 \u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510   \u2502\n"
"\u2502 \u2502 \U0001f527 Scrcpy \u9ad8\u7ea7\u63a7\u5236              \u2502   \u2502\n"
"\u2502 \u2502 [\U0001f50d\u68c0\u67e5] [\U0001f680\u542f\u52a8] [\U0001f4cb\u83b7\u53d6\u547d\u4ee4]   \u2502   \u2502\n"
"\u2502 \u2502 \u2705 Scrcpy \u8fd0\u884c\u4e2d (PID: 12345)   \u2502   \u2502\n"
"\u2502 \u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518   \u2502\n"
"\u2502                                      \u2502\n"
"\u2502 \u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510   \u2502\n"
"\u2502 \u2502 \U0001f517 ADB \u8fde\u63a5\u7ba1\u7406                  \u2502   \u2502\n"
"\u2502 \u2502 [\U0001f50c\u8fde\u63a5ADB] [\U0001f50b\u65ad\u5f00ADB]         \u2502   \u2502\n"
"\u2502 \u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518   \u2502\n"
"\u2502                                      \u2502\n"
"\u2502 \u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510   \u2502\n"
"\u2502 \u2502 \U0001f3ae \u9065\u63a7\u5668\u6309\u94ae                    \u2502   \u2502\n"
"\u2502 \u2502      [\u2b06\u4e0a]                      \u2502   \u2502\n"
"\u2502 \u2502 [\u2b05\u5de6] [\u23ce\u786e\u5b9a] [\u27a1\u53f3]             \u2502   \u2502\n"
"\u2502 \u2502      [\u2b07\u4e0b]                      \u2502   \u2502\n"
"\u2502 \u2502 [\u21a9\u8fd4\u56de] [\U0001f3e0\u4e3b\u9875] [\U0001f4cb\u83dc\u5355]        \u2502   \u2502\n"
"\u2502 \u2502 [\U0001f50a\u97f3\u91cf+] [\U0001f509\u97f3\u91cf-] [\u23fb\u7535\u6e90]     \u2502   \u2502\n"
"\u2502 \u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518   \u2502\n"
"\u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518"
    )

    doc.add_heading("4.4  ADB 控制台模态框", level=2)
    add_prototype_box(doc,
"\u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510\n"
"\u2502 \U0001f4bb ADB \u547d\u4ee4\u63a7\u5236\u53f0                                  [\u2715\u5173\u95ed]\u2502\n"
"\u2502                                                          \u2502\n"
"\u2502 [\u8bbe\u5907\u9009\u62e9\u25bc]  [\u547d\u4ee4\u8f93\u5165\u6846____________] [\u25b6\u6267\u884c] [\u23f9\u505c\u6b62]    \u2502\n"
"\u2502 \u23f0 \u5b9a\u65f6: [5]\u79d2/\u6b21  [0]\u6b21(0=\u65e0\u9650)  [\u5f00\u59cb\u5b9a\u65f6]             \u2502\n"
"\u2502                                                          \u2502\n"
"\u2502 \u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510  \u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510  \u2502\n"
"\u2502 \u2502\u5e38\u7528\u547d\u4ee4(\u53cc\u51fb)  \u2502  \u2502\u8f93\u51fa\u7ed3\u679c                    [\u6e05\u7a7a]  \u2502  \u2502\n"
"\u2502 \u2502              \u2502  \u2502                                   \u2502  \u2502\n"
"\u2502 \u2502 \u7535\u6c60\u4fe1\u606f      \u2502  \u2502 > adb shell dumpsys battery      \u2502  \u2502\n"
"\u2502 \u2502 WiFi \u4fe1\u606f     \u2502  \u2502 Current Battery Manager state:   \u2502  \u2502\n"
"\u2502 \u2502 \u7cfb\u7edf\u7248\u672c      \u2502  \u2502   AC powered: false              \u2502  \u2502\n"
"\u2502 \u2502 \u8bbe\u5907\u578b\u53f7      \u2502  \u2502   USB powered: true              \u2502  \u2502\n"
"\u2502 \u2502 \u7b2c\u4e09\u65b9\u5e94\u7528    \u2502  \u2502   level: 100                     \u2502  \u2502\n"
"\u2502 \u2502 ...          \u2502  \u2502                                   \u2502  \u2502\n"
"\u2502 \u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518  \u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518  \u2502\n"
"\u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518"
    )

    doc.add_heading("4.5  电视 APP 界面", level=2)
    add_prototype_box(doc,
"\u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510\n"
"\u2502 10:35                                      \u2502\n"
"\u2502                                            \u2502\n"
"\u2502                                  [\u2699]       \u2502\n"
"\u2502                                            \u2502\n"
"\u2502            \u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510            \u2502\n"
"\u2502            \u2502   \u542f\u52a8\u6295\u5c4f       \u2502            \u2502\n"
"\u2502            \u2502 \u70b9\u51fb\u7acb\u5373\u6267\u884c\u7b56\u7565  \u2502            \u2502\n"
"\u2502            \u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518            \u2502\n"
"\u2502                                            \u2502\n"
"\u2502\u6295\u5c4f\u8f6f\u4ef6              [APP\u6a21\u5f0f] [\u25b6 \u8fd0\u884c\u4e2d]  \u2502\n"
"\u2502                         \u25cf \u5df2\u8fde\u63a5           \u2502\n"
"\u2502                         v0.2.0             \u2502\n"
"\u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518"
    )

    # === 5 API 接口清单 ===
    doc.add_heading("5  API 接口清单", level=1)
    make_table(doc,
        ["方法", "路径", "说明"],
        [
            ["GET", "/health", "健康检查"],
            ["GET", "/", "管理后台页面"],
            ["POST", "/api/v1/devices/register", "设备注册"],
            ["POST", "/api/v1/devices/heartbeat", "设备心跳"],
            ["GET", "/api/v1/devices", "设备列表"],
            ["GET", "/api/device/status?ip=", "按 IP 查询设备状态"],
            ["DELETE", "/api/v1/devices/{id}", "移除设备"],
            ["POST", "/api/v1/devices/{id}/bind-policy/{pid}", "绑定策略"],
            ["POST", "/api/v1/devices/{id}/room", "修改会议室"],
            ["POST", "/api/v1/devices/{id}/pause-policy", "暂停策略"],
            ["POST", "/api/v1/devices/{id}/resume-policy", "恢复策略"],
            ["POST", "/api/v1/deploy-tv?ip=", "一键部署"],
            ["GET", "/api/v1/devices/{id}/screenshot", "截图"],
            ["POST", "/api/v1/devices/{id}/input", "发送按键/点击"],
            ["POST", "/api/v1/devices/{id}/adb/shell", "执行 ADB 命令"],
            ["POST", "/api/v1/devices/{id}/adb/connect", "ADB 连接"],
            ["POST", "/api/v1/devices/{id}/adb/disconnect", "ADB 断开"],
            ["POST", "/api/v1/devices/{id}/adb-install", "ADB 安装 APK"],
            ["POST", "/api/v1/devices/{id}/uninstall", "卸载应用"],
            ["POST", "/api/v1/devices/{id}/install-uploaded", "安装上传的 APK"],
            ["POST", "/api/v1/devices/{id}/mouse", "鼠标控制"],
            ["GET", "/api/v1/scrcpy/check", "检查 Scrcpy"],
            ["POST", "/api/v1/devices/{id}/scrcpy/start", "启动 Scrcpy"],
            ["POST", "/api/v1/devices/{id}/scrcpy/stop", "停止 Scrcpy"],
            ["GET", "/api/v1/devices/{id}/scrcpy/status", "Scrcpy 状态"],
            ["GET", "/api/v1/devices/{id}/scrcpy/command", "获取 Scrcpy 命令"],
            ["POST", "/api/v1/scrcpy/connect-by-ip?ip=", "IP 直连 Scrcpy"],
            ["POST", "/api/v1/upload-apk", "上传 APK"],
            ["GET", "/api/v1/ota/check", "OTA 检查更新"],
            ["GET", "/api/v1/policies", "策略列表"],
            ["POST", "/api/v1/policies", "创建策略"],
            ["DELETE", "/api/v1/policies/{id}", "删除策略"],
            ["GET", "/api/v1/logs", "操作日志"],
            ["DELETE", "/api/v1/logs", "清空日志"],
        ], col_widths=[2, 6.5, 5])

    # === 6 数据库设计 ===
    doc.add_heading("6  数据库设计", level=1)

    doc.add_heading("6.1  devices 表", level=2)
    make_table(doc,
        ["字段", "类型", "约束", "说明"],
        [
            ["id", "Integer", "PK, 自增", "设备 ID"],
            ["device_sn", "String(128)", "UNIQUE, NOT NULL", "设备序列号(ANDROID_ID)"],
            ["device_name", "String(128)", "NOT NULL", "设备名称(如 TV-9E52)"],
            ["room_name", "String(128)", "", "会议室名称"],
            ["model_name", "String(128)", "", "电视型号"],
            ["wifi_ip", "String(64)", "", "WiFi IP"],
            ["eth_ip", "String(64)", "", "有线 IP"],
            ["wifi_mac", "String(32)", "", "WiFi MAC"],
            ["eth_mac", "String(32)", "", "有线 MAC"],
            ["network_ssid", "String(128)", "", "WiFi 名称"],
            ["network_type", "String(16)", "", "当前活动网络类型(wifi/ethernet/none)"],
            ["installed_apps", "Text", "", "已安装应用(JSON 数组)"],
            ["ram_usage", "String(64)", "", "内存使用"],
            ["storage_usage", "String(64)", "", "存储使用"],
            ["online", "Boolean", "默认 False", "是否在线"],
            ["token", "String(128)", "UNIQUE, NOT NULL", "认证令牌"],
            ["policy_id", "Integer", "FK \u2192 policies.id", "绑定策略 ID"],
            ["policy_paused", "Boolean", "默认 False", "策略是否暂停"],
            ["updated_at", "DateTime", "自动更新", "更新时间"],
        ], highlight_col=[0], col_widths=[3, 2.8, 3.2, 4.5])

    doc.add_heading("6.2  policies 表", level=2)
    make_table(doc,
        ["字段", "类型", "约束", "说明"],
        [
            ["id", "Integer", "PK, 自增", "策略 ID"],
            ["name", "String(128)", "UNIQUE, NOT NULL", "策略名称"],
            ["mode", "String(16)", "NOT NULL", "模式(app/hdmi)"],
            ["target_app_package", "String(256)", "", "目标 APP 包名"],
            ["target_hdmi_port", "Integer", "", "HDMI 端口号"],
            ["is_active", "Boolean", "默认 True", "是否启用"],
            ["created_at", "DateTime", "自动", "创建时间"],
        ], highlight_col=[0], col_widths=[3, 2.8, 3.2, 4.5])

    doc.add_heading("6.3  device_heartbeats 表", level=2)
    make_table(doc,
        ["字段", "类型", "约束", "说明"],
        [
            ["id", "Integer", "PK, 自增", "记录 ID"],
            ["device_id", "Integer", "FK \u2192 devices.id", "设备 ID"],
            ["status", "String(32)", "默认 ok", "心跳状态"],
            ["message", "Text", "", "状态消息"],
            ["created_at", "DateTime", "自动", "时间戳"],
        ], highlight_col=[0], col_widths=[3, 2.8, 3.2, 4.5])

    doc.add_heading("6.4  operation_logs 表", level=2)
    make_table(doc,
        ["字段", "类型", "约束", "说明"],
        [
            ["id", "Integer", "PK, 自增", "记录 ID"],
            ["device_id", "Integer", "", "设备 ID"],
            ["device_name", "String(128)", "", "设备名称"],
            ["action", "String(64)", "NOT NULL", "操作类型"],
            ["detail", "Text", "", "操作详情"],
            ["operator", "String(128)", "默认 admin", "操作人"],
            ["created_at", "DateTime", "自动", "时间戳"],
        ], highlight_col=[0], col_widths=[3, 2.8, 3.2, 4.5])

    path = os.path.join(OUT_DIR, "02-产品需求文档PRD.docx")
    doc.save(path)
    print(f"已生成: {path}")


# ================================================================
# 文档3：原型图与技术设计文档
# ================================================================
def create_design_doc():
    doc = Document()
    setup_styles(doc)
    setup_page(doc)

    add_cover(doc, "原型图与技术设计文档", "Meeting TV Launcher\n会议室电视集中管控系统", "TVL-DES-001", "V1.0", "2026-04-23")
    add_revision_table(doc)

    add_toc(doc, [
        "1  系统架构  3",
        "2  通信协议  4",
        "3  数据流设计  5",
        "4  界面原型图  7",
        "5  Android APP 架构  11",
        "6  运行环境  13",
    ])

    # === 1 系统架构 ===
    doc.add_heading("1  系统架构", level=1)
    doc.add_heading("1.1  整体架构图", level=2)
    add_prototype_box(doc,
"\u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510     \u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510     \u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510\n"
"\u2502  \u7ba1\u7406\u5458\u6d4f\u89c8\u5668  \u2502\u2500\u2500\u25b6\u2502    FastAPI \u540e\u7aef\u670d\u52a1      \u2502\u2540\u2500\u2500\u2500\u2500\u2502  \u5c0f\u7c73\u7535\u89c6 APP  \u2502\n"
"\u2502  (Web \u63a7\u5236\u53f0)  \u2502HTTP \u2502    (0.0.0.0:8000)      \u2502\u5fc3\u8df3 \u2502  (TV Launcher) \u2502\n"
"\u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518     \u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518     \u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518\n"
"                           \u2502           \u25b2                   \u2502\n"
"                           \u2502ADB        \u2502ADB \u63a8\u9001            \u2502\u81ea\u52a8\u6ce8\u518c\n"
"                           \u25bc           \u2502                   \u25bc\n"
"                     \u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510      \u2502             \u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510\n"
"                     \u2502 Scrcpy   \u2502      \u2502             \u2502   SQLite DB   \u2502\n"
"                     \u2502 (\u672c\u5730\u8fdb\u7a0b) \u2502      \u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2502  (\u8bbe\u5907/\u7b56\u7565)   \u2502\n"
"                     \u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518                    \u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518"
    )

    doc.add_heading("1.2  子系统说明", level=2)
    make_table(doc,
        ["子系统", "技术栈", "部署位置", "职责"],
        [
            ["后端服务", "Python 3.9+ / FastAPI / SQLAlchemy / SQLite", "服务器(0.0.0.0:8000)", "REST API、Web 页面、ADB 控制、Scrcpy 管理"],
            ["Web 前端", "HTML / CSS / JavaScript", "后端内嵌模板", "管理后台界面、设备管理、远程控制"],
            ["Android APP", "Kotlin / Android SDK 34 / OkHttp / WorkManager", "小米电视", "策略执行、心跳上报、自注册、保活"],
        ], col_widths=[2.5, 4.5, 3.5, 3])

    # === 2 通信协议 ===
    doc.add_heading("2  通信协议", level=1)
    make_table(doc,
        ["通信方向", "协议", "说明"],
        [
            ["浏览器 \u2192 后端", "HTTP REST", "管理操作 API"],
            ["APP \u2192 后端", "HTTP REST", "注册、心跳 API"],
            ["后端 \u2192 电视", "ADB 无线", "策略推送、远程控制、APK 安装"],
            ["后端 \u2192 本地", "进程启动", "Scrcpy 窗口"],
        ], col_widths=[3.5, 3, 7])

    # === 3 数据流设计 ===
    doc.add_heading("3  数据流设计", level=1)

    doc.add_heading("3.1  设备注册", level=2)
    add_prototype_box(doc,
"APP --POST /register--> \u540e\u7aef --> SQLite(devices \u8868)\n"
"\u540e\u7aef\u8fd4\u56de token + device_id\n"
"APP \u4fdd\u5b58 token \u7528\u4e8e\u540e\u7eed\u5fc3\u8df3\u8ba4\u8bc1"
    )

    doc.add_heading("3.2  心跳上报", level=2)
    add_prototype_box(doc,
"APP --POST /heartbeat (Header: X-Device-Token)--> \u540e\u7aef\n"
"\u540e\u7aef\u66f4\u65b0\u8bbe\u5907\u4fe1\u606f(\u7f51\u7edc/\u8d44\u6e90/\u5e94\u7528) + \u8fd4\u56de\u5f53\u524d\u7b56\u7565\u548c policy_paused \u72b6\u6001\n"
"APP \u5bf9\u6bd4\u7b56\u7565\u53d8\u5316 \u2192 \u53d1\u9001 POLICY_UPDATED \u5e7f\u64ad \u2192 \u91cd\u65b0\u6267\u884c\u7b56\u7565"
    )
    add_body(doc, '后端每 60 秒执行离线检测：将 online=True 且 updated_at 超过 30 分钟未更新的设备自动标记为 offline。设备关机或断网后，最长 30 分钟内后台状态自动更新为离线。')

    doc.add_heading("3.3  策略下发", level=2)
    add_prototype_box(doc,
"\u7ba1\u7406\u5458 --> \u540e\u7aef --ADB broadcast POLICY_UPDATED--> APP\n"
"APP \u6536\u5230\u5e7f\u64ad \u2192 \u540c\u6b65\u5fc3\u8df3 \u2192 \u83b7\u53d6\u6700\u65b0\u7b56\u7565 \u2192 \u6267\u884c\u7b56\u7565"
    )

    doc.add_heading("3.4  一键部署流程", level=2)
    add_prototype_box(doc,
"\u7ba1\u7406\u5458              \u540e\u7aef\u670d\u52a1\u5668                    \u7535\u89c6\n"
"  \u2502                    \u2502                         \u2502\n"
"  \u2502 \u8f93\u5165IP,\u70b9\u51fb\u90e8\u7f72     \u2502                         \u2502\n"
"  \u2502\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u25b6\u2502  ADB connect IP:5555    \u2502\n"
"  \u2502                    \u2502\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u25b6\u2502\n"
"  \u2502                    \u2502  pm clear (\u6e05\u65e7\u6570\u636e)     \u2502\n"
"  \u2502                    \u2502\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u25b6\u2502\n"
"  \u2502                    \u2502  install APK            \u2502\n"
"  \u2502                    \u2502\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u25b6\u2502\n"
"  \u2502                    \u2502  \u914d\u7f6e\u670d\u52a1\u5668URL           \u2502\n"
"  \u2502                    \u2502  (settings put global)  \u2502\n"
"  \u2502                    \u2502\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u25b6\u2502\n"
"  \u2502                    \u2502  am start \u542f\u52a8APP       \u2502\n"
"  \u2502                    \u2502\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u25b6\u2502\n"
"  \u2502                    \u2502                         \u2502\u2500\u2500\u2590 APP\u81ea\u6ce8\u518c\n"
"  \u2502                    \u2502\u2540\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2502\u2540\u2500\u2500\u2518 POST /register\n"
"  \u2502                    \u2502  \u7b49\u5f8510\u79d2,\u67e5\u8be2\u6570\u636e\u5e93      \u2502\n"
"  \u2502                    \u2502  \u7ed1\u5b9a\u9ed8\u8ba4\u7b56\u7565             \u2502\n"
"  \u2502 \u2705 \u90e8\u7f72\u6210\u529f        \u2502                         \u2502\n"
"  \u2502\u2540\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518\u2502                         \u2502"
    )

    # === 4 界面原型图 ===
    doc.add_heading("4  界面原型图", level=1)

    doc.add_heading("4.1  管理后台主页", level=2)
    add_body(doc, "页面分为两大卡片区域：")
    add_body(doc, "上方：设备管理卡片 \u2014 标题栏含全局搜索框，表格含 7 列（会议室/设备、型号/资源、网络、在线状态、策略、控制、操作），每列支持筛选和排序，每行设备含操作按钮组。")
    add_body(doc, "下方：策略库卡片 \u2014 策略列表 + 新增策略按钮。")
    add_prototype_box(doc,
"\u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510\n"
"\u2502 Meeting TV Launcher \u63a7\u5236\u53f0                                       \u2502\n"
"\u2502 [IP] [\u4e00\u952e\u90e8\u7f72\u4e0a\u7ebf]  [IP][5555][Scrcpy\u8fde\u63a5]  [ADB\u63a7\u5236\u53f0] [\u64cd\u4f5c\u65e5\u5fd7] \u2502\n"
"\u251c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2524\n"
"\u2502 \u8bbe\u5907\u7ba1\u7406 (\u4f1a\u8bae\u5ba4\u89c6\u89d2)                                  [\U0001f50d \u641c\u7d22] \u2502\n"
"\u2502 \u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u253c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u253c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u253c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u253c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u253c\u2500\u2500\u2500\u2500\u2500\u253c\u2500\u2500\u2500\u2510 \u2502\n"
"\u2502 \u2502\u4f1a\u8bae\u5ba4/\u8bbe\u5907\u2502\u578b\u53f7/\u8d44\u6e90 \u2502\u7f51\u7edc\u8be6\u60c5  \u2502\u5728\u7ebf\u72b6\u6001  \u2502\u6267\u884c\u7b56\u7565  \u2502\u7b56\u7565\u63a7\u2502\u64cd\u4f5c\u2502 \u2502\n"
"\u2502 \u2502  \u21c5 \u25bc   \u2502  \u21c5 \u25bc   \u2502  \u21c5 \u25bc   \u2502  \u21c5 \u25bc   \u2502  \u21c5 \u25bc   \u2502\u21c5 \u25bc  \u2502    \u2502 \u2502\n"
"\u2502 \u251c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u253c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u253c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u253c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u253c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u253c\u2500\u2500\u2500\u2500\u2500\u253c\u2500\u2500\u2500\u2524 \u2502\n"
"\u2502 \u2502[811]    \u2502MiTV4   \u2502\u6709\u7ebfIP   \u2502\u25cf \u5728\u7ebf   \u2502\u7b56\u7565\u4e0b\u62c9  \u2502\u25b6\u8fd0\u884c \u2502\U0001f4fa  \u2502 \u2502\n"
"\u2502 \u2502TV-DCF8  \u2502RAM:\u2588\u2588\u2590 \u2502WiFi IP \u2502SSID     \u2502[\u66f4\u65b0]    \u2502[\u23f8\u6682\u505c]\u2502\U0001f4e5\U0001f4f1\U0001f5d1\u2502 \u2502\n"
"\u2502 \u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2534\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2534\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2534\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2534\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2534\u2500\u2500\u2500\u2500\u2500\u2534\u2500\u2500\u2500\u2518 \u2502\n"
"\u2502                                                                  \u2502\n"
"\u2502 \u7b56\u7565\u5e93                                            [+ \u65b0\u589e\u7b56\u7565]  \u2502\n"
"\u2502 \u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u253c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u253c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u253c\u2500\u2500\u2500\u2500\u2500\u2510                 \u2502\n"
"\u2502 \u2502\u7b56\u7565\u540d\u79f0   \u2502\u6a21\u5f0f      \u2502\u76ee\u6807\u5185\u5bb9           \u2502\u7ba1\u7406   \u2502                 \u2502\n"
"\u2502 \u2502\u6295\u5c4f\u8f6f\u4ef6   \u2502APP      \u2502com.tcly.share..  \u2502[\u5220\u9664] \u2502                 \u2502\n"
"\u2502 \u2502HDMI1     \u2502HDMI     \u2502HDMI 1            \u2502[\u5220\u9664] \u2502                 \u2502\n"
"\u2502 \u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2534\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2534\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2534\u2500\u2500\u2500\u2500\u2500\u2518                 \u2502\n"
"\u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518"
    )

    doc.add_heading("4.2  远程控制模态框", level=2)
    add_prototype_box(doc,
"\u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510\n"
"\u2502 \u8fdc\u7a0b\u63a7\u5236: TV-9E52              [\u2715\u5173\u95ed]\u2502\n"
"\u2502                                      \u2502\n"
"\u2502 \u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510   \u2502\n"
"\u2502 \u2502 \U0001f527 Scrcpy \u9ad8\u7ea7\u63a7\u5236              \u2502   \u2502\n"
"\u2502 \u2502 [\U0001f50d\u68c0\u67e5] [\U0001f680\u542f\u52a8] [\U0001f4cb\u83b7\u53d6\u547d\u4ee4]   \u2502   \u2502\n"
"\u2502 \u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518   \u2502\n"
"\u2502                                      \u2502\n"
"\u2502 \u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510   \u2502\n"
"\u2502 \u2502 \U0001f517 ADB \u8fde\u63a5\u7ba1\u7406                  \u2502   \u2502\n"
"\u2502 \u2502 [\U0001f50c\u8fde\u63a5] [\U0001f50b\u65ad\u5f00]               \u2502   \u2502\n"
"\u2502 \u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518   \u2502\n"
"\u2502                                      \u2502\n"
"\u2502 \u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510   \u2502\n"
"\u2502 \u2502 \U0001f3ae \u9065\u63a7\u5668\u6309\u94ae                    \u2502   \u2502\n"
"\u2502 \u2502      [\u2b06\u4e0a]                      \u2502   \u2502\n"
"\u2502 \u2502 [\u2b05\u5de6] [\u23ce\u786e\u5b9a] [\u27a1\u53f3]             \u2502   \u2502\n"
"\u2502 \u2502      [\u2b07\u4e0b]                      \u2502   \u2502\n"
"\u2502 \u2502 [\u21a9\u8fd4\u56de] [\U0001f3e0\u4e3b\u9875] [\U0001f4cb\u83dc\u5355]        \u2502   \u2502\n"
"\u2502 \u2502 [\U0001f50a+] [\U0001f509-] [\u23fb\u7535\u6e90]            \u2502   \u2502\n"
"\u2502 \u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518   \u2502\n"
"\u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518"
    )

    doc.add_heading("4.3  ADB 命令控制台", level=2)
    add_prototype_box(doc,
"\u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510\n"
"\u2502 \U0001f4bb ADB \u547d\u4ee4\u63a7\u5236\u53f0                                  [\u2715\u5173\u95ed]\u2502\n"
"\u2502 [\u8bbe\u5907\u9009\u62e9\u25bc]  [\u547d\u4ee4\u8f93\u5165\u6846____________] [\u25b6\u6267\u884c]             \u2502\n"
"\u2502 \u23f0 \u5b9a\u65f6: [5]\u79d2/\u6b21  [0]\u6b21(0=\u65e0\u9650)  [\u5f00\u59cb\u5b9a\u65f6]             \u2502\n"
"\u2502 \u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510  \u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510  \u2502\n"
"\u2502 \u2502\u5e38\u7528\u547d\u4ee4(\u53cc\u51fb)  \u2502  \u2502\u8f93\u51fa\u7ed3\u679c                    [\u6e05\u7a7a]  \u2502  \u2502\n"
"\u2502 \u2502 \u7535\u6c60\u4fe1\u606f      \u2502  \u2502 > adb shell dumpsys battery      \u2502  \u2502\n"
"\u2502 \u2502 WiFi \u4fe1\u606f     \u2502  \u2502   level: 100                     \u2502  \u2502\n"
"\u2502 \u2502 \u7cfb\u7edf\u7248\u672c      \u2502  \u2502                                   \u2502  \u2502\n"
"\u2502 \u2502 \u8bbe\u5907\u578b\u53f7      \u2502  \u2502                                   \u2502  \u2502\n"
"\u2502 \u2502 \u7b2c\u4e09\u65b9\u5e94\u7528    \u2502  \u2502                                   \u2502  \u2502\n"
"\u2502 \u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518  \u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518  \u2502\n"
"\u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518"
    )

    doc.add_heading("4.4  应用管理", level=2)
    add_prototype_box(doc,
"\u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510\n"
"\u2502 \u5e94\u7528\u7ba1\u7406 - TV-9E52                   \u2502\n"
"\u2502                                      \u2502\n"
"\u2502 \u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510   \u2502\n"
"\u2502 \u2502 com.tcly.sharescreen   [\u5378\u8f7d]  \u2502   \u2502\n"
"\u2502 \u2502 com.company.tvlauncher [\u5378\u8f7d]  \u2502   \u2502\n"
"\u2502 \u2502 com.xiaomi.mitv.tvplayer[\u5378\u8f7d] \u2502   \u2502\n"
"\u2502 \u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518   \u2502\n"
"\u2502                                      \u2502\n"
"\u2502 \u5b89\u88c5\u65b0\u5e94\u7528                           \u2502\n"
"\u2502 [\u9009\u62e9APK\u6587\u4ef6___________] [\u4e0a\u4f20\u5e76\u5b89\u88c5] \u2502\n"
"\u2502                                      \u2502\n"
"\u2502                            [\u5173\u95ed]     \u2502\n"
"\u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518"
    )

    doc.add_heading("4.5  电视 APP 界面", level=2)
    add_prototype_box(doc,
"\u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510\n"
"\u2502 10:35                                      \u2502\n"
"\u2502                                            \u2502\n"
"\u2502                                  [\u2699]       \u2502\n"
"\u2502                                            \u2502\n"
"\u2502            \u250c\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2510            \u2502\n"
"\u2502            \u2502   \u542f\u52a8\u6295\u5c4f       \u2502            \u2502\n"
"\u2502            \u2502 \u70b9\u51fb\u7acb\u5373\u6267\u884c\u7b56\u7565  \u2502            \u2502\n"
"\u2502            \u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518            \u2502\n"
"\u2502                                            \u2502\n"
"\u2502\u6295\u5c4f\u8f6f\u4ef6              [APP\u6a21\u5f0f] [\u25b6 \u8fd0\u884c\u4e2d]  \u2502\n"
"\u2502                         \u25cf \u5df2\u8fde\u63a5           \u2502\n"
"\u2502                         v0.2.0             \u2502\n"
"\u2514\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2518"
    )

    # === 5 Android APP 架构 ===
    doc.add_heading("5  Android APP 架构", level=1)

    doc.add_heading("5.1  源文件清单", level=2)
    make_table(doc,
        ["文件名", "行数", "职责"],
        [
            ["MainActivity.kt", "625", "主界面、策略执行入口、广播接收、保活检测"],
            ["RemoteApi.kt", "193", "HTTP 通信：注册、心跳、OTA"],
            ["PolicyStore.kt", "137", "SharedPreferences 存储策略和配置"],
            ["LauncherExecutor.kt", "392", "策略执行引擎：APP 启动、HDMI 切换、进程检测"],
            ["KeepAliveForegroundService.kt", "182", "前台保活服务：5 秒检测目标 APP"],
            ["NetworkInfoProvider.kt", "85", "网络信息采集：IP/MAC/SSID/活动网络类型"],
            ["SyncWorker.kt", "51", "WorkManager 后台定时同步（15 分钟）"],
            ["BootReceiver.kt", "70", "开机自启：立即启动Activity+延迟1.5秒执行策略"],
            ["HomeKeyService.kt", "76", "无障碍服务：Kiosk 模式拦截 HOME/BACK"],
            ["SettingsActivity.kt", "157", "设置界面：服务器地址、策略、企业 WiFi"],
        ], col_widths=[6, 1.5, 5.5])

    doc.add_heading("5.2  关键机制", level=2)
    make_table(doc,
        ["机制", "实现方式", "说明"],
        [
            ["策略执行", "BroadcastReceiver + 心跳对比", "收到 POLICY_UPDATED 广播后同步最新策略并立即执行"],
            ["保活检测", "前台服务 + 5 秒定时", "检测目标 APP 是否运行，不运行则重新启动"],
            ["启动冷却", "SharedPreferences 记录时间", "启动 APP 后 15 秒内判定为运行中，避免重启循环"],
            ["前后台判断", "onResume/onPause 写状态", "Launcher 在前台 \u2192 目标 APP 没在运行；Launcher 在后台 \u2192 目标 APP 在运行"],
            ["Kiosk 模式", "AccessibilityService", "拦截 HOME/BACK，三击 HOME 临时退出 10 分钟"],
            ["HDMI 切换", "TvView API + 多芯片适配", "通过TvInputManager/TvView直接调谐HDMI输入，适配Amlogic/MediaTek/MStar/Realtek"],
            ["冷启动优化", "轻量Splash主题 + 立即启动", "Manifest使用android:Theme.NoTitleBar.Fullscreen避免MaterialComponents白屏，onCreate中切回完整主题"],
        ], col_widths=[2.5, 4.5, 6.5])

    # === 6 运行环境 ===
    doc.add_heading("6  运行环境", level=1)
    make_table(doc,
        ["组件", "要求"],
        [
            ["后端 Python", "3.9 - 3.13\uff083.14+ \u4e0e Pydantic \u4e0d\u517c\u5bb9\uff09"],
            ["后端框架", "FastAPI 0.115 + Uvicorn 0.30 + SQLAlchemy 2.0"],
            ["数据库", "SQLite\uff08\u6587\u4ef6 tv_launcher.db\uff09"],
            ["电视系统", "Android 6.0+\uff0c\u5c0f\u7c73\u54c1\u724c\uff0cADB \u65e0\u7ebf\u8c03\u8bd5\u5f00\u542f"],
            ["电视 APP", "minSdk 23\uff0ctargetSdk 34\uff0cKotlin 1.9"],
            ["网络", "\u670d\u52a1\u5668\u4e0e\u7535\u89c6\u540c\u4e00\u5c40\u57df\u7f51\uff0cADB \u7aef\u53e3 5555"],
            ["Scrcpy", "\u53ef\u9009\uff0c\u5b89\u88c5\u5728\u670d\u52a1\u5668\u7535\u8111\u4e0a\uff0cv2.4"],
            ["浏览器", "Chrome / Edge / Firefox 最新版"],
        ], col_widths=[3, 10])

    path = os.path.join(OUT_DIR, "03-原型图与技术设计文档.docx")
    doc.save(path)
    print(f"已生成: {path}")


# ========== 执行生成 ==========
if __name__ == "__main__":
    create_requirement_doc()
    create_prd_doc()
    create_design_doc()
    print("\n全部文档生成完成！")
    print(f"输出目录: {OUT_DIR}")
