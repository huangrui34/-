#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成电视部署上线操作文档 - 供其他IT人员使用"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

OUTPUT_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "电视部署上线操作文档.docx")

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
style.paragraph_format.line_spacing = 1.35
style.paragraph_format.space_after = Pt(6)

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_styled_table(headers, rows, header_color="1F4E79"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(10)
        set_cell_shading(cell, header_color)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    return table

# ==================== 文档标题 ====================
title = doc.add_heading('电视部署上线操作文档', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('同程旅行总部会议室电视管理系统')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

doc.add_paragraph()

# ==================== 一、前置条件 ====================
doc.add_heading('一、前置条件（必须满足）', level=1)

p = doc.add_paragraph()
run = p.add_run('以下条件必须全部满足，否则无法完成部署：')
run.bold = True

add_styled_table(
    ['序号', '前置条件', '检查方法', '不满足的后果'],
    [
        ['1', '电视已开启ADB调试', '设置 → 关于本机 → 连点「版本号」7次 → 返回开发者选项 → 开启USB调试和无线调试', '无法远程连接和控制电视'],
        ['2', '电视与管理电脑在同一网络', '在电脑上 ping 电视IP，能通即可', 'ADB无法连接'],
        ['3', '电视已授权ADB调试', '首次ADB连接时电视弹出授权提示，需点击「允许」', 'ADB连接被拒绝'],
        ['4', '电视已连接网络（WiFi或有线）', '电视设置 → 网络 → 查看连接状态', '无法与管理后台通信'],
    ]
)

doc.add_paragraph()

# ==================== 二、部署步骤 ====================
doc.add_heading('二、部署步骤', level=1)

doc.add_heading('步骤1：确认电视IP地址', level=2)
doc.add_paragraph('在电视上进入：设置 → 关于本机 → 查看IP地址（有线网络优先）', style='List Bullet')
doc.add_paragraph('记录IP地址，后续步骤会用到', style='List Bullet')

doc.add_heading('步骤2：连接ADB', level=2)
doc.add_paragraph('在管理后台顶部找到「远程连接」输入框，输入电视IP，点击连接。')
p = doc.add_paragraph()
run = p.add_run('或手动操作：')
run.bold = True
doc.add_paragraph('打开命令行，执行：adb connect <电视IP>:5555', style='List Bullet')
doc.add_paragraph('首次连接电视会弹出授权提示，务必点击「允许」', style='List Bullet')
doc.add_paragraph('验证连接：adb devices，看到设备IP即成功', style='List Bullet')

doc.add_heading('步骤3：安装管理APP', level=2)
doc.add_paragraph('在管理后台找到对应设备，点击「一键部署上线」按钮，系统将自动完成以下操作：')

add_styled_table(
    ['操作', '说明', '预期结果'],
    [
        ['检查ADB连接', '验证电视是否可达且已授权', '显示「ADB已连接」'],
        ['安装APP', '推送APK到电视并安装', '显示「安装成功」'],
        ['卸载广告应用', '智能匹配并停用系统广告应用', '显示卸载数量'],
        ['设置默认桌面', '将管理APP设为默认HOME', '按HOME键进入管理界面'],
    ]
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('手动安装方式：')
run.bold = True
doc.add_paragraph('adb install -r <APK文件路径>', style='List Bullet')
doc.add_paragraph('例：adb install -r tv-launcher-1.0.apk', style='List Bullet')

doc.add_heading('步骤4：设置默认桌面', level=2)
doc.add_paragraph('一键部署会自动设置。如需手动操作：')
doc.add_paragraph('设置默认HOME：adb shell pm set-home-activity com.company.tvlauncher', style='List Bullet')
doc.add_paragraph('按电视遥控器HOME键，选择「TV Launcher」并设为始终', style='List Bullet')

doc.add_paragraph()

# ==================== 三、常见问题处理 ====================
doc.add_heading('三、常见问题处理', level=1)

add_styled_table(
    ['问题', '可能原因', '解决方案'],
    [
        ['ADB连接不上', '1. 电视未开启ADB\n2. 不在同一网络\n3. 电视未授权', '1. 检查开发者选项中ADB已开启\n2. 确认电脑和电视在同一网段\n3. 重新连接并在电视上点击「允许」'],
        ['安装失败 INSTALL_FAILED_USER_RESTRICTED', '电视安全设置阻止安装', '电视设置 → 安全 → 允许USB安装应用（不同型号路径不同）'],
        ['HOME键没反应', '未设置默认桌面', '手动设置：adb shell pm set-home-activity com.company.tvlauncher'],
        ['电视重启后APP没有自启动', '自启动权限被系统限制', '电视设置 → 应用 → TV Launcher → 开启自启动权限'],
        ['广告应用又出现了', '系统更新恢复了停用的应用', '重新停用：管理后台 → 应用管理 → 停用广告应用'],
        ['WiFi连接不上', '企业WiFi需要802.1X认证', '目前仅支持WPA2-PSK，企业WiFi需联系IT配置'],
    ],
    header_color="C0392B"
)

doc.add_paragraph()

# ==================== 四、系统广告应用清单 ====================
doc.add_heading('四、系统广告应用清单', level=1)
doc.add_paragraph('以下应用为小米电视系统预装广告应用，部署时建议停用：')

ad_apps = [
    ['com.miui.msa.global', '小米系统广告'],
    ['com.xiaomi.mipicks', '小米应用商店'],
    ['com.mitv.payment', '小米电视支付'],
    ['com.xiaomi.tweather', '小米天气'],
    ['com.xiaomi.gamecenter', '小米游戏中心'],
    ['com.miui.weather2', 'MIUI天气'],
    ['com.xiaomi.shop', '小米商城'],
]
add_styled_table(
    ['包名', '说明'],
    ad_apps,
    header_color="8E44AD"
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('注意：停用系统应用属于危险操作，请确认无误后再操作。停用后可通过恢复出厂设置恢复。')
run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
run.bold = True

doc.add_paragraph()

# ==================== 五、功能前置条件速查 ====================
doc.add_heading('五、功能前置条件速查', level=1)

add_styled_table(
    ['功能', '需要的前置条件', '如何开启'],
    [
        ['远程连接（scrcpy）', 'ADB已连接且已授权', '管理后台 → 设备详情 → 连接ADB'],
        ['应用管理（卸载/停用）', 'ADB已连接且已授权', '管理后台 → 设备详情 → 连接ADB'],
        ['一键部署上线', 'ADB已连接且已授权', '管理后台 → 设备详情 → 一键部署'],
        ['推送WiFi配置', 'APP已安装且在线', 'APP安装后自动连接后台'],
        ['策略执行', 'APP已安装且在线', 'APP安装后自动同步策略'],
        ['OTA升级', 'APP已安装且在线', '后台推送更新后APP自动下载'],
    ],
    header_color="27AE60"
)

doc.add_paragraph()

# ==================== 六、联系方式 ====================
doc.add_heading('六、联系方式', level=1)
doc.add_paragraph('遇到无法解决的问题，请联系IT部门负责人。')
doc.add_paragraph('管理后台地址部署在内网，请联系IT获取访问地址。')

doc.save(OUTPUT_PATH)
print(f"文档已生成: {OUTPUT_PATH}")
