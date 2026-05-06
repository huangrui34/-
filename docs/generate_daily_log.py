"""生成 TV Launcher 每日工作记录文档 (.docx)"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

OUT_DIR = os.path.dirname(os.path.abspath(__file__))
os.makedirs(OUT_DIR, exist_ok=True)

# ========== 色彩体系 ==========
COLOR_PRIMARY = "1B3A5C"
COLOR_SECONDARY = "2E75B6"
COLOR_ACCENT = "4A90D9"
COLOR_TEXT = "2D3436"
COLOR_TEXT_LIGHT = "636E72"
COLOR_BORDER = "D5DDE5"
COLOR_BG_GRAY = "F5F6F8"
COLOR_BG_GREEN = "E8F8F5"
COLOR_BG_RED = "FDEDEC"
COLOR_BG_YELLOW = "FEF9E7"
COLOR_BG_BLUE = "EBF5FB"

# ========== 通用样式工具 ==========

def set_cell_shading(cell, color_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_styled_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor.from_string(
            COLOR_PRIMARY if level == 1 else COLOR_SECONDARY if level == 2 else COLOR_ACCENT
        )
        run.font.bold = True
    return h

def add_body_text(doc, text, bold=False, color=None, indent=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(color or COLOR_TEXT)
    run.font.bold = bold
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(4)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if indent:
        pf.left_indent = Cm(1)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(COLOR_TEXT)
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    pf.left_indent = Cm(1 + level * 0.8)
    return p

def set_table_style(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(3)
                paragraph.paragraph_format.space_after = Pt(3)
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = "Microsoft YaHei"
            if row_idx == 0:
                set_cell_shading(cell, COLOR_SECONDARY)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor.from_string("FFFFFF")
                        run.font.bold = True

# ========== 工作记录数据 ==========

WORK_LOG = [
    {
        "date": "2026-05-06",
        "summary": "HDMI策略按HOME/BACK后立即恢复(绕过冷却期)+DPAD按键泄漏防护",
        "tasks": [
            {
                "type": "Bug修复",
                "title": "HDMI策略按HOME/BACK后不立即恢复",
                "description": "用户误按HOME/BACK后HDMI策略需要3-5秒才能恢复，冷却期阻止了立即恢复",
                "details": [
                    "MainActivity.onResume: HDMI模式下hdmi_foreground=false时立即恢复，绕过5秒防抖和30秒冷却期",
                    "forceExecutePolicy: HDMI恢复场景绕过防抖和冷却期检查",
                    "保活检查: hdmi_foreground=false时直接恢复，不受冷却期限制",
                    "KeepAliveForegroundService: HDMI模式检查hdmi_foreground并重新执行策略",
                    "关键思路: 冷却期防止策略切换时重复启动，但用户误离开时必须立即恢复",
                ],
                "status": "已完成",
                "files": ["MainActivity.kt", "KeepAliveForegroundService.kt"]
            },
            {
                "type": "Bug修复",
                "title": "DPAD按键泄漏到其他Activity",
                "description": "HdmiActivity的scheduleConfirmDialog发送的DPAD按键可能泄漏到MainActivity点击设置按钮",
                "details": [
                    "scheduleConfirmDialog添加hasFocus检查",
                    "发送DPAD_CENTER/DPAD_RIGHT前检查hasFocus，失去焦点时跳过",
                    "防止按键在Activity切换时泄漏到其他Activity",
                ],
                "status": "已完成",
                "files": ["HdmiActivity.kt"]
            },
        ]
    },
    {
        "date": "2026-05-03",
        "summary": "HDMI暂停按键误触修复、策略切换闪烁修复、HDMI拔插黑屏修复、离线设备显示优化、管理后台按钮loading状态",
        "tasks": [
            {
                "type": "Bug修复",
                "title": "HDMI策略暂停后按键误触设置和密码键盘",
                "description": "安卓6电视上暂停HDMI策略后，系统自动点击设置按钮和密码键盘。Runtime.exec按键注入无法取消导致",
                "details": [
                    "三层防护: isLeaving标志阻止handler待发按键, dispatchKeyEvent拦截已exec按键, bringLauncherToFront不再启动HdmiActivity",
                    "HdmiActivity.onPause: 设置isLeaving=true, 清理handler队列",
                    "MainActivity: 策略暂停后5秒内拦截DPAD按键(HDMI_PAUSE_KEY_SUPPRESS_MS)",
                    "KeepAliveForegroundService: 策略暂停时只更新通知，不拉回主页",
                ],
                "status": "已完成",
                "files": ["HdmiActivity.kt", "MainActivity.kt", "KeepAliveForegroundService.kt"]
            },
            {
                "type": "Bug修复",
                "title": "HDMI策略切换(HDMI1→HDMI2)屏幕闪烁黑屏",
                "description": "安卓6电视切换HDMI端口时闪烁2-3次最终停留在旧端口，onNewIntent用inputId比较端口失效",
                "details": [
                    "onNewIntent改用端口号比较: 直接比较EXTRA_HDMI_PORT，不再从inputId提取端口",
                    "安卓6 Amlogic只有1个HDMI InputService，所有端口同一inputId，inputId比较失效",
                    "forceExecutePolicy检查hdmi_foreground: HdmiActivity已在前台时发送onNewIntent复用实例",
                    "LauncherExecutor添加FLAG_ACTIVITY_CLEAR_TOP|FLAG_ACTIVITY_SINGLE_TOP",
                ],
                "status": "已完成",
                "files": ["HdmiActivity.kt", "MainActivity.kt", "LauncherExecutor.kt"]
            },
            {
                "type": "Bug修复",
                "title": "HDMI线拔插后黑屏不恢复",
                "description": "安卓6电视拔插HDMI线后画面黑屏，BroadcastReceiver不触发，改用TvInputCallback",
                "details": [
                    "TvInputManager.TvInputCallback替代HDMI_PLUGGED BroadcastReceiver",
                    "onInputAdded: HDMI输入添加时根据targetPort重新解析inputId并调谐",
                    "onInputRemoved: 当前HDMI输入移除时重置hasTuned和currentInputId",
                    "onInputStateChanged: HDMI信号恢复(state=0)时调度调谐",
                    "添加targetPort变量跟踪策略指定端口",
                    "安卓6上HDMI_PLUGGED广播从不触发，TvInputCallback可靠工作",
                ],
                "status": "已完成",
                "files": ["HdmiActivity.kt"]
            },
            {
                "type": "功能优化",
                "title": "离线设备管理页面显示优化",
                "description": "离线设备隐藏信号质量(良好/5GHz/Mbps)和延迟/丢包数据，保留IP/MAC地址",
                "details": [
                    "后端: 离线设备清除wifi_rssi/frequency/link_speed/latency/loss字段",
                    "前端: renderWifiQuality对离线设备返回空字符串",
                    "删除WiFi推送功能(WifiConfigManager/相关API/弹窗)",
                ],
                "status": "已完成",
                "files": ["main.py", "schemas.py", "models.py", "index.html"]
            },
            {
                "type": "功能优化",
                "title": "管理后台更新/暂停按钮添加loading状态",
                "description": "按钮点击后无反馈导致重复点击，添加禁用+状态文字防止重复操作",
                "details": [
                    "更新按钮: 点击后显示\"下发中...\"并禁用",
                    "暂停/继续按钮: 点击后显示\"暂停中...\"/\"继续中...\"并禁用",
                    "失败时恢复按钮原文字，避免状态错乱",
                ],
                "status": "已完成",
                "files": ["index.html"]
            },
        ]
    },
    {
        "date": "2026-04-30",
        "summary": "设置页面焦点高亮、WiFi企业级自动识别、一键部署SSE进度流、WiFi页面焦点高亮",
        "tasks": [
            {
                "type": "功能开发",
                "title": "SettingsActivity焦点高亮：遥控器操作视觉反馈",
                "description": "为SettingsActivity所有可交互项添加OnFocusChangeListener，焦点获得时显示浅蓝背景+蓝色边框，失去焦点时恢复白色背景+灰色边框",
                "details": [
                    "焦点样式: 浅蓝背景(#E3F2FD) + 蓝色边框(#2196F3) + 2dp描边",
                    "普通样式: 白色背景(#FFFFFF) + 灰色边框(#E0E0E0)",
                    "使用GradientDrawable动态切换背景，roundedBg()工具方法",
                    "覆盖所有菜单项、返回按钮、编辑框等可交互元素",
                ],
                "status": "已完成",
                "files": ["SettingsActivity.kt"]
            },
            {
                "type": "功能开发",
                "title": "WifiConnectActivity焦点高亮及企业级WiFi自动识别",
                "description": "WiFi连接页面添加焦点高亮，同时自动识别企业级WiFi(802.1X/EAP)并显示对应的输入框",
                "details": [
                    "焦点高亮: 与SettingsActivity相同的样式，所有WiFi项、按钮、开关均有焦点反馈",
                    "企业级检测: 检查scanResult.capabilities是否包含'EAP'字段",
                    "企业级标签: WiFi列表中为企业级WiFi显示紫色「企业级」标签",
                    "企业级连接: 点击企业级WiFi时弹出用户名(Identity)+密码输入框，使用WifiEnterpriseConfig PEAP/MSCHAPv2",
                    "普通WiFi: 仍只显示密码输入框，使用WifiConfiguration.preSharedKey",
                    "返回按钮: 添加←返回按钮，方便遥控器操作",
                    "Toast替代AlertDialog: 以太网警告等改用Toast，避免弹窗焦点问题",
                ],
                "status": "已完成",
                "files": ["WifiConnectActivity.kt"]
            },
            {
                "type": "功能开发",
                "title": "一键部署SSE流式进度显示",
                "description": "将一键部署API从同步响应改为SSE流式推送，前端实时显示每个步骤的进度",
                "details": [
                    "后端新增 /api/v1/deploy-tv-stream 端点，text/event-stream格式",
                    "步骤消息: adb_connect, adb_auth, adb_ok, install_apk, install_ok, configure, launch, wait_register, done",
                    "ADB授权重试: 3次重试，每次5秒间隔，推送「等待电视授权」消息",
                    "前端使用EventSource消费SSE流，实时更新步骤指示器和运行日志",
                ],
                "status": "已完成",
                "files": ["main.py", "index.html"]
            },
            {
                "type": "部署",
                "title": "新版APP部署到测试电视",
                "description": "包含焦点高亮和企业级WiFi识别的新APK部署",
                "details": [
                    "10.181.34.29 - APK安装成功",
                    "10.181.185.247 - APK安装成功",
                ],
                "status": "已完成",
                "files": []
            },
        ]
    },
    {
        "date": "2026-04-29",
        "summary": "811电视开机白屏修复、管理后台下拉框Bug修复、HDMI弹窗确认自动误操作修复、WiFi SSID显示修复",
        "tasks": [
            {
                "type": "Bug修复",
                "title": "811安卓6电视开机白屏闪烁优化",
                "description": "811会议室小米电视(Amlogic, 安卓6)开机后白屏闪黑好几秒，用户体验极差",
                "details": [
                    "根因1: BootReceiver延迟3秒启动，期间系统显示默认Launcher白色界面",
                    "根因2: MaterialComponents主题在Android 6冷启动初始化慢，短暂显示白色预览窗口",
                    "根因3: Activity启动与策略执行耦合在同一延迟中",
                    "修复: 新增Theme.TvLauncher.Splash轻量启动主题，继承android:Theme.NoTitleBar.Fullscreen",
                    "修复: Manifest中MainActivity使用Splash主题，onCreate中setTheme切回完整主题",
                    "修复: BootReceiver去掉3秒延迟，立即启动Activity，策略执行单独延迟1.5秒",
                    "修复: 补充ACTION_REBOOT广播处理",
                    "效果: 开机预览窗口直接显示深色(#0A0E14)，消除白屏闪烁",
                ],
                "status": "已完成",
                "files": ["themes.xml", "AndroidManifest.xml", "MainActivity.kt", "BootReceiver.kt"]
            },
            {
                "type": "Bug修复",
                "title": "管理后台策略下拉框选择时消失、选择值自动回退",
                "description": "管理后台执行策略栏的下拉框在选择时消失，选择后还没点下发就自动回退到原值",
                "details": [
                    "根因: setInterval(loadData, 5000)每5秒用innerHTML重建整个设备表DOM",
                    "问题1: 打开的select被重建的DOM销毁，下拉框瞬间消失",
                    "问题2: 用户已选但未保存的值被服务器旧数据覆盖，选择回退",
                    "修复: 添加isSelectActive()检查，select获焦时跳过刷新",
                    "修复: 添加pendingSelections字典，记录用户已修改但未保存的选择",
                    "修复: loadData重建select时优先使用pendingSelections中的本地值",
                    "修复: bindPolicy下发成功后清除对应的pendingSelections",
                ],
                "status": "已完成",
                "files": ["index.html"]
            },
            {
                "type": "Bug修复",
                "title": "HDMI弹窗确认代码重复发送DPAD_CENTER导致自动误操作",
                "description": "安卓6电视(10.181.184.254)自动点击设置按钮和密码输入框，scheduleConfirmDialog()每轮检查都发DPAD_CENTER",
                "details": [
                    "根因1: scheduleConfirmDialog()在每次!hasFocus时都发送DPAD_CENTER，最多8次",
                    "根因2: 窗口持续失焦时重复发送确认键，导致误点击UI元素",
                    "根因3: onDestroy未取消confirmDialogRunnable，Activity销毁后仍在发按键",
                    "修复: 改为检测焦点变化(有→无)的瞬间才发确认键，只发一次",
                    "修复: 用hasConfirmedDialog标志位防止重复发送",
                    "修复: onDestroy中同时取消confirmDialogRunnable",
                ],
                "status": "已完成",
                "files": ["HdmiActivity.kt"]
            },
            {
                "type": "功能修复",
                "title": "WiFi SSID显示优化：显示WiFi名称而非'WiFi'文字",
                "description": "管理后台在线状态列，WiFi设备只显示'WiFi'而非具体WiFi名称",
                "details": [
                    "根因: Android 6+读取WiFi SSID需要ACCESS_COARSE_LOCATION权限，APP未声明",
                    "修复: AndroidManifest.xml添加ACCESS_COARSE_LOCATION和ACCESS_FINE_LOCATION",
                    "修复: MainActivity.onCreate中请求运行时位置权限",
                    "修复: 已部署设备通过adb shell pm grant授权",
                    "说明: 前端getNetLabel()已有SSID显示逻辑，只是数据源缺少SSID值",
                ],
                "status": "已完成",
                "files": ["AndroidManifest.xml", "MainActivity.kt"]
            },
            {
                "type": "部署",
                "title": "新版本APP(v0.2.0)部署到电视设备",
                "description": "包含白屏修复、弹窗确认修复、位置权限的新APK部署",
                "details": [
                    "10.181.184.254 (MiTV4-ANSM0, Android 6) - APK已安装，但ADB连接断开待恢复",
                    "10.181.184.226 (MiTV-ASTP0, Android 9) - APK已安装，位置权限已授权",
                ],
                "status": "进行中",
                "files": []
            },
        ]
    },
    {
        "date": "2026-04-28",
        "summary": "HDMI TvView API实现、810电视服务器URL修复、APP部署",
        "tasks": [
            {
                "type": "功能开发",
                "title": "HdmiActivity TvView API实现",
                "description": "创建HdmiActivity，使用Android TV Input Framework的TvView API直接显示HDMI输入信号",
                "details": [
                    "支持Amlogic/Droidlogic芯片: com.droidlogic.tvinput/.services.Hdmi1InputService/HW5",
                    "支持MediaTek芯片: com.mediatek.tvinput/.hdmi.HDMIInputService/HW5",
                    "支持MStar芯片: com.mstar.tvinput/.service.Hdmi1InputService/HW5",
                    "支持Realtek芯片: com.realtek.tvinput/.services.Hdmi1InputService/HW5",
                    "5级匹配策略: 已知ID → HW编号匹配 → 名称模糊匹配 → 单HDMI回退 → 构造ID尝试",
                    "HDMI1/2/3端口均测试通过(811会议室Amlogic Android 6)",
                    "测试电视(10.181.184.226, Amlogic Android 9) HDMI2切换测试通过",
                ],
                "status": "已完成",
                "files": ["HdmiActivity.kt", "LauncherExecutor.kt", "AndroidManifest.xml", "themes.xml"]
            },
            {
                "type": "Bug修复",
                "title": "810电视服务器URL错误导致策略不生效",
                "description": "810会议室电视(MiTV-MFTP0, MediaTek, Android 9)服务器URL配置为http://10.181.99.23:8000，但实际服务器在http://10.181.153.84:8000",
                "details": [
                    "通过run-as修改SharedPreferences中的server_base_url",
                    "同时通过settings put global设置全局配置作为备份",
                    "重启APP后心跳成功，策略切换生效",
                ],
                "status": "已完成",
                "files": []
            },
            {
                "type": "Bug修复",
                "title": "RemoteApi心跳异常导致APP崩溃",
                "description": "当服务器不可达时，OkHttp抛出SocketTimeoutException，未被捕获导致APP崩溃",
                "details": [
                    "添加try-catch包裹heartbeat网络调用",
                    "捕获Exception后记录日志并返回false，不崩溃",
                    "影响所有电视设备，尤其是网络不稳定时",
                ],
                "status": "已完成",
                "files": ["RemoteApi.kt"]
            },
            {
                "type": "部署",
                "title": "新版本APP部署到所有电视",
                "description": "构建包含HdmiActivity和心跳修复的新APK，部署到测试电视和810电视",
                "details": [
                    "10.181.184.226 (MiTV-ASTP0, Android 9) - 旧版APP没有HdmiActivity，安装后HDMI切换正常",
                    "10.181.185.247 (MiTV-MFTP0, Android 9) - 修复服务器URL + 安装新版APP",
                ],
                "status": "已完成",
                "files": []
            },
        ]
    },
    {
        "date": "2026-04-27",
        "summary": "HDMI自动切换策略功能实现",
        "tasks": [
            {
                "type": "功能开发",
                "title": "HDMI热插拔自动切换策略",
                "description": "检测HDMI信号接入/断开，自动切换策略",
                "details": [
                    "新增HdmiReceiver监听android.intent.action.HDMI_PLUGGED广播",
                    "HDMI接入时保存当前策略并切换到HDMI1",
                    "HDMI断开时恢复之前的策略",
                    "PolicyStore新增pre_hdmi_policy相关存储方法",
                    "支持HDMI自动切换开关(hdmi_auto_switch_enabled)",
                ],
                "status": "已完成",
                "files": ["HdmiReceiver.kt", "MainActivity.kt", "PolicyStore.kt", "AndroidManifest.xml"]
            },
        ]
    },
]

# ========== 生成文档 ==========

def generate():
    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # 标题
    title = doc.add_heading("TV Launcher 每日工作记录", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor.from_string(COLOR_PRIMARY)
        run.font.size = Pt(26)

    # 副标题
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Meeting TV Launcher 项目开发日志")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor.from_string(COLOR_TEXT_LIGHT)

    doc.add_paragraph()  # 间距

    # 目录说明
    add_body_text(doc, "本文档记录TV Launcher项目每日的开发工作、问题修复和部署变更。", color=COLOR_TEXT_LIGHT)

    # 按日期倒序排列
    for day in sorted(WORK_LOG, key=lambda x: x["date"], reverse=True):
        # 日期标题
        add_styled_heading(doc, f"📅 {day['date']}", level=1)
        add_body_text(doc, f"今日概要: {day['summary']}", bold=True, color=COLOR_SECONDARY)

        for idx, task in enumerate(day["tasks"], 1):
            # 任务标题
            type_color = {
                "功能开发": COLOR_BG_GREEN,
                "Bug修复": COLOR_BG_RED,
                "部署": COLOR_BG_BLUE,
                "研究": COLOR_BG_YELLOW,
            }.get(task["type"], COLOR_BG_GRAY)

            status_color = {
                "已完成": "27AE60",
                "进行中": "F39C12",
                "未开始": COLOR_TEXT_LIGHT,
            }.get(task["status"], COLOR_TEXT_LIGHT)

            add_styled_heading(doc, f"{idx}. [{task['type']}] {task['title']}", level=2)

            # 状态
            p = doc.add_paragraph()
            run = p.add_run(f"状态: {task['status']}")
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor.from_string(status_color)
            run.font.bold = True

            # 描述
            add_body_text(doc, task["description"])

            # 详细条目
            if task.get("details"):
                add_body_text(doc, "详细:", bold=True, color=COLOR_SECONDARY)
                for detail in task["details"]:
                    add_bullet(doc, detail)

            # 涉及文件
            if task.get("files"):
                add_body_text(doc, "涉及文件: " + ", ".join(task["files"]), color=COLOR_TEXT_LIGHT, indent=True)

        doc.add_paragraph()  # 日期间距

    # 保存
    out_path = os.path.join(OUT_DIR, "05-每日工作记录.docx")
    doc.save(out_path)
    print(f"文档已生成: {out_path}")

if __name__ == "__main__":
    generate()
